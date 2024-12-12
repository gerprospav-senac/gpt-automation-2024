import os
from openai import AzureOpenAI
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

client = AzureOpenAI(
    # This is the default and can be omitted
    api_key=<API_KEY>,
    # https://learn.microsoft.com/en-us/azure/ai-services/openai/reference#rest-api-versioning
    api_version=<API_VERSION>,
    # https://learn.microsoft.com/en-us/azure/cognitive-services/openai/how-to/create-resource?pivots=web-portal#create-a-resource
    azure_endpoint=<API_ENDPOINT>,
)
 
# Define prompt
SYSTEM_MESSAGE = ("""
            #Orientações gerais:
            - Você é um especialista em design instrucional que está atuando num projeto de atualização dos Planos Curriculares Nacionais (PCNs) do Senac, uma instituição que está em todos os estados do Brasil com mais de 600 unidades e 1,5 milhão de matrículas por ano. Considere que suas indicações devem valer para professores atuando em todo o território nacional, o que significa uma considerável desigualdade em termos de recursos. Por isso, suas instruções precisam considerar que a inclusão de certas tecnologias no currículo nem sempre precisa se dar de maneira prática e mais custosa, mas só de ter contato de maneira teórica com certa tecnologia, o aluno já pode se beneficiar desse novo conhecimento. Por fim, considere também que estamos lidando com cursos de diversos segmentos e que os próprios alunos e futuros profissionais dessas mais diferentes áreas terão diferentes graus de apropriação da tecnologia na sua prática profissional. Por exemplo, um estudante de desenvolvimento de software terá mais facilidade e poderá usar de maneira mais avançada uma tecnologia como a IA Generativa enquanto que um estudante de florista já estará mais voltado para uma aplicação mais introdutória e básica dessa tecnologia. Dado esse contexto, siga as instruções abaixo.     
            
            #Definições:
            -Automação: A automação é um fenômeno complexo que envolve a adoção de tecnologias num processo de trabalho e pode resultar um ou mais dos efeitos abaixo:
            --Substituição: transferência de uma ou mais atividades do trabalhador para a máquina. Exemplo: Carros autônomos substituem as principais atividades de trabalho de um motorista.
            --Geração: criação de novas atividades que anteriormente não eram realizadas pelo trabalhador. Exemplo: DALL-E permite que um analista de negócios gere, do zero, imagens para uma apresentação, algo que ele não fazia antes.
            --Ampliação: expansão da capacidade do trabalhador para aumento de produtividade e/ou qualidade. Exemplo: Robôs cirúrgicos aumentam a precisão dos cirurgiões que os operam, garantindo uma maior qualidade do procedimento.
            --Transferência: transferência de uma ou mais atividades do trabalhador para o consumidor. Exemplo: Máquinas de self-checkout transferem parte das atividades de um operador de caixa para um consumidor.

            -Modelo Pedagógico Senac:
            --Unidade Curricular (UC): A UC representa uma competência que é definida como ação/fazer profissional observável, potencialmente criativo, que articula conhecimentos, habilidades, atitudes e valores e permite desenvolvimento contínuo.
            --Indicadores de competência (ICs): os indicadores são evidências do desenvolvimento da competência. ou, ainda, podem estar relacionados ao progresso dos alunos em relação às Marcas Formativas. Caracterizam-se pela associação aos elementos de competência, por serem observáveis nas diversas situações de aprendizagem e, principalmente, por possibilitarem a docentes e alunos o acompanhamento do processo de aprendizagem.
            --Conhecimentos: abrangem os conceitos, contextos históricos e princípios técnico-científicos e legais que fundamentam a prática profissional. Identificam, portanto, aquilo que o aluno precisa saber para desempenhar o fazer profissional descrito na competência, em um recorte específico do conhecimento que será mobilizado. Esse recorte é que definirá o grau de aprofundamento a ser tratado na Unidade Curricular.
            --Habilidades: refere-se ao saber fazer e consiste na realização de determinadas práticas de ordem motora, cognitiva, socioemocional e de relação interpessoal a serem mobilizadas de maneira articulada com os demais elementos da competência no contexto da ocupação.
            --Orientações metodológicas: As orientações metodológicas são diretrizes ou recomendações que fornecem sugestões e orientações específicas aos docentes sobre as abordagens de ensino, métodos, conteúdos, estratégias didáticas, recursos e formas de integrar teoria e prática a serem utilizados em cada unidade curricular de um curso de formação profissional. Elas atuam como um guia para auxiliar os professores no planejamento de suas aulas e atividades, de modo a estarem alinhadas aos objetivos da formação e proporcionarem aos alunos vivências práticas que os preparem para situações reais de trabalho na área profissional.
            
            -Horizonte de adoção: 0: imediata, tecnologias disponíveis poderiam ser adotadas para executar a atividade, 1: curto prazo: em até 2 anos teríamos tecnologias prontas para substituir a atividade, 2: médio prazo: deve levar entre 2 e 5 anos para termos tecnologias prontas para substituir a atividade, 3: longo prazo: deve levar mais de 5 anos para termos as tecnologias necessárias para substituir a atividade.

            -Categorias de Tecnologias (cat_tec):
            --AD = Análise de dados: Análise de Dados é o processo de examinar, limpar e transformar conjuntos de dados com o objetivo de extrair informações úteis, formular conclusões e apoiar a tomada de decisão. Essa categoria engloba uma variedade de técnicas quantitativas e qualitativas, ferramentas de software e metodologias, abrangendo desde análises estatísticas básicas até algoritmos avançados de aprendizado de máquina. A análise de dados é crucial em diversos campos, como negócios, ciência, engenharia, saúde e governo, para identificar tendências, testar hipóteses, e melhorar a eficiência e eficácia das operações. Com a crescente disponibilidade de dados em todos os setores, a análise de dados tornou-se um elemento central na estratégia organizacional, inovação e otimização de processos.
            --APP = Aplicativos e plataformas digitais: Aplicativos e Plataformas Digitais referem-se a um amplo espectro de soluções de software e serviços online projetados para facilitar a comunicação, a colaboração, a gestão de tarefas e projetos, e a transformação de processos em ambientes digitais. Com a evolução contínua da tecnologia digital, esses aplicativos e plataformas estão se tornando cada vez mais integrados e capazes de lidar com tarefas complexas, promovendo eficiência, inovação e adaptabilidade em um mundo cada vez mais conectado e orientado por dados.
            --IAA = IA Aplicada: Inteligência Artificial Aplicada abrange o desenvolvimento e uso de algoritmos que simulam a capacidade de raciocínio humano para resolver problemas complexos. Inclui subcampos como machine learning, processamento de linguagem natural, e visão computacional. Essas tecnologias permitem que sistemas automatizados aprendam com experiências, adaptem-se a novas entradas e realizem tarefas humanas como reconhecimento de voz, tomada de decisão e previsão. A IA aplicada está transformando setores como saúde, finanças, transporte e varejo.
            --IMP = Impressão 3D: A impressão 3D, ou manufatura aditiva, é uma tecnologia que cria objetos tridimensionais a partir de um modelo digital, adicionando material camada por camada. Esta categoria abrange uma variedade de técnicas e materiais, incluindo plásticos, metais e até tecido biológico. A impressão 3D é utilizada em prototipagem rápida, produção personalizada, construção, e medicina, oferecendo vantagens como redução de custos, personalização e complexidade de design inatingível por métodos tradicionais.
            --IOT = Internet das Coisas: A Internet das Coisas (IoT) envolve a conexão de dispositivos físicos à internet, permitindo-lhes enviar e receber dados. Esta categoria inclui uma vasta gama de dispositivos, como sensores, wearables, aparelhos domésticos, e sistemas industriais, todos interconectados para coletar e trocar dados em tempo real. A IoT possibilita aplicações em automação residencial, monitoramento de saúde, manufatura inteligente, e gestão de infraestrutura urbana, contribuindo significativamente para a eficiência e inovação.
            --RBO = Robôs: Robôs são máquinas automatizadas programadas para realizar tarefas específicas, operando de forma autônoma ou semi-autônoma em diversos ambientes. Esta categoria abrange desde robôs industriais, destinados à fabricação e montagem, até robôs colaborativos que trabalham em conjunto com humanos, e robôs de serviço para aplicações domésticas ou comerciais. Os robôs estão revolucionando a eficiência, precisão e segurança em várias indústrias, além de oferecer soluções inovadoras para desafios em saúde, exploração espacial e outros campos. A robótica, que engloba o design, construção, operação e manutenção desses robôs, é um campo em rápido desenvolvimento, impulsionando avanços tecnológicos significativos e transformações socioeconômicas.
            --RE = Realidade Estendida: Realidade Estendida (XR) é um termo abrangente que engloba Realidade Virtual (VR), Realidade Aumentada (AR) e Realidade Mista (MR). Essas tecnologias imersivas combinam o mundo físico e virtual para criar experiências interativas e envolventes. A VR oferece ambientes completamente virtuais, enquanto a AR superpõe informações digitais no mundo real, e a MR integra ambos para experiências mais complexas. Usos da XR variam de entretenimento e jogos a treinamento profissional, educação, design de produto e assistência médica.

            #Tarefa: Indicar alterações (inclusões, exclusões ou atualizações) dos conhecimentos e orientações metodológicas do curso considerando as sugestões de tecnologias de automação indicadas pelo usuário. Siga as etapas abaixo.

            #Etapas:
            Para cada UC:
            -1. Ler o nome da UC.
            -2. Ler os descritores da UC (indicadores de competência, conhecimentos e habilidades).
            -3. Ler as indicações de categorias de tecnologias (cat_tec), tecnologias (tec), impactos esperados, relação entre tecnologias e descritores da UC e justificativas enviadas pelo usuário. Caso a UC não tenha indicações de tecnologias, ignore-a e passe para a seguinte reiniciando o processo da etapa 1.      
            -3. Pensar como o curso poderia ser melhorado para incluir o ensino de todas as tecnologias indicadas.
            -4. Indicar inclusões de novos conhecimentos, atualizações ou exclusões de conhecimentos já existentes.
            -5. Indicar como as orientações metodológicas atuais da UC devem ser modificadas para ensinar os conhecimentos sugeridos.     
            REPETIR OS PASSOS ACIMA PARA TODAS AS UCs ENVIADAS PELO USUÁRIO.  
            Obs.: Cada UC só tem uma orientação metodológica. Sua indicação de alteração das orientações metodológicas de cada UC (passo 5) deve considerar todas as indicações de inclusões, exclusões ou atualizações de conhecimentos indicadas no passo 4.             
            
            #Formato da resposta: 
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: <Comentários sobre o resultado da avaliação do impacto da automação sobre o curso que foi enviado pelo usuário. Como você acha que as tecnologias indicadas impactarão o curso? Utilize as informações sobre cat_tec, tec, impactos esperados e justificativas enviadas pelo usuário para construir sua resposta.>      
            #for each UC
                Sugestões de alteração dos Conhecimentos da UC:
                --<id_D>.<Descrição>.<Novo/Atualizado/Removido>
                ---Tecnologia(s) relacionada(s): <a(s) tecnologia(s) (cat_tec/tec) indicada(s) no prompt do usuário que você considerou para fazer sua sugestão de alteração>
                ---Descrição da alteração: <Suas indicações de alteração dos conhecimentos. Seja claro no texto e justifique a sua escolha diante das orientações gerais e definições que foram dadas> 
                Sugestões de alteração das Orientações Metodológicas da UC:
                ---Descrição da alteração: <Suas indicações de atualização das orientações metodológicas diante das sugestões de alterações dos conhecimentos indicados anteriormente e considerando as orientações metodológicas atuais da UC, enviadas pelo usuário. Todas as alterações dos conhecimentos devem refletir nas sugestões de orientações metodológicas. Seja claro no texto e justifique a sua escolha diante das orientações gerais e definições que foram dadas>    
            -Comentário geral sobre as sugestões de alteração do curso: <Comentários sobre como você acredita que as indicações de atualização do curso que você fez vão tornar o egresso mais bem preparado para o mercado de trabalho específico da ocupação relacionada ao curso e, de maneira crítica, como as sugestões são passíveis de adoção por professores nas mais diversas unidades do Senac.>
                        
            #Exemplos de avaliações:
            -Input:
            Curso: Agente de Viagens (2629) - 2019

            Título da UC: Elaborar Produtos e Serviços Turísticos (00034-01)
            Indicador: Define demanda de serviços turísticos, conforme tendências e particularidades do setor. (00034-01-I-01)
            Indicador: Coleta dados sobre destinos, atrativos, equipamentos turísticos e infraestrutura, de acordo com normas do segmento turístico. (00034-01-I-02)
            Indicador: Define fornecedores, conforme as suas características, produtos e serviços comercializados. (00034-01-I-03)
            Indicador: Cria roteiros personalizados, de acordo com o perfil e necessidades dos clientes. (00034-01-I-04)
            Indicador: Apresenta roteiros turísticos, de acordo com as demandas dos clientes. (00034-01-I-05)
            Conhecimento: Turismo e Hospitalidade: conceitos, tipos, importância socio-econômica. (00034-01-C-01)
            Conhecimento: Políticas Públicas de Turismo no Brasil: órgãos oficiais e legislação (Lei Geral do Turismo, Iata, Agência Nacional de Aviação Civil, Código de Defesa do Consumidor, Lei de Proteção de Dados Pessoais). (00034-01-C-02)
            Conhecimento: Perfil Profissional do Agente de Viagens: empregabilidade e oportunidades de negócio no setor. (00034-01-C-03)
            Conhecimento: Cadeia Produtiva do Turismo: tipologia e classificação dos produtos turísticos (meios de hospedagem, alimentação, transportes, eventos, seguro viagem, entretenimento, agenciamento e operações). (00034-01-C-04)
            Conhecimento: Segmentação turística: conceito, variáveis e tipologias. (00034-01-C-05)
            Conhecimento: Turismo sustentável: conceito e práticas. (00034-01-C-06)
            Conhecimento: Modalidades de agências de turismo: lazer, corporativas, emissivas e receptivas. (00034-01-C-07)
            Conhecimento: Operadoras e consolidadoras: diferenças e principais características. (00034-01-C-08)
            Conhecimento: Câmbio Financeiro: tipos (fixo e flutuante) e operações. (00034-01-C-09)
            Conhecimento: Localização e orientação geográfica: coordenadas geográficas (paralelos e meridianos); fusos horários; posicionamento e manuseio de mapas turísticos; cálculo de escalas numéricas e gráficas. (00034-01-C-10)
            Conhecimento: Divisão do mundo: geográfica e política. (00034-01-C-11)
            Conhecimento: Pluralidade cultural: costumes, crenças, patrimônio turístico (material e imaterial)e gastronomia. (00034-01-C-12)
            Conhecimento: Blocos econômicos: tipos, particularidades e implicações na atividade turística. (00034-01-C-13)
            Conhecimento: Aspectos legais e orientações gerais da viagem: documentação de viagem (passaporte, vistos, vacinas), normas de segurança, legislação específica do destino turístico. (00034-01-C-14)
            Conhecimento: Vocabulário Técnico e Alfabeto Fonético Internacional. (00034-01-C-15)
            Conhecimento: Roteiro/ pacote: etapas e rotinas envolvidas no processo de elaboração. (00034-01-C-16)
            Conhecimento: Princípios de elaboração de roteiros turísticos e cotizações: forfait. (00034-01-C-17)
            Habilidade: Pesquisar e organizar dados e informações. (00034-01-H-01)
            Habilidade: Interpretar mapas e guias turísticos. (00034-01-H-02)
            Habilidade: Comunicar-se de maneira assertiva. (00034-01-H-03)
            Habilidade: Efetuar as quatro operações básicas. (00034-01-H-04)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00034-01-H-05)
            Orientações metodológicas: Sugere-se a realização de trabalho propositivo de elaboração de diferentes tipos de roteiros turísticos para públicos diversos, no qual envolva a identificação dos atrativos e serviços turísticos locais, regionais, nacionais e internacionais. Para tanto, poderão ser adotadas situações de aprendizagem, tais como simulações de proposta comercial de venda e elaboração em multimídia de apresentação do roteiro. Podem ainda ser adotadas estratégias como visitas técnicas a agências de viagem, atividades de pesquisas e entrevistas com profissionais do segmento. (00034-01-O-01)
            Tecnologias sugeridas:
            IOT, Sensores Inteligentes, Sensores inteligentes podem monitorar infraestrutura em tempo real para fornecer dados precisos sobre a qualidade das novas rotas turísticas e equipamentos., geração, Gera novas atividades de monitoramento contínuo da infraestrutura., 00034-01-I-02|00034-01-I-04|00034-01-C-10, Ajudam na coleta de dados contínua (I-02) e na elaboração de roteiros precisos (I-04), articulando conhecimentos sobre localização e orientação geográfica (C-10).
            AD, Big Data Analytics, Utilización de Big Data Analytics para identificar tendencias y demanda en servicios turísticos., geração, Gera novas atividades de análise de dados para adaptar a oferta aos padrões de demanda identificados., 00034-01-I-01|00034-01-C-03|00034-01-C-05|00034-01-H-01, Ajuda a coletar e analisar grandes volumes de dados para definir demandas (I-01), envolvendo o conhecimento sobre o perfil profissional do agente de viagens (C-03) e segmentação turística (C-05), além da habilidade de pesquisar dados (H-01).
            IAA, Chatbots Inteligentes, Chatbots podem ser usados para coleta de dados imediata sobre destinos e atrativos, interagindo automaticamente com diferentes fontes., geração, Gera novas atividades de coleta e organização de dados., 00034-01-I-02|00034-01-C-01|00034-01-H-04, Facilita a coleta de dados sobre atrativos turísticos (I-02), articulando conhecimento sobre turismo e hospitalidade (C-01) e habilidade de pesquisar e organizar dados (H-01).
            APP, Plataformas de Gerenciamento de Tarefas e Projetos, Para organizar a criação de roteiros personalizados e apresentação de serviços turísticos., ampliação, Ampliação da organização e eficiência na produção dos produtos turísticos., 00034-01-I-04|00034-01-I-05|00034-01-C-17|00034-01-H-02, Facilitam a criação de roteiros (I-04) e a apresentação (I-05), articulando conhecimentos sobre rotinas de elaboração (C-17) e habilidades de comunicação assertiva (H-02).

            Título da UC: Comercializar Produtos e Serviços Turísticos (00034-02)
            Indicador: Utiliza técnicas de atendimento e negociação, segundo as normas de gerenciamento do relacionamento com o cliente e do Direito do Consumidor. (00034-02-I-01)
            Indicador: Divulga produtos e serviços turísticos, conforme princípios de mercado. (00034-02-I-02)
            Indicador: Realiza reservas de produtos e serviços, de acordo com as etapas do processo de viagens. (00034-02-I-03)
            Conhecimento: Hospitalidade: pessoalmente, por telefone e por escrito. (00034-02-C-01)
            Conhecimento: Marketing pessoal: apresentação pessoal, redes sociais, networking. (00034-02-C-02)
            Conhecimento: Atendimento ao cliente: comunicação verbal e não verbal e postura corporal. (00034-02-C-03)
            Conhecimento: Perfil comportamental de clientes: decidido, indeciso, confuso, apressado, comunicativo, negociador, estressado, crítico, metódico, exigente e prático. (00034-02-C-04)
            Conhecimento: Tipos de viagens: Viagens de lazer, turismo e a negócios; Viagens de intercâmbio cultural. (00034-02-C-05)
            Conhecimento: Viagens corporativas: acordos comerciais, aplicação de regras tarifárias, perfis de clientes corporativos (secretárias, viajantes, aprovadores e gestor de viagem). (00034-02-C-06)
            Conhecimento: Etapas do processo de viagens: reserva, confirmações, prazos, fechamentos, pagamentos, entrega de documentos (bilhetes, vouchers etc.). (00034-02-C-07)
            Conhecimento: Softwares de reserva e emissão de passagens aéreas, rodoviárias, meios de hospedagem, locação de veículo, cruzeiros marítimos e cartões de assistência. (00034-02-C-08)
            Conhecimento: Procedimentos técnicos: embarque e desembarque em terminais rodoviários, marítimos e aéreos. (00034-02-C-09)
            Conhecimento: Técnicas de negociação e venda: abordagem, atendimento e vendas, negociação de Produtos e Serviços, Superação de objeções, Estratégias Persuasivas. (00034-02-C-10)
            Conhecimento: Direitos e deveres do consumidor: princípios do Código de Defesa do Consumidor, formas de atuação como prestador de serviços, autônomo ou contratado/colaborador, impostos e sindical. (00034-02-C-11)
            Conhecimento: Políticas de tarifas, descontos e comissões: contratos, cálculo, venda, fechamento e conferência. (00034-02-C-12)
            Conhecimento: Divulgação de produtos e destinos turísticos: meios e ferramentas tecnológicas; uso das mídias sociais. (00034-02-C-13)
            Conhecimento: Ferramentas do ambiente virtual: sites de busca, Google Earth, Google Maps. (00034-02-C-14)
            Habilidade: Comunicar-se de maneira assertiva. (00034-02-H-01)
            Habilidade: Interpretar linguagem de sistemas de reservas, emissão, mapas e guias turísticos. (00034-02-H-02)
            Habilidade: Utilizar recursos tecnológicos para a divulgação de atrativos e destinos turísticos. (00034-02-H-03)
            Habilidade: Orientar-se geograficamente. (00034-02-H-04)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00034-02-H-05)
            Habilidade: Mediar conflitos nas situações de trabalho. (00034-02-H-06)
            Habilidade: Calcular estimativas e percentual. (00034-02-H-07)
            Habilidade: Orientar sobre vistos, documentos e efeitos fisiológicos. (00034-02-H-08)
            Orientações metodológicas: Sugere-se ao docente promover situações de aprendizagem que evidenciem as práticas operacionais em agências de turismo levando em consideração as tendências do mercado de viagem, segmentação do turismo e situações reais do mundo do trabalho. Podem ser adotadas estratégias como pesquisas, visitas técnicas em agências de turismo, situações-problemas de atendimento ao cliente e vendas de pacotes turísticos, simulação das etapas do processo de viagem e elaboração, promoção de roteiros e pacotes turísticos e sugerimos, ainda, que o docente apresente e faça treinamento em algum GDS, a depender da viabilidade de cada DR.
            O desenvolvimento destas estratégias permite colocar o aluno em movimento, possibilitando a vivência de situações do dia a dia, nas quais os participantes têm a oportunidade de realizar a autoavaliação e o aperfeiçoamento das habilidades e atitudes profissionais.  (00034-02-O-01)
            Tecnologias sugeridas:
            IAA, Sistemas de Recomendação, Algoritmos de recomendação para sugerir produtos e serviços ao cliente com base no histórico de compras., geração, Gera novas atividades de marketing direto e automatizado., 00034-02-I-01|00034-02-C-03|00034-02-C-05|00034-02-H-01, Cria novas possibilidades de prestar informações sobre os produtos (I-01) que independem do trabalhador enquanto articula técnicas de vendas (H-01) e conhecimentos sobre o perfil comportamental dos clientes (C-04) e tipos de viagem (C-05).

            Título da UC: Assessorar o Viajante (00034-03)
            Indicador: Atende o cliente durante a viagem, conforme demanda e venda realizada. (00034-03-I-01)
            Indicador: Presta assistência especializada nas intercorrências ao longo da viagem, de acordo com a necessidade do cliente e os recursos locais. (00034-03-I-02)
            Conhecimento: Assessoria de viagem: sugestão de restaurantes, passeios, atrativos, ajuste de roteiro. (00034-03-C-01)
            Conhecimento: Gerenciamento de conflitos e crises em viagens: overbooking, passageiro atrasado ou desaparecido, conflitos armados, guerras civis, terrorismo, condições meteorológicas adversas, endemias, pandemias, acidente e morte de passageiros. (00034-03-C-02)
            Conhecimento: Órgãos oficiais: consulados, bombeiros, delegacias de atendimento ao turista, hospitais, embaixadas, redes de assistência à saúde. (00034-03-C-03)
            Conhecimento: Pós-venda: avaliação da viagem e estratégias de fidelização. (00034-03-C-04)
            Habilidade: Mediar conflitos na situação de trabalho. (00034-03-H-01)
            Habilidade: Comunicar-se de maneira assertiva. (00034-03-H-02)
            Habilidade: Organizar documentos e local de trabalho. (00034-03-H-03)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na organização. (00034-03-H-04)
            Habilidade: Analisar resultados de pesquisas de satisfação do cliente. (00034-03-H-05)
            Orientações metodológicas: Recomenda-se que sejam propiciados aos alunos experiências de simulação como clientes em situações do tipo city-tour e visitas técnicas, possibilitando a prática de colocar-se no lugar do viajante. Propõe-se a realização de análise de casos reais de conflitos e crises em viagens como: overbooking, passageiro atrasado ou desaparecido ou demais situações problema que possam ocorrer em viagens. (00034-03-O-01)
            Tecnologias sugeridas:
            IAA, Assistentes Virtuais, Assistentes virtuais podem ajudar na comunicação contínua com os clientes durante a viagem, oferecendo assistência personalizada e imediata., ampliação, Amplia a capacidade de resposta e comunicação com o cliente durante a viagem., 00034-03-I-02|00034-03-H-02|00034-03-C-01, Melhora a assistência especializada durante a viagem (I-02), articulando comunicação assertiva (H-02) e conhecimentos de assessoria de viagem (C-01).
            APP, Plataformas de Comunicação e Colaboração Online, Facilitam gestão e acompanhamento de todas as demandas e necessidades do cliente em tempo real ao longo da viagem., ampliação, A ampliação da tecnologia facilita a comunicação em tempo real, organizando melhor a assistência ao viajante., 00034-03-I-01|00034-03-I-02|00034-03-H-02, Facilitam diretamente a prestação de assistência durante a viagem (I-01/I-02), articulando a comunicação assertiva (H-02).

            Curso: Aprendizagem Profissional Técnica em Segurança do Trabalho (2528) - 2019

            Título da UC: Elaborar, Implantar e Implementar a Política de Saúde e Segurança do Trabalho (00082-01)
            Indicador: Identifica os indicadores dos modelos de gestão e da cultura organizacional, conforme legislação, literatura técnica e diretrizes da organização. (00082-01-I-01)
            Indicador: Identifica riscos à saúde e segurança dos trabalhadores, conforme legislação e normas técnicas. (00082-01-I-02)
            Indicador: Define metas, prioridades e responsabilidades, conforme diretrizes da política e legislação. (00082-01-I-03)
            Indicador: Define novos programas e procedimentos, conforme normas, legislação e diretrizes da organização. (00082-01-I-04)
            Indicador: Divulga política de saúde e segurança entre os colaboradores, de acordo com as diretrizes da organização. (00082-01-I-05)
            Indicador: Atualiza política de saúde e segurança da organização, conforme avaliação dos resultados da política. (00082-01-I-06)
            Conhecimento: Legislação e normativas do trabalho: Normas regulamentadoras, Consolidação das Leis do Trabalho (CLT), Normas Brasileiras (NBRs), anuários de estatísticas de acidentes do trabalho, diretrizes sobre sistemas de gestão de segurança e saúde no trabalho Organização Internacional do Trabalho (OIT). (00082-01-C-01)
            Conhecimento: Indicadores de modelo de gestão: tipos e características. (00082-01-C-02)
            Conhecimento: Política de saúde e segurança do trabalho: conceitos, estrutura, importância. (00082-01-C-03)
            Conhecimento: Serviço Especializado em Engenharia de Segurança e Medicina do Trabalho (SESMT): dimensionamento de profissionais. (00082-01-C-04)
            Conhecimento: Atuação do SESMT: atribuições na elaboração, implantação e implementação da política e da gestão de saúde e segurança do trabalho. (00082-01-C-05)
            Conhecimento: Comissão Interna de Prevenção de Acidentes (CIPA): conceito e dimensionamento de membros e atribuições. (00082-01-C-06)
            Conhecimento: Conceitos e princípios de administração: modelos de gestão e cultura organizacional nas relações de trabalho. (00082-01-C-07)
            Conhecimento: Gestão de saúde e segurança do trabalho nas Organizações: objetivos, implantação, melhoria contínua, PDCA (plan, do, check, act) e monitoramento. (00082-01-C-08)
            Conhecimento: Princípios e diretrizes da OIT e Organização Mundial da Saúde (OMS): melhoria das condições de segurança do ambiente de trabalho e saúde do trabalhador. (00082-01-C-09)
            Conhecimento: Plano Nacional de Saúde e Segurança do Trabalho. (00082-01-C-10)
            Conhecimento: Diretrizes do Ministério do Trabalho, Normas Regulamentadoras (NRs) sobre política e gestão de saúde e segurança do trabalho. (00082-01-C-11)
            Conhecimento: Normas da Associação Brasileira de Normas Técnicas (ABNT) sobre gestão de saúde e segurança do trabalho. (00082-01-C-12)
            Conhecimento: História e Evolução do trabalho: a segurança do trabalho na revolução agrícola, na revolução industrial e na revolução da informação. (00082-01-C-13)
            Conhecimento: Introdução à Portaria n° 3.214/78 e à Lei n° 6.514/1977, definição, estrutura e hierarquia da legislação vigente: leis, decretos, resoluções, portarias, instruções normativas, súmulas do Técnico em Segurança do Trabalho e outros. (00082-01-C-14)
            Conhecimento: Atribuições do Técnico em Segurança do Trabalho. (00082-01-C-15)
            Conhecimento: Definição de Responsabilidade Civil e Criminal aplicada à Saúde e Segurança do Trabalho. (00082-01-C-16)
            Conhecimento: Definição e características da Legislação Previdenciária relacionada à saúde e à segurança no trabalho: PPP e aposentadoria especial, E-social. (00082-01-C-17)
            Conhecimento: Conceitos de acidentes de trabalho: CAT, NTEP e FAP; tipos de acidente de trabalho. (00082-01-C-18)
            Conhecimento: Definição e aplicabilidade de Inspeção Prévia, Embargo e Interdição. (00082-01-C-19)
            Conhecimento: Classificação de riscos ambientais de acordo com a portaria vigente. (00082-01-C-20)
            Conhecimento: Planejamento de vida e carreira: desenvolvimento de metas pessoais, profissionais e econômicas e uma proposta de guia para alcançá-las; mundo do trabalho; empreendedorismo e outras formas de inserção no mercado de trabalho; construção de itinerário de profissionalização. (00082-01-C-21)
            Habilidade: Comunicar-se de maneira assertiva. (00082-01-H-01)
            Habilidade: Interpretar textos e procedimentos técnicos. (00082-01-H-02)
            Habilidade: Elaborar textos e apresentações técnicas. (00082-01-H-03)
            Habilidade: Utilizar técnicas de medição e controle. (00082-01-H-04)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na organização. (00082-01-H-05)
            Habilidade: Operar planilhas de cálculos, editores de texto e mídias para criação e exibição de apresentação. (00082-01-H-06)
            Habilidade: Analisar as etapas do processo de trabalho. (00082-01-H-07)
            Orientações metodológicas: 	Elaborar, implantar e implementar a Política de Saúde e Segurança do Trabalho
            Nessa Unidade Curricular o docente poderá planejar atividades de pesquisas na internet de organizações e suas políticas de saúde e segurança do trabalho, possibilitando aos alunos conhecerem a importância e as vantagens de uma política de segurança coerente com o ambiente de atuação do aprendiz Técnico em Segurança do Trabalho. Em relação ao Mapa de Risco, sugere-se que na Unidades Curricular 1 seja realizada uma introdução e trabalhados os conhecimentos básicos sobre o tema. 
            Os conteúdos indicados nos incisos X, XI e XII das diretrizes relacionadas no art. 336 da Portaria n.º 671/2021, relacionados a essa UC, deverão ser contextualizados quando o docente abordar os elementos da competência. (00082-01-O-01)
            Tecnologias sugeridas:
            APP, Sistemas de Gerenciamento de Tarefas e Projetos, Facilitam a organização e controle das tarefas relacionadas à implementação da política de saúde e segurança do trabalho., ampliação, Ampliação da eficiência na definição de metas, prioridades e responsabilidades e na atualização de políticas., 00082-01-I-03|00082-01-I-06|00082-01-C-08|00082-01-H-06, Essas ferramentas auxiliam diretamente na definição de metas (I-03) e atualização de políticas (I-06), envolvendo conhecimentos sobre gestão de saúde e segurança (C-08) e melhorando a organização de documentos e controle (H-06).

            Título da UC: Realizar Avaliação e Medidas de Controle de Riscos Físicos, Químicos e Biológicos (00082-02)
            Indicador: Identifica e classifica riscos ambientais, conforme literatura técnica, normas e legislações aplicáveis. (00082-02-I-01)
            Indicador: Representa graficamente o leiaute dos ambientes de trabalho, conforme técnica para elaboração de croquis e mapa de risco. (00082-02-I-02)
            Indicador: Identifica os agentes ambientais e define as avaliações qualitativas e quantitativas dos riscos, conforme seus limites de tolerância. (00082-02-I-03)
            Indicador: Estabelece medidas de controle, conforme manuais, normas e legislações aplicáveis. (00082-02-I-04)
            Conhecimento: Normas regulamentadoras relativas aos riscos físicos, químicos e biológicos. (00082-02-C-01)
            Conhecimento: Legislação trabalhista e previdenciária: aspectos referentes à avaliação de riscos. (00082-02-C-02)
            Conhecimento: Elaboração de croquis e desenho técnico: representações gráficas, escalas de redução e de ampliação, leitura e interpretação de plantas, representação de postos de trabalho, normas técnicas e literatura técnica. (00082-02-C-03)
            Conhecimento: Fundamentos de higiene ocupacional: princípios, conceitos, definições e noções de fisiologia humana. (00082-02-C-04)
            Conhecimento: Riscos físicos, químicos e biológicos: definições, tipos, categorias e grupos. (00082-02-C-05)
            Conhecimento: Agentes físicos “ruído, vibração, calor, frio, umidade, radiação, pressões anormais”. (00082-02-C-06)
            Conhecimento: Processos produtivos e suas características. (00082-02-C-07)
            Conhecimento: Agentes químicos: aerodispersóides, gases, névoas vapores, neblinas, poeiras e fumos. (00082-02-C-08)
            Conhecimento: Produtos químicos: Ficha de Informação de Segurança de Produtos Químicos (FISPQ) e ficha de emergência. (00082-02-C-09)
            Conhecimento: Agentes biológicos: fungos, bactérias, vírus e protozoários; Biossegurança e Normas regulamentadoras. (00082-02-C-10)
            Conhecimento: Análise de riscos ocupacionais: checklist e inspeção. (00082-02-C-11)
            Conhecimento: Doenças ocupacionais: legislação vigente, conceito, tipos, vias de penetração e prevenção. (00082-02-C-12)
            Conhecimento: Conceitos de limites de exposição ocupacional: Limite de Tolerância (LT) NR pertinente; Thereshold Limit Values (TLV) American Conference of Governmental Industrial Hygienists (ACGIH); Permissible Exposion Limits (PEL) Ocupational Safety and Health Administration (OSHA); Recommended Exposure Limit (REL); National Institute for Occupational Safety and Health (NIOSH); Valor Teto, Valor Máximo, Nível de ação, ppm e mg/m3. (00082-02-C-13)
            Conhecimento: Técnicas de controle de agentes físicos, químicos e biológicos; medidas administrativas. (00082-02-C-14)
            Conhecimento: Equipamento de Proteção Coletiva (EPC) e Equipamento de Proteção Individual (EPI). (00082-02-C-15)
            Conhecimento: Normas de Higiene Ocupacional (NHO) Fundacentro, normas nacionais e internacionais: (ACGIH -NIOSH) e procedimentos. (00082-02-C-16)
            Conhecimento: Avaliações qualitativas e quantitativas: tipos e metodologias específicas das NHOs. (00082-02-C-17)
            Conhecimento: Descrição de atividades laborais e elaboração de Ordens de Serviço. (00082-02-C-18)
            Habilidade: Comunicar-se de maneira assertiva. (00082-02-H-01)
            Habilidade: Interpretar textos e procedimentos técnicos. (00082-02-H-02)
            Habilidade: Mediar conflitos nas situações de trabalho. (00082-02-H-03)
            Habilidade: Utilizar técnicas de medição e controle. (00082-02-H-04)
            Habilidade: Realizar análise de risco. (00082-02-H-05)
            Habilidade: Operar planilhas de cálculos, editores de texto e de apresentação. (00082-02-H-06)
            Habilidade: Analisar e propor soluções por meio de raciocínio lógico e crítico dos processos de trabalho. (00082-02-H-07)
            Orientações metodológicas: Realizar avaliação e medidas de controle de riscos físicos, químicos e biológicos
            Nesta Unidade Curricular são tratadas as Normas Nacionais, estas devem ser priorizadas em relação às internacionais, conforme a seguinte recomendação descrita na NR 9: “quando os resultados das avaliações quantitativas da exposição dos trabalhadores excederem os valores dos limites previstos na NR 15 ou na ausência destes os valores limites de exposição ocupacional adotados pela ACGIH – American Conference of Govermmental Industrial Higyenist, ou aqueles que venham a ser estabelecidos em negociação coletiva de trabalho, desde que mais rigorosos do que os critérios técnico-legais estabelecidos. 
            O desenho técnico poderá ser trabalho pelo docente de forma manual ou com a utilização de programas específicos (como o SketchUp, que é de uso livre) para elaboração de croquis, para compor o relatório de análise dos riscos físicos, químicos e biológicos. Quando houver necessidades de fazer cálculos, o docente poderá acessar calculadoras online, em diversos sites, por exemplo:
            1 - http://www.calculadoraonline.com.br/cientifica
            2- http://www.calculadoraonline.com.br/calculadora-virtual-gratis
            3 - http://www.alcula.com/es/calculadoras/calculadora-cientifica/
            Em relação ao Mapa de Risco, sugere-se que nas Unidades Curriculares 1 e 2 seja realizada uma introdução e trabalhados os conhecimentos básicos sobre o tema. Poderão ser trabalhados os conhecimentos em relação aos instrumentos utilizados em Higiene Ocupacional. Nesta Unidade, poderá ser utilizado software visualizador de plantas AutoCAD, que é de uso livre.
            Os conteúdos indicados nos incisos X, XI e XII das diretrizes relacionadas no art. 336 da Portaria n.º 671/2021, relacionados a essa UC, deverão ser contextualizados quando o docente abordar os elementos da competência. (00082-02-O-01)

            Título da UC: Realizar Avaliação e Medidas de Controle de Riscos Ergonômicos e de Acidentes (00082-03)
            Indicador: Identifica os riscos ergonômicos, conforme literatura técnica, normas e demais legislações aplicáveis. (00082-03-I-01)
            Indicador: Identifica os riscos de acidentes, conforme literatura técnica, normas e legislações aplicáveis. (00082-03-I-02)
            Indicador: Avalia os riscos ergonômicos e de acidentes por meio de instrumentos de avaliação e metodologia específica. (00082-03-I-03)
            Indicador: Estabelece medidas de controle, conforme literatura técnica, normas e demais legislações aplicáveis. (00082-03-I-04)
            Indicador: Realiza análise de leiaute a partir da leitura da planta baixa, considerando as normas de segurança e saúde do trabalho. (00082-03-I-05)
            Conhecimento: Normas Regulamentadoras relacionadas a riscos ergonômicos, de acidentes e Normas ABNT. (00082-03-C-01)
            Conhecimento: Legislação trabalhista e previdenciária: aspectos referentes à avaliação de riscos. (00082-03-C-02)
            Conhecimento: Ergonomia: conceito, segmentos, agentes ergonômicos e consequências à saúde do trabalhador - Lesão por Esforço Repetitivo (LER) e Distúrbio Osteomuscular Relacionado ao Trabalho (DORT). (00082-03-C-03)
            Conhecimento: Principais riscos de acidentes no ambiente de trabalho: eletricidade, trabalho em altura, construção civil, espaço confinado, máquinas e equipamentos e demais riscos em segmentos específicos. (00082-03-C-04)
            Conhecimento: Avaliações de riscos de acidentes e medidas de controle no ambiente de trabalho. (00082-03-C-05)
            Conhecimento: Conceito de antropometria estática e dinâmica. (00082-03-C-06)
            Conhecimento: Biomecânica ocupacional: reação do corpo humano em relação aos aspectos ergonômicos. (00082-03-C-07)
            Conhecimento: Iluminamento e seus efeitos sobre o organismo humano. (00082-03-C-08)
            Conhecimento: Definição, objetivos, métodos, técnicas e recomendações da Análise Ergonômica do Trabalho AET; Psicologia do trabalho: da ergonomia cognitiva; Ergonomia física e organizacional. (00082-03-C-09)
            Conhecimento: Avaliação do ambiente/atividade de trabalho: organização do trabalho, esforço físico (levantamento, transporte e movimentação manual de materiais), trabalho noturno; ritmo de trabalho. (00082-03-C-10)
            Conhecimento: Medidas de controle para riscos ergonômicos e de acidentes. (00082-03-C-11)
            Conhecimento: Elaboração de croquis: Representação dos ambientes e organização do trabalho; construção de arranjos físicos adaptados ao trabalhador, normas técnicas e literatura técnica. (00082-03-C-12)
            Conhecimento: Projetos de adequações: definição e responsabilidade. (00082-03-C-13)
            Conhecimento: Incidentes e acidentes de trabalho: conceitos e diretrizes. (00082-03-C-14)
            Conhecimento: Tipos de manutenção para a prevenção de acidentes: preventiva, corretiva, preditiva, detectiva. (00082-03-C-15)
            Conhecimento: Tipos de inspeções de segurança para a prevenção de acidentes e doenças ocupacionais: relatórios, planilhas, formulários e checklist. (00082-03-C-16)
            Conhecimento: Análise de processos tecnológicos e características de locais e atividades profissionais. (00082-03-C-17)
            Conhecimento: Procedimentos para ordens de serviço: tipo Procedimento Operacional Padrão (POP) e elaboração. (00082-03-C-18)
            Habilidade: Comunicar-se de maneira assertiva. (00082-03-H-01)
            Habilidade: Interpretar textos e procedimentos técnicos. (00082-03-H-02)
            Habilidade: Mediar conflitos nas situações de trabalho. (00082-03-H-03)
            Habilidade: Utilizar técnicas de medição e controle. (00082-03-H-04)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na organização. (00082-03-H-05)
            Habilidade: Identificar posturas e movimentos incorretos na execução de tarefas. (00082-03-H-06)
            Habilidade: Realizar análise de riscos. (00082-03-H-07)
            Habilidade: Operar planilhas de cálculos, editores de texto e de apresentação. (00082-03-H-08)
            Habilidade: Analisar as etapas do processo de trabalho. (00082-03-H-09)
            Orientações metodológicas: Realizar avaliação e medidas de controle de riscos ergonômicos e de acidentes
            Nesta Unidade, o desenho técnico poderá ser trabalhado pelo docente de forma manual ou com a utilização de recursos tecnológicos, com programas específicos como o SketchUp, que é de uso livre na análise dos riscos ergonômicos. Sugere-se que o docente realize ações relacionadas à produção de Mapas de Riscos. Poderão ser trabalhados os conhecimentos em relação aos instrumentos utilizados em Higiene Ocupacional. Nesta Unidade, poderá ser utilizado software visualizador de plantas AutoCAD, que é de uso livre.
            Os conteúdos indicados nos incisos X, XI e XII das diretrizes relacionadas no art. 336 da Portaria n.º 671/2021, relacionados a essa UC, deverão ser contextualizados quando o docente abordar os elementos da competência. (00082-03-O-01)
            Tecnologias sugeridas:
            IAA, Visão Computacional, Tecnologias de visão computacional podem avaliar posturas e movimentos, identificando riscos ergonômicos automaticamente., substituição, Substituir a observação manual por ferramentas automatizadas., 00082-03-I-01|00082-03-I-02|00082-03-C-03|00082-03-C-04|00082-03-C-05, A visão computacional permite uma análise automatizada de parte dos riscos ergonômicos (I-01) e de acidentes (I-02), envolvendo conhecimentos sobre ergonomia (C-03) e riscos de acidentes (C-04, C-05).

            Título da UC: Monitorar Riscos Ocupacionais (00082-04)
            Indicador: Seleciona os equipamentos de medição conforme o risco ocupacional. (00082-04-I-01)
            Indicador: Calibra, prepara e utiliza equipamentos conforme procedimentos técnicos. (00082-04-I-02)
            Indicador: Aplica técnicas de amostragem no monitoramento quantitativo em cada grupo homogêneo de risco, conforme manuais de higiene ocupacional. (00082-04-I-03)
            Indicador: Avalia os níveis de exposição aos riscos ocupacionais, conforme legislação. (00082-04-I-04)
            Indicador: Propõe e acompanha a implantação de melhorias a partir dos cronogramas definidos nos procedimentos organizacionais. (00082-04-I-05)
            Conhecimento: Normas regulamentadoras: monitoramento dos riscos ambientais, ocupacionais e as Normas de Higiene Ocupacional da Fundacentro. (00082-04-C-01)
            Conhecimento: Avaliação ambiental: instrumentos, etapas, procedimentos, metodologia qualitativa e quantitativa, tipos de planilhas, relatórios e checklist da norma. (00082-04-C-02)
            Conhecimento: Equipamentos de medição: tipos, características, funcionamento (softwares), dados e resultados apresentados. (00082-04-C-03)
            Conhecimento: Procedimentos para aferição e calibração de equipamentos e arquivamento de certificados. (00082-04-C-04)
            Conhecimento: Entrevistas para monitoramento das condições de trabalho: coleta de dados e organização das informações. (00082-04-C-05)
            Conhecimento: Tipos e técnicas de amostragem: grupo homogêneo de exposição, amostragem ativa e passiva, técnicas da HSE, técnicas da NIOSH, amostragens rápidas, curto prazo, longo prazo Time Weighted Average (TWA), contínua e BULK. (00082-04-C-06)
            Habilidade: Comunicar-se de maneira assertiva. (00082-04-H-01)
            Habilidade: Interpretar textos e procedimentos técnicos. (00082-04-H-02)
            Habilidade: Registrar informações das avaliações ambientais. (00082-04-H-03)
            Habilidade: Realizar cálculos nos processos de monitoramento de riscos. (00082-04-H-04)
            Habilidade: Categorizar as etapas do processo de monitoramento. (00082-04-H-05)
            Habilidade: Manusear equipamentos e acessórios de monitoramentos. (00082-04-H-06)
            Habilidade: Pesquisar dados sobre as condições de trabalho. (00082-04-H-07)
            Habilidade: Interpretar dados sobre as condições de trabalho. (00082-04-H-08)
            Habilidade: Realizar análise de risco. (00082-04-H-09)
            Habilidade: Operar planilhas de cálculos, editores de texto e de apresentação. (00082-04-H-10)
            Habilidade: Analisar as etapas do processo de trabalho. (00082-04-H-11)
            Orientações metodológicas: Monitorar riscos ocupacionais
            Na Unidade Curricular indica-se a realização de no mínimo uma demonstração de monitoramento para cada agente. No caso de agentes químicos, no mínimo uma avaliação para: gases, vapores, poeiras e tubos colorimétricos. Nesta Unidade o desenho técnico poderá ser trabalhado pelo docente de forma manual ou com a utilização de recursos tecnológicos, com programas específicos, como o SketchUp, que é de uso livre, na análise dos riscos ergonômicos. Deverão ser realizadas demonstrações sobre os instrumentos utilizados em Higiene Ocupacional e suas funcionalidades pelos alunos; Nesta Unidade, poderá ser utilizado software visualizador de plantas AutoCAD, que é de uso livre.
            Os conteúdos indicados nos incisos X, XI e XII das diretrizes relacionadas no art. 336 da Portaria n.º 671/2021, relacionados a essa UC, deverão ser contextualizados quando o docente abordar os elementos da competência. (00082-04-O-01)
            Tecnologias sugeridas:
            IOT, Plataformas de Gerenciamento de IoT, Automatiza o monitoramento e controle dos equipamentos para coleta de dados de ambientes de trabalho., substituição, Substitui a monitoramento manual com sistemas automatizados., 00082-04-I-01|00082-04-I-03, Permite controle mais eficiente e detalhado dos sensores de medições ambientais (I-01) e acompanhamento de amostragem de riscos ocupacionais (I-03).

            Título da UC: Executar Ações de Investigação, Registro e Controle de Incidentes, Acidentes de Trabalho e Doenças Ocupacionais (00082-05)
            Indicador: Analisa o incidente e/ou acidente, conforme procedimentos técnicos. (00082-05-I-01)
            Indicador: Define medidas corretivas, conforme procedimentos técnicos. (00082-05-I-02)
            Indicador: Avalia potencial de perdas e danos do acidente/incidente, conforme procedimentos da organização. (00082-05-I-03)
            Indicador: Calcula e registra perdas e danos do acidente, conforme procedimentos da organização e os formulários dos Quadros III e IV da Norma Regulamentadora. (00082-05-I-04)
            Indicador: Utiliza método específico para investigação de cada tipo de incidente ou acidente de trabalho, conforme manuais e literatura técnica. (00082-05-I-05)
            Indicador: Emprega ferramentas para auxiliar a investigação de acidentes e incidentes de trabalho, conforme procedimentos da organização e manuais de investigação. (00082-05-I-06)
            Indicador: Coleta e tabula informações dos acidentes e incidentes, conforme metodologia. (00082-05-I-07)
            Conhecimento: Normas regulamentadoras: investigação, registro e controle de incidentes, acidentes de trabalho e doenças ocupacionais. (00082-05-C-01)
            Conhecimento: Legislação previdenciária sobre acidentes do trabalho: aspectos referentes aos direitos dos trabalhadores. (00082-05-C-02)
            Conhecimento: Acidente e incidentes: conceitos, tipos e teorias (Heinrich e Bird), metodologias de investigação e acompanhamento de ocorrências e elaboração de relatórios e formulários. (00082-05-C-03)
            Conhecimento: Procedimentos legais nos acidentes de trabalho: perícias e fiscalizações (tipos, acompanhamento e assessorias), definições e aplicação Nexo Técnico Epidemiológico (NTEP), Fator Acidentário de Prevenção (FAP), Seguro de Acidente do Trabalho (SAT), formulários dos quadros III, IV, V e VI da NR 4, e-Social. (00082-05-C-04)
            Conhecimento: Metodologia de avaliação e controle de perdas e danos. (00082-05-C-05)
            Conhecimento: Cálculos: de perdas e danos, estatísticas, taxa de frequência e gravidade. (00082-05-C-06)
            Conhecimento: Conceitos de adicional, insalubridade e periculosidade. (00082-05-C-07)
            Conhecimento: Fiscalizações de saúde e segurança do trabalho e penalidades. (00082-05-C-08)
            Conhecimento: Tipos e características de metodologias: árvores de causas, Heinrich, Bird, Couto, Costella e Saurin. (00082-05-C-09)
            Conhecimento: Tipos de ferramentas: câmaras, trenas, entrevistas, checklist. (00082-05-C-10)
            Conhecimento: Técnicas de investigação: tipos e características. (00082-05-C-11)
            Conhecimento: Análise para reabilitação de funcionários após acidentes de trabalho. (00082-05-C-12)
            Habilidade: Comunicar-se de maneira assertiva. (00082-05-H-01)
            Habilidade: Redigir textos, relatórios e procedimentos. (00082-05-H-02)
            Habilidade: Mediar conflitos nas situações de trabalho. (00082-05-H-03)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na organização. (00082-05-H-04)
            Habilidade: Demonstrar criatividade nas propostas de ações corretivas. (00082-05-H-05)
            Habilidade: Realizar análise de risco. (00082-05-H-06)
            Habilidade: Analisar as etapas do processo de trabalho. (00082-05-H-07)
            Orientações metodológicas: Executar ações de investigação, registro e controle de incidentes, acidentes de trabalho e doenças ocupacionais
            O aluno deve conhecer as principais metodologias disponíveis na literatura, em relação a investigação de incidentes e acidentes de trabalho, porém o docente poderá definir apenas uma metodologia para o foco da Unidade Curricular. É importante que o docente planeje atividades, como visitas técnicas, exemplos reais para que o aluno analise criticamente as situações de acidentes e incidentes. O docente poderá solicitar uma pesquisa sobre softwares disponíveis no mercado, que são desenvolvidos para gerenciar ocorrências de acidentes e incidentes nas organizações.
            Os conteúdos indicados nos incisos X, XI e XII das diretrizes relacionadas no art. 336 da Portaria n.º 671/2021, relacionados a essa UC, deverão ser contextualizados quando o docente abordar os elementos da competência. (00082-05-O-01)
            Tecnologias sugeridas:
            APP, Ferramentas de Anotação e Organização de Conteúdos, Facilitam o registro de informações e elaboração de relatórios de investigação de acidentes., ampliação, Ampliam a eficiência na coleta e organização de dados de incidentes e acidentes., 00082-05-I-01|00082-05-I-07|00082-05-H-02|00082-05-H-06, Facilitam a investigação e registro de informações (I-01, I-07), incluindo a redação de relatórios e procedimentos (H-02) e análise de riscos (H-06).

            Título da UC: Auxiliar e Executar Ações de Elaboração de Programas de Saúde e Segurança do Trabalho (00082-06)
            Indicador: Verifica objetivo e campo de aplicação de programas de Saúde e Segurança do Trabalho, conforme as normas vigentes. (00082-06-I-01)
            Indicador: Estrutura os programas de gerenciamento de riscos, considerando suas particularidades e parâmetros da legislação. (00082-06-I-02)
            Indicador: Elabora os planos de ação para gerenciamento de riscos, conforme normas estabelecidas para os programas. (00082-06-I-03)
            Indicador: Gerencia e controla documentos em meio físico e eletrônico, de acordo com os procedimentos da organização e legislação. (00082-06-I-04)
            Conhecimento: Normas regulamentadoras do trabalho relacionadas à programas de saúde e segurança do trabalho. (00082-06-C-01)
            Conhecimento: Estrutura de programas e planos de ação de gerenciamento de riscos e saúde ocupacional: Gerenciamento de Riscos Ocupacionais (GRO) e Programa de Prevenção de Riscos, Programa de Controle Médico de Saúde Ocupacional (PCMSO). (00082-06-C-02)
            Conhecimento: Documentos: tipos, características, preenchimento e arquivamento em meio físico e eletrônico. (00082-06-C-03)
            Conhecimento: Estratégias de resolução de problemas: identificação, diagnóstico e negociação. (00082-06-C-04)
            Conhecimento: Técnica do PDCA na formatação e planejamento do Plano de ação dos programas definidos em legislação. (00082-06-C-05)
            Habilidade: Comunicar-se de maneira assertiva. (00082-06-H-01)
            Habilidade: Interpretar textos e procedimentos técnicos. (00082-06-H-02)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na organização. (00082-06-H-03)
            Habilidade: Elaborar relatórios técnicos. (00082-06-H-04)
            Habilidade: Utilizar técnicas de medição e controle. (00082-06-H-05)
            Habilidade: Realizar análise de risco. (00082-06-H-06)
            Habilidade: Analisar as etapas do processo de trabalho. (00082-06-H-07)
            Orientações metodológicas: Auxiliar e executar ações de elaboração dos programas de saúde e segurança do trabalho
            O docente poderá planejar atividades para que os alunos conheçam os principais programas de saúde e segurança do trabalho, recomenda-se priorizar a elaboração do Gerenciamento de Riscos Ocupacionais (GRO) e Programa de Prevenção de Riscos, utilizando ambiente que tenha computadores para a realização dos trabalhos. Existem softwares disponíveis no mercado para gerenciar a saúde e segurança do trabalho nas organizações, é importante que os alunos tenham conhecimentos sobre as funcionalidades, desta forma uma demonstração poderá ser realizada à turma. É importante que o docente planeje atividades, como visitas técnicas, pesquisas e entrevista com profissionais em organizações de prestação de serviços em medicina e segurança do trabalho.
            Os conteúdos indicados nos incisos X, XI e XII das diretrizes relacionadas no art. 336 da Portaria n.º 671/2021, relacionados a essa UC, deverão ser contextualizados quando o docente abordar os elementos da competência. (00082-06-O-01)
            Tecnologias sugeridas:
            APP, Ferramentas de Formulários Online, Ferramentas como Google Forms e JotForm facilitam a coleta e análise de dados para avaliação e implementação de programas de saúde e segurança., ampliação, Aumentam a eficiência na coleta e análise de dados, melhorando a gestão e controle de programas., 00082-06-I-01|00082-06-I-04|00082-06-C-03|00082-06-H-04, Facilita a verificação de objetivos e campo de aplicação (I-01) e controle de documentos (I-04), envolvendo conhecimentos de documentação de SST (C-03) e habilidades de redação de relatórios técnicos (H-04).

            Curso: Assistente de Pessoal (1513) - 2014

            Título da UC: Apoiar e Executar Ações Referentes Às Rotinas de Admissão e Demissão de Colaboradores (00202-01)
            Indicador: Abre o prontuário do colaborador, guardando a documentação e atualizando os dados e as informações, sempre que necessário, garantindo seu sigilo. (00202-01-I-01)
            Indicador: Preenche formulários, recolhe, organiza e guarda documentos relacionados às rotinas de admissão e demissão, garantindo sua integralidade e respeitando a temporalidade conforme requisitos legais e fluxo estabelecido pela organização. (00202-01-I-02)
            Indicador: Tabula dados provenientes dos questionários de desligamento e efetua o levantamento de pendências financeiras e dos materiais que necessitam ser devolvidos, dando os encaminhamentos devidos conforme procedimentos da organização para elaboração do termo de rescisão do contrato de trabalho. (00202-01-I-03)
            Indicador: Recolhe documentação, agenda exames admissionais e demissionais, bem como a homologação nos sindicatos, conforme requisitos legais e procedimentos da organização. (00202-01-I-04)
            Conhecimento: Planejamento de carreira: mundo do trabalho, formas de inserção no mercado de trabalho, marketing e apresentação pessoal, preparação de currículos, entrevista de emprego. (00202-01-C-01)
            Conhecimento: Fundamentos da legislação trabalhista e poder disciplinador do empregador: conceitos de empregado, empregador, autônomo, avulso, estagiário, relação de emprego; admissão e demissão de empregados, remuneração, jornada de trabalho, férias; poder disciplinador do empregador, advertência, suspensão, demissão por justa causa, demissão sem justa causa, aviso prévio, dispensa do empregado estável, apuração de faltas graves. (00202-01-C-02)
            Conhecimento: Consolidação das Leis do Trabalho (CLT) Decreto-lei 5.452/23 e alterações. Definição e consulta. (00202-01-C-03)
            Conhecimento: Constituição da República Federativa do Brasil, artigo 7º: dos direitos e garantias fundamentais dos trabalhadores. Relações de trabalho: conceito, histórico, relação homem e trabalho. Tipos de contrato de trabalho e formas de rescisão. (00202-01-C-04)
            Conhecimento: Estrutura do Departamento de Pessoal. (00202-01-C-05)
            Conhecimento: Processo de admissão: documentos necessários, ficha de registro de empregado, registro/anotações/ atualizações em carteira de trabalho, cadastro do Programa de Integração Social/Programa de Formação do Patrimônio do Servidor Público (PIS/Pasep), declaração de encargos de família para fins de imposto de renda, ficha de salário-família, vale-transporte, exame admissional. (00202-01-C-06)
            Conhecimento: Processo de demissão: documentos necessários, modalidades de saque do FGTS, comunicação e homologação da rescisão contratual, prazos, direito do empregado conforme o tipo de contrato, cálculo das verbas rescisórias, seguro-desemprego, exame demissional. (00202-01-C-07)
            Conhecimento: Prontuário do colaborador: conceito e documentos a serem arquivados. (00202-01-C-08)
            Conhecimento: Poder disciplinar do empregador: advertência, suspensão, demissão por justa causa, demissão sem justa causa, aviso prévio, dispensa do empregado estável, apuração de faltas graves. (00202-01-C-09)
            Conhecimento: Noções da organização: conceito, tipologia, classificação, natureza jurídica, organograma (estrutura e níveis). (00202-01-C-10)
            Habilidade: Comunicar-se de forma oral e escrita com clareza e assertividade. (00202-01-H-01)
            Habilidade: Gerenciar tempo e atividades de trabalho. (00202-01-H-02)
            Habilidade: Pesquisar e organizar dados e informações. (00202-01-H-03)
            Habilidade: Resolver conflitos inerentes ao processo de trabalho. (00202-01-H-04)
            Habilidade: Trabalhar em equipe. (00202-01-H-05)
            Habilidade: Ler e interpretar textos legais. (00202-01-H-06)
            Orientações metodológicas: Os alunos devem ser orientados a realizar o planejamento de sua carreia tendo em vista a análise dos cenários de trabalho e emprego em sua região e conforme seu investimento pessoal e profissional.
            Recomenda-se ao docente no desenvolvimento das Unidades Curriculares 1, 2 e 3 que contemple atividades como dramatizações, estudos de casos, situações-problema e pesquisas, após as quais possa realizar debates e júri simulado. Esse exercício possibilita ao docente avaliar as Marcas Formativas, como o domínio técnico e científico e a visão crítica, que ficam evidentes na fundamentação de uma análise de estudo de caso, na resolução de uma situação-problema ou na defesa de um ponto de vista, bem como subsidia a avaliação do desenvolvimento das competências por meio de seus indicadores.
            Da mesma forma, trabalhos em grupo, como projetos que contemplem pesquisa, problemas a serem solucionados e propostas de melhoria, fornecem boas possibilidades de desenvolvimento das atitudes colaborativas, sustentáveis e empreendedoras.
            Na descrição de estudos de casos, dramatização de situações de trabalho, visitas técnicas e entrevistas com profissionais da área que envolvam conflitos e aspectos relacionados a diversidade de pessoas que compõem os quadros funcionais das organizações, é possível também observar as atitudes e os valores adotados pelos discentes, bem como o desenvolvimento das Marcas Formativas.
            Entrevistas com profissionais da área, bem roteirizadas pelo docente em conjunto com os discentes, realizadas em ambientes das organizações, proporcionam bons momentos de síntese e aplicação, quando seus resultados são apresentados e debatidos em sala de aula. Além disso, permitem, igualmente, boas possibilidades de desenvolvimento das Marcas Formativas e avaliação do desenvolvimento das competências.
            O docente deve proporcionar, na medida do possível, atividades em laboratórios de informática para a produção e edição de textos e de planilhas eletrônicas, efetuando as correções gramaticais necessárias.
            Para o desenvolvimento de atividades que envolvam conhecimentos relacionados às quatro operações matemáticas, sistemas de medidas e razão e proporção, deve-se trabalhar com questões contextualizadas tendo como referência a competência a ser desenvolvida na Unidade Curricular.
            Orientações metodológicas da Unidade Curricular 4 – Projeto Integrador
            Recomenda-se que o docente responsável apresente o tema gerador na primeira semana de contato com os discentes. Estes, por sua vez, devem validar o tema, podendo sugerir modificações ou acréscimos para a proposta, cabendo aos docentes avaliar juntamente com os discentes a pertinência e a viabilidade das adequações. É essencial estabelecer o cronograma de trabalho, com etapas e prazos das entregas, apresentando-o formalmente aos discentes.
            Caso o docente opte por trabalhar com um tema gerador diferente daqueles sugeridos no Plano de Curso, recomenda-se priorizar pesquisas de campo por meio de vivências, práticas, visitas técnicas, entrevistas com pessoas do mercado de trabalho, entre outros. Quando não for possível a vivência em ambiente real de trabalho, sugere-se o uso de estratégias como resolução de situações-problema e estudo de casos, por meio de recursos como vídeos, reportagens e casos fictícios baseados na realidade. As pesquisas e visitas técnicas realizadas nas demais Unidades Curriculares também servem de subsídio para o desenvolvimento do projeto.
            É fundamental que o docente responsável pelo Projeto Integrador realize seu planejamento conjuntamente com os demais docentes do curso, no sentido de incentivar a participação ativa dos envolvidos e reforçar as contribuições de cada Unidade Curricular para o Projeto. Os docentes devem acompanhar as entregas parciais conforme previsto no cronograma, auxiliando os grupos na realização e consolidação das pesquisas. É importante que todos os docentes do curso participem da elaboração, execução e apresentação dos respectivos resultados parciais e finais.
            No momento de síntese, é realizada a sistematização das informações e referências pesquisadas e das atividades desenvolvidas no decorrer do Projeto, de modo que a análise desse processo subsidie a construção das respostas e a apresentação das soluções encontradas pelos discentes. Aspectos como criatividade e inovação devem estar presentes tanto nos produtos/soluções desenvolvidos quanto na forma de apresentação dos resultados. (00202-01-O-01)
            Tecnologias sugeridas:
            APP, Formulários Online, Facilitam o preenchimento, organização e armazenamento de formulários relacionados às rotinas de admissão e demissão., ampliação, Amplia a precisão e eficiência no preenchimento e organização de formulários., 00202-01-I-02|00202-01-I-03, Facilitam a organização e armazenamento de documentos (I-02), além de ajudar na tabulação de dados e levantamento de pendências (I-03).
            IAA, Chatbots, Chatbots podem ser usados para coletar documentação e realizar agendamentos como exames admissionais e demissionais., transferência, Transfere atividades que eram feitas pelo trabalhador para a interação máquina-consumidor., 00202-01-I-04|00202-01-H-01, Facilita a comunicação (H-01) e automatiza a execução de tarefas relacionadas a exames e homologações (I-04).

            Título da UC: Acompanhar e Controlar a Entrega de Benefícios Legais e Espontâneos Concedidos Pela Organização (00202-02)
            Indicador: Recebe, protocola e encaminha a documentação necessária para que o benefício seja concedido, respeitando as normas legais e os procedimentos da organização. (00202-02-I-01)
            Indicador: Atualiza dados e informações dos colaboradores, sempre que necessário, garantindo o sigilo e o registro conforme informações recebidas. (00202-02-I-02)
            Indicador: Preenche, recolhe, organiza e arquiva documentos pertinentes ao processo de controle da entrega de benefícios legais e espontâneos, garantindo sigilo e integridade conforme requisitos estabelecidos. (00202-02-I-03)
            Indicador: Monitora e registra a entrega dos benefícios concedidos em razão de determinação legal e benefícios espontâneos concedidos pela organização, indicando aqueles colaboradores que tenham direito ao recebimento, observando os prazos para concessão, os procedimentos, as políticas internas da organização e a legislação vigente. (00202-02-I-04)
            Conhecimento: Benefícios: conceitos, fontes, princípios, características e direitos, previdência social, benefícios obrigatórios e requisitos para concessão e gozo, benefícios espontâneos. (00202-02-C-01)
            Conhecimento: Benefícios legais: férias conceito, período de férias, perda do direito de férias, abono pecuniário, parcelamento, comunicação do período de férias, pagamento das férias, adicional de 1/3 de férias, férias coletivas, encargos sociais; atraso e não concessão de férias; gratificação natalina (13º salário): cálculo, prazo e forma de pagamento, multa. (00202-02-C-02)
            Conhecimento: Demais benefícios: vale-transporte, depósito do FGTS (fundo de garantia por tempo de serviço), hora extra, hora noturna, verbas rescisórias e adicionais de periculosidade e insalubridade. (00202-02-C-03)
            Conhecimento: Benefícios espontâneos: política da organização, benefícios mais comuns. (00202-02-C-04)
            Habilidade: Comunicar-se de forma oral e escrita com clareza e assertividade. (00202-02-H-01)
            Habilidade: Gerenciar tempo e atividades de trabalho. (00202-02-H-02)
            Habilidade: Pesquisar e organizar dados e informações. (00202-02-H-03)
            Habilidade: Resolver conflitos inerentes ao processo de trabalho. (00202-02-H-04)
            Habilidade: Trabalhar em equipe. (00202-02-H-05)
            Habilidade: Ler e interpretar textos legais. (00202-02-H-06)
            Orientações metodológicas: Os alunos devem ser orientados a realizar o planejamento de sua carreia tendo em vista a análise dos cenários de trabalho e emprego em sua região e conforme seu investimento pessoal e profissional.
            Recomenda-se ao docente no desenvolvimento das Unidades Curriculares 1, 2 e 3 que contemple atividades como dramatizações, estudos de casos, situações-problema e pesquisas, após as quais possa realizar debates e júri simulado. Esse exercício possibilita ao docente avaliar as Marcas Formativas, como o domínio técnico e científico e a visão crítica, que ficam evidentes na fundamentação de uma análise de estudo de caso, na resolução de uma situação-problema ou na defesa de um ponto de vista, bem como subsidia a avaliação do desenvolvimento das competências por meio de seus indicadores.
            Da mesma forma, trabalhos em grupo, como projetos que contemplem pesquisa, problemas a serem solucionados e propostas de melhoria, fornecem boas possibilidades de desenvolvimento das atitudes colaborativas, sustentáveis e empreendedoras.
            Na descrição de estudos de casos, dramatização de situações de trabalho, visitas técnicas e entrevistas com profissionais da área que envolvam conflitos e aspectos relacionados a diversidade de pessoas que compõem os quadros funcionais das organizações, é possível também observar as atitudes e os valores adotados pelos discentes, bem como o desenvolvimento das Marcas Formativas.
            Entrevistas com profissionais da área, bem roteirizadas pelo docente em conjunto com os discentes, realizadas em ambientes das organizações, proporcionam bons momentos de síntese e aplicação, quando seus resultados são apresentados e debatidos em sala de aula. Além disso, permitem, igualmente, boas possibilidades de desenvolvimento das Marcas Formativas e avaliação do desenvolvimento das competências.
            O docente deve proporcionar, na medida do possível, atividades em laboratórios de informática para a produção e edição de textos e de planilhas eletrônicas, efetuando as correções gramaticais necessárias.
            Para o desenvolvimento de atividades que envolvam conhecimentos relacionados às quatro operações matemáticas, sistemas de medidas e razão e proporção, deve-se trabalhar com questões contextualizadas tendo como referência a competência a ser desenvolvida na Unidade Curricular.
            Orientações metodológicas da Unidade Curricular 4 – Projeto Integrador
            Recomenda-se que o docente responsável apresente o tema gerador na primeira semana de contato com os discentes. Estes, por sua vez, devem validar o tema, podendo sugerir modificações ou acréscimos para a proposta, cabendo aos docentes avaliar juntamente com os discentes a pertinência e a viabilidade das adequações. É essencial estabelecer o cronograma de trabalho, com etapas e prazos das entregas, apresentando-o formalmente aos discentes.
            Caso o docente opte por trabalhar com um tema gerador diferente daqueles sugeridos no Plano de Curso, recomenda-se priorizar pesquisas de campo por meio de vivências, práticas, visitas técnicas, entrevistas com pessoas do mercado de trabalho, entre outros. Quando não for possível a vivência em ambiente real de trabalho, sugere-se o uso de estratégias como resolução de situações-problema e estudo de casos, por meio de recursos como vídeos, reportagens e casos fictícios baseados na realidade. As pesquisas e visitas técnicas realizadas nas demais Unidades Curriculares também servem de subsídio para o desenvolvimento do projeto.
            É fundamental que o docente responsável pelo Projeto Integrador realize seu planejamento conjuntamente com os demais docentes do curso, no sentido de incentivar a participação ativa dos envolvidos e reforçar as contribuições de cada Unidade Curricular para o Projeto. Os docentes devem acompanhar as entregas parciais conforme previsto no cronograma, auxiliando os grupos na realização e consolidação das pesquisas. É importante que todos os docentes do curso participem da elaboração, execução e apresentação dos respectivos resultados parciais e finais.
            No momento de síntese, é realizada a sistematização das informações e referências pesquisadas e das atividades desenvolvidas no decorrer do Projeto, de modo que a análise desse processo subsidie a construção das respostas e a apresentação das soluções encontradas pelos discentes. Aspectos como criatividade e inovação devem estar presentes tanto nos produtos/soluções desenvolvidos quanto na forma de apresentação dos resultados. (00202-02-O-01)
            Tecnologias sugeridas:
            IOT, Sensores Inteligentes, Sensores podem emitir alertas automáticos para prazos de concessão de benefícios, assegurando que os benefícios sejam concedidos no prazo., geração, Geram alertas automáticos para prazos de concessão de benefícios., 00202-02-I-04, Acompanhamento dos prazos para concessão (I-04) é facilitado pela integração de sensores inteligentes.
            APP, Sistemas de Gerenciamento de Benefícios, Esses sistemas facilitam o controle, atualização e monitoramento da entrega de benefícios aos colaboradores., ampliação, Ampliam a eficiência na organização e controle dos benefícios concedidos., 00202-02-I-04|00202-02-C-01|00202-02-H-03, Auxilia no monitoramento dos benefícios concedidos (I-04) e envolve conhecimentos sobre benefícios (C-01), bem como a habilidade de organizar dados e informações (H-03).

            Título da UC: Auxiliar a Elaboração da Folha de Pagamento (00202-03)
            Indicador: Monitora, recebe e coleta dados e informações para atualização do cadastro de colaboradores, efetuando o seu registro no sistema operacional da organização, com agilidade e exatidão. (00202-03-I-01)
            Indicador: Recolhe, organiza e arquiva documentos pertinentes ao processo de elaboração de folha de pagamento, garantindo sigilo e integridade, conforme requisitos estabelecidos. (00202-03-I-02)
            Indicador: Atende clientes internos com cordialidade, registrando necessidades, dando retorno às solicitações, garantindo o fluxo das informações e cumprimento da legislação vigente e das normas adotadas pela organização. (00202-03-I-03)
            Indicador: Efetua, sob supervisão, cálculo dos proventos e descontos, de tributos trabalhistas e previdenciários, considerando a legislação vigente, os dados resultantes do controle de frequência e do banco de horas. (00202-03-I-04)
            Conhecimento: Controle de ponto: sanções, marcação manual, mecânica, eletrônica e magnética. (00202-03-C-01)
            Conhecimento: Salário: remuneração, adicionais, benefícios e encargos sociais, intervalos, Fundo de Garantia por Tempo de Serviço (FGTS) e PIS; contrato por prazo determinado e indeterminado. (00202-03-C-02)
            Conhecimento: Suspensão e interrupção do contrato de trabalho e ausências legais. (00202-03-C-03)
            Conhecimento: Jornada de trabalho: conceito, duração, intervalos, prorrogação, compensação de horas, escala de revezamento, categorias diferenciadas, Descanso Semanal Remunerado (DSR), trabalho noturno, jornada de trabalho mensalista, jornada de trabalho horista. (00202-03-C-04)
            Conhecimento: Processo de controle de frequência: formas de registro, empregados desobrigados da marcação, registro de serviço externo, quadro de horários, registro dos intervalos de refeição e descanso, registro das horas extras, registro de faltas e atrasos. (00202-03-C-05)
            Conhecimento: Remuneração: conceitos, diferença de remuneração e salário, adicionais (horas extras, trabalho noturno, insalubridade, periculosidade, abonos, gratificações, prêmios ou comissões, salário utilidade, gorjetas, diárias de viagem), prazo e forma para pagamento, recibo/holerite do pagamento; políticas de remuneração, prêmios e campanhas, programa motivacional. (00202-03-C-06)
            Conhecimento: Processo de elaboração da folha de pagamento: cálculo de proventos (salário, salário-família, horas extras, comissões, abonos, gratificações, prêmios, adicional de insalubridade, adicional de periculosidade, adicional noturno), afastamentos, salário-maternidade, cálculo de descontos Instituto Nacional do Seguro Social (INSS), imposto de renda, descontos para sindicato, faltas e atrasos, desconto do repouso semanal remunerado, vale-transporte, pensão alimentícia, adiantamento de salário, seguro de vida, vale--refeição e plano de saúde. (00202-03-C-07)
            Conhecimento: Obrigações mensais do empregador: recolhimento e repasse dos encargos para o INSS, Imposto de Renda Retido na Fonte (IRRF), FGTS, prestar informações ao Cadastro Geral de Empregados e Desempregados (Caged). (00202-03-C-08)
            Conhecimento: Obrigações anuais do empregador: Declaração de Informações de Rendimento na Fonte (Dirf), Relação Anual de Informações Sociais (Rais), comprovante de rendimentos pagos e de retenção de imposto de renda na fonte (informe de rendimento), Livro de Inspeção e Fiscalização do Trabalho. (00202-03-C-09)
            Conhecimento: Operações matemáticas: quatro operações básicas, porcentagem, razão e proporção. (00202-03-C-10)
            Habilidade: Comunicar-se de forma oral e escrita com clareza e assertividade. (00202-03-H-01)
            Habilidade: Gerenciar tempo e atividades de trabalho. (00202-03-H-02)
            Habilidade: Pesquisar e organizar dados e informações. (00202-03-H-03)
            Habilidade: Resolver conflitos inerentes ao processo de trabalho. (00202-03-H-04)
            Habilidade: Trabalhar em equipe. (00202-03-H-05)
            Habilidade: Ler e interpretar textos legais. (00202-03-H-06)
            Habilidade: Executar cálculos matemáticos. (00202-03-H-07)
            Orientações metodológicas: Os alunos devem ser orientados a realizar o planejamento de sua carreia tendo em vista a análise dos cenários de trabalho e emprego em sua região e conforme seu investimento pessoal e profissional.
            Recomenda-se ao docente no desenvolvimento das Unidades Curriculares 1, 2 e 3 que contemple atividades como dramatizações, estudos de casos, situações-problema e pesquisas, após as quais possa realizar debates e júri simulado. Esse exercício possibilita ao docente avaliar as Marcas Formativas, como o domínio técnico e científico e a visão crítica, que ficam evidentes na fundamentação de uma análise de estudo de caso, na resolução de uma situação-problema ou na defesa de um ponto de vista, bem como subsidia a avaliação do desenvolvimento das competências por meio de seus indicadores.
            Da mesma forma, trabalhos em grupo, como projetos que contemplem pesquisa, problemas a serem solucionados e propostas de melhoria, fornecem boas possibilidades de desenvolvimento das atitudes colaborativas, sustentáveis e empreendedoras.
            Na descrição de estudos de casos, dramatização de situações de trabalho, visitas técnicas e entrevistas com profissionais da área que envolvam conflitos e aspectos relacionados a diversidade de pessoas que compõem os quadros funcionais das organizações, é possível também observar as atitudes e os valores adotados pelos discentes, bem como o desenvolvimento das Marcas Formativas.
            Entrevistas com profissionais da área, bem roteirizadas pelo docente em conjunto com os discentes, realizadas em ambientes das organizações, proporcionam bons momentos de síntese e aplicação, quando seus resultados são apresentados e debatidos em sala de aula. Além disso, permitem, igualmente, boas possibilidades de desenvolvimento das Marcas Formativas e avaliação do desenvolvimento das competências.
            O docente deve proporcionar, na medida do possível, atividades em laboratórios de informática para a produção e edição de textos e de planilhas eletrônicas, efetuando as correções gramaticais necessárias.
            Para o desenvolvimento de atividades que envolvam conhecimentos relacionados às quatro operações matemáticas, sistemas de medidas e razão e proporção, deve-se trabalhar com questões contextualizadas tendo como referência a competência a ser desenvolvida na Unidade Curricular.
            Orientações metodológicas da Unidade Curricular 4 – Projeto Integrador
            Recomenda-se que o docente responsável apresente o tema gerador na primeira semana de contato com os discentes. Estes, por sua vez, devem validar o tema, podendo sugerir modificações ou acréscimos para a proposta, cabendo aos docentes avaliar juntamente com os discentes a pertinência e a viabilidade das adequações. É essencial estabelecer o cronograma de trabalho, com etapas e prazos das entregas, apresentando-o formalmente aos discentes.
            Caso o docente opte por trabalhar com um tema gerador diferente daqueles sugeridos no Plano de Curso, recomenda-se priorizar pesquisas de campo por meio de vivências, práticas, visitas técnicas, entrevistas com pessoas do mercado de trabalho, entre outros. Quando não for possível a vivência em ambiente real de trabalho, sugere-se o uso de estratégias como resolução de situações-problema e estudo de casos, por meio de recursos como vídeos, reportagens e casos fictícios baseados na realidade. As pesquisas e visitas técnicas realizadas nas demais Unidades Curriculares também servem de subsídio para o desenvolvimento do projeto.
            É fundamental que o docente responsável pelo Projeto Integrador realize seu planejamento conjuntamente com os demais docentes do curso, no sentido de incentivar a participação ativa dos envolvidos e reforçar as contribuições de cada Unidade Curricular para o Projeto. Os docentes devem acompanhar as entregas parciais conforme previsto no cronograma, auxiliando os grupos na realização e consolidação das pesquisas. É importante que todos os docentes do curso participem da elaboração, execução e apresentação dos respectivos resultados parciais e finais.
            No momento de síntese, é realizada a sistematização das informações e referências pesquisadas e das atividades desenvolvidas no decorrer do Projeto, de modo que a análise desse processo subsidie a construção das respostas e a apresentação das soluções encontradas pelos discentes. Aspectos como criatividade e inovação devem estar presentes tanto nos produtos/soluções desenvolvidos quanto na forma de apresentação dos resultados. (00202-03-O-01)
            Tecnologias sugeridas:
            IAA, Assistente Virtual de Voz, Assistentes virtuais de voz ajudam a coletar informações para atualização de cadastro e registro no sistema operacional., ampliação, Ampliam a eficiência e precisão na coleta de dados e informações para cadastro., 00202-03-I-01|00202-03-I-03|00202-03-H-01|00202-03-C-02, Assistentes virtuais podem ajudar na coleta e atualização de dados de cadastro (I-01) e no atendimento aos clientes internos (I-03), utilizando conhecimento sobre salário e benefícios (C-02) e habilidades de comunicação assertiva (H-01).
            APP, Sistemas de Gerenciamento de Folha de Pagamento, Facilitam o cálculo e organização da folha de pagamento, considerando proventos, descontos e tributos., ampliação, Ampliam a eficiência e precisão na elaboração da folha de pagamento., 00202-03-I-04|00202-03-C-07|00202-03-H-03|00202-03-H-07, Facilitam o cálculo de proventos e descontos (I-04), envolvendo conhecimentos sobre folha de pagamento (C-07) e habilidades de organização de dados (H-03) e execução de cálculos matemáticos (H-07).

            Curso: Cuidador de Idoso (2454) - 2018

            Título da UC: Estimular a Independência e Autonomia do Idoso em Suas Atividades de Vida Diária (00140-01)
            Indicador: Acompanha e auxilia o idoso nas suas atividades de vida diária, respeitando as diferenças individuais do processo de envelhecimento. (00140-01-I-01)
            Indicador: Realiza ações de prevenção frente a situações de vulnerabilidade social, psicológica e física a que estão expostas o idoso, considerando as diretrizes do Estatuto do Idoso, de acordo com as orientações da equipe multiprofissional. (00140-01-I-02)
            Indicador: Orienta atividades de lazer e de ocupação do tempo livre de acordo com interesse do idoso e orientações da equipe multiprofissional. (00140-01-I-03)
            Indicador: Providencia adaptações no ambiente de acordo com as orientações da equipe multiprofissional, condições de mobilidade, segurança e valores do idoso e da família. (00140-01-I-04)
            Conhecimento: Legislação vigente relacionada ao idoso: Estatuto do Idoso (Lei nº 10.741 de outubro de 2003 e suas atualizações) conceitos, direitos e deveres do idoso; Política Municipal do Idoso - especificidades locais (como definição de idade para acesso aos direitos); Lei Orgânica de Assistência Social (LOAS) e Política Nacional de Saúde da Pessoa Idosa: princípios e diretrizes. (00140-01-C-01)
            Conhecimento: Política Nacional de Humanização: conceito, diretrizes, definição de acolhimento. (00140-01-C-02)
            Conhecimento: Envelhecimento ativo: definição e determinantes - culturais, comportamentais, pessoais, sociais e econômicos. (00140-01-C-03)
            Conhecimento: Perfil da população idosa: aspectos demográficos do envelhecimento; preconceitos, mitos e estereótipos; senescência e senilidade: processos naturais e patológicos do envelhecimento; diversidade no processo de envelhecimento; conceito de identidade, curso de vida e life span; espiritualidade e valores culturais; trabalho e aposentadoria; aspectos da sexualidade na velhice; infecções sexualmente transmissíveis: dados epidemiológicos de prevalência na população idosa, prevenção e cuidados. (00140-01-C-04)
            Conhecimento: Mercado de trabalho para o Cuidador de Idoso: atribuições, campos e limites de atuação; Classificação Brasileira de Ocupações - CBO (5162-10), Projetos de Lei vigentes; empregabilidade, empreendedorismo, apresentação pessoal, e planejamento de carreira. (00140-01-C-05)
            Conhecimento: O Cuidador de Idoso: perfil profissional principais características (escuta ativa, paciência ativa, controle emocional, organização e planejamento, atuação ética como elo entre idoso, família, sociedade e equipe multiprofissional); saúde do Cuidador - qualidade de vida, ergonomia, vacinação, cuidados e acompanhamento de saúde física e psicológica, autocuidado, fatores de risco (estresse, tabagismo, obesidade, automedicação, alcoolismo, sedentarismo). (00140-01-C-06)
            Conhecimento: O Cuidador de Idoso e a equipe multiprofissional: médico, enfermeiro, técnico e auxiliar de enfermagem, fisioterapeuta, terapeuta ocupacional, psicólogo, nutricionista, fonoaudiólogo, odontólogo, educador físico, gerontólogo, assistente social, farmacêutico; funções e acesso aos profissionais; relações de trabalho; busca por orientação e limites de atuação. (00140-01-C-07)
            Conhecimento: Atividades de vida diária (AVDs): definição, classificação e tipos: básicas (relacionadas ao autocuidado, como banhar-se, alimentar-se, vestir-se, caminhar), instrumentais (mantém o idoso ativo na comunidade como fazer compras, utilizar transporte, preparar refeições, ir ao banco) e avançadas (realizar viagens, planejamento financeiro, atividades de ocupação do tempo livre). (00140-01-C-08)
            Conhecimento: Autocuidado do idoso: definição, finalidade e ações para o desenvolvimento. (00140-01-C-09)
            Conhecimento: Dependência, independência e autonomia do idoso: conceitos, ações e importância do estímulo e manutenção. (00140-01-C-10)
            Conhecimento: Fatores de risco para pessoa idosa, prevenção e cuidados: sedentarismo, obesidade, quedas, atividades físicas, deficiências visual e auditiva, amputações, consumo de drogas lícitas e ilícitas, desatualização vacinal. (00140-01-C-11)
            Conhecimento: Mobilidade funcional reduzida: conceito, ações para o desenvolvimento da mobilidade, equipamentos relacionados (bengala, muletas, andador, cadeira de rodas, entre outros), fatores de riscos para quedas e prevenção. (00140-01-C-12)
            Conhecimento: Ambiência: conceito, mobilidade, segurança, higiene, organização e adaptações no ambiente respeitando valores morais, culturais, éticos e religiosos. (00140-01-C-13)
            Conhecimento: Comunicação verbal e não verbal: barreiras comunicacionais (idioma, inibição, estereótipo, deficiências, agressividade, dentre outros), estratégias de comunicação com o idoso. (00140-01-C-14)
            Conhecimento: Relação de ajuda: conceito, importância e como construí-la. (00140-01-C-15)
            Conhecimento: Vulnerabilidade: social, psicológica e física definições, ações de prevenção e encaminhamentos. (00140-01-C-16)
            Conhecimento: Violência contra o idoso: tipos; indicadores de maus tratos; encaminhamentos. (00140-01-C-17)
            Conhecimento: Programas voltados ao idoso disponíveis no Sistema Único de Saúde (SUS) e no Sistema Único de Assistência Social (SUAS): definições, ações e acesso. (00140-01-C-18)
            Conhecimento: Qualidade de vida do idoso: organização da rotina, sono e repouso, alimentação, hidratação, higiene e imagem pessoal, cuidados estéticos; aspectos culturais e valores morais e espirituais (gostos, preferências, hábitos); atividades de convívio social: equipamentos de lazer do município, ocupação do tempo livre, atividades de lazer e entretenimento; uso de aplicativos, softwares e jogos eletrônicos: cuidados, indicações e orientações da equipe multiprofissional. (00140-01-C-19)
            Conhecimento: Hospitalidade: conceito, relação hóspede e anfitrião, regras de convivência. (00140-01-C-20)
            Habilidade: Interpretar as orientações da equipe multiprofissional. (00140-01-H-01)
            Habilidade: Comunicar-se de maneira assertiva. (00140-01-H-02)
            Habilidade: Auxiliar na locomoção e movimentação do idoso. (00140-01-H-03)
            Habilidade: Organizar o ambiente de permanência do idoso. (00140-01-H-04)
            Habilidade: Identificar situações de risco para o idoso. (00140-01-H-05)
            Habilidade: Mediar conflitos nas situações de trabalho. (00140-01-H-06)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na residência. (00140-01-H-07)
            Orientações metodológicas: Para esta UC, que tem foco nos aspectos sociais e de convivência, bem como no resgate da independência e autonomia dos idosos, recomenda-se que sejam planejadas situações de aprendizagem que promovam a vivência dos fazeres profissionais voltados à prática do Cuidador, dentro de seus limites de atuação, voltadas ao:
                    Acompanhamento do idoso nas suas atividades de vida diária;
                    Acompanhamento e inserção do idoso nas políticas sociais e de saúde;
                    Realização de atividades não terapêuticas de lazer e de ocupação do tempo livre, de interesse do idoso; 
                    Sugestão de adaptações no ambiente melhorando a mobilidade e segurança, respeitando os valores do idoso; 
                    Identificação, prevenção, comunicação e busca de auxilio em situações de risco de vulnerabilidade física, psicológica e social.
            Para tal, podem ser propostas atividades como pesquisas, estudos de situações-problema, simulações, visitas técnicas em espaços de lazer públicos e instituições de longa permanência com o objetivo de identificar adaptações e possibilidades de inclusão do idoso em atividades de ocupação do tempo livre, minimizando os fatores de risco estimulando a independência e autonomia. Também podem ser exibidos documentários e veiculados materiais de pesquisa para discussões sobre o processo de envelhecimento, políticas públicas e estudos de casos visando a discussão das atribuições do Cuidador e seu papel na equipe multiprofissional, além de entrevistas com idosos para reconhecer quais mitos e estereótipos impactam mais nas atividades de vida diária. Outra possibilidade é a promoção de encontros com profissionais que dão suporte ao idoso como geriatra, gerontólogo, fisioterapeuta, terapeuta ocupacional, assistente social, educador físico, entre outros, a fim de contextualizar a equipe multiprofissional.
            As questões comportamentais e valorativas, extremamente relevantes a essa ocupação para que o Cuidador se reconheça como elo entre o idoso, a família, sociedade e equipe multiprofissional, podem ser exploradas em atividades como estudos de situações-problema, simulações, dinâmicas de grupo, palestras com profissionais da área, desenvolvimentos de ações socioeducativas.   (00140-01-O-01)
            Tecnologias sugeridas:
            IAA, Assistente Virtual de Voz, Assistentes virtuais de voz são tecnologias adequadas para ajudar no acompanhamento e auxílio nas atividades de vida diária do idoso, fornecendo lembretes e orientações., transferência, Transfere parte das atividades para a interação entre cliente e máquina., 00140-01-I-01|00140-01-I-03|00140-01-H-02|00140-01-C-08|00140-01-C-09|00140-01-C-10, Assistentes virtuais podem ajudar nas suas atividades de vida (I-01) e de ocupação do tempo livre (I-03), com a comunicação assertiva (H-02) e conhecimentos ligados à AVDs (C-08), autocuidado (C-09) e autonomia (C-10).
            IOT, Sensores Inteligentes, Sensores que monitoram a movimentação e condições ambientais são cruciais para garantir segurança e adaptar o ambiente., geração, Geram atividades automatizadas de monitoramento contínuo da segurança do idoso., 00140-01-I-01|00140-01-I-02|00140-01-I-04|00140-01-C-12|00140-01-C-16|00140-01-H-05, Ajuda a monitorar continuamente a movimentação do idoso (I-01/I-02/H-05) e providenciar adaptações no ambiente (I-04), utilizando conhecimentos de mobilidade funcional (C-12) e vulnerabilidade física (C-16).

            Título da UC: Cuidar da Pessoa Idosa em Suas Atividades de Vida Diária (00140-02)
            Indicador: Monitora o estado de saúde do idoso de acordo com orientações da equipe multiprofissional e premissas do cuidado humanizado. (00140-02-I-01)
            Indicador: Realiza a higiene corporal e bucal de acordo com autonomia e independência do idoso, respeitando as orientações da equipe multiprofissional. (00140-02-I-02)
            Indicador: Proporciona medidas de conforto de acordo com autonomia e independência do idoso, respeitando as orientações da equipe multiprofissional. (00140-02-I-03)
            Indicador: Auxilia na alimentação de acordo com autonomia e independência do idoso, respeitando as orientações da equipe multiprofissional. (00140-02-I-04)
            Indicador: Auxilia o idoso na administração de medicamentos de acordo com orientação e prescrição. (00140-02-I-05)
            Indicador: Registra a rotina e comunica a família e equipe multiprofissional, considerando as alterações do estado de saúde do idoso. (00140-02-I-06)
            Indicador: Presta cuidados de primeiros socorros de acordo com a situação de emergência, solicitando auxílio aos serviços ou profissionais especializados, quando necessário. (00140-02-I-07)
            Conhecimento: Medidas de segurança: contaminação, infecção e infecção cruzada - conceitos e responsabilidades; higienização das mãos - definição, finalidade e passo a passo de acordo com o Ministério da Saúde; equipamentos de proteção individual: conceito, tipos (luvas, máscaras, avental), indicação e formas de utilização. (00140-02-C-01)
            Conhecimento: Descarte de resíduos: tipos de resíduos e destino. (00140-02-C-02)
            Conhecimento: Equipamentos e materiais utilizados pelo Cuidador de Idoso finalidade, tipos, utilização e limites de atuação: termômetro digital, aparelho de pressão arterial digital, bolsa térmica, compressas, comadre, papagaio, dispositivo para incontinência urinária masculino, bolsa coletora de urina, bolsa de colostomia, fralda, luva para higiene corporal, cadeira higiênica, cama, colchões e lençóis, aspirador nasal, novas tecnologias facilitadoras para atividades diárias de vida (como porta comprimido digital, aplicativos de agenda do idoso). (00140-02-C-03)
            Conhecimento: Equipamentos e materiais apresentados e orientados pelos serviços de saúde: seringas, agulhas e caneta de aplicação de insulina; aparelho de glicemia capilar; sonda enteral e de gastrostomia; equipos e frascos para dietas. (00140-02-C-04)
            Conhecimento: Monitoramento do estado de saúde do idoso: cuidados com a pele (hidratação, prevenção de lesões, curativos simples para escoriações, feridas pequenas, não profundas e sem secreções); aplicação de calor e frio (manuseio da bolsa térmica, tempo e cuidados na aplicação); lesão por pressão (conceito, sinais de identificação, prevenção); controles da pressão arterial e temperatura corporal (parâmetros de normalidade, procedimento de verificação, registro, informação das alterações), glicemia capilar (parâmetros de normalidade). (00140-02-C-05)
            Conhecimento: Senilidade e síndromes geriátricas - varizes, hipertensão, diabetes, desnutrição, desidratação, catarata, pneumonia, incontinência urinária e fecal, infecção urinária, artrose, Parkinson, Acidente Vascular Encefálico (AVE), demências (declínio cognitivo leve, Alzheimer, Corpos de Lewy, demência frontotemporal): definição, principais sinais e sintomas, cuidados relacionados. (00140-02-C-06)
            Conhecimento: Comunicação para idoso com demência: características específicas e abordagem. (00140-02-C-07)
            Conhecimento: Medidas de higiene: definição, tipos, finalidade e execução - cuidado corporal: cabelos, unhas, pele, barba, banho (chuveiro, banheira e na cama), higiene oral, higiene íntima, remoção de secreções de boca e nariz. (00140-02-C-08)
            Conhecimento: Medidas de conforto: definição, tipos, finalidade, execução, cuidados relacionados e anotação - hidratação da pele, mudança da posição do corpo, utilização de almofadas e travesseiros, organização do ambiente, cuidados pelos objetos do idoso, transferência cama/cadeira higiênica, limpeza de boca e nariz, eliminações fisiológicas (troca de fralda, esvaziamento e troca de bolsa de colostomia, esvaziamento de bolsa coletora de urina). (00140-02-C-09)
            Conhecimento: Nutrição, hidratação e alimentação: propriedades nutricionais dos alimentos (energéticos, construtores e reguladores), vias (oral, nasoenteral); tipos de dieta (hipossódica, pastosa, líquida, específica para diabético, hipogordurosa, entre outras), restrições hídricas e alimentares; estímulo e auxílio na hidratação e alimentação por via oral. (00140-02-C-10)
            Conhecimento: Cuidados com medicação: medicamentos - diferenças entre comercial (ou de referência), genérico e similar; administração de medicamentos - leitura e interpretação da prescrição, tempo de tratamento, horários, dosagem, vias (oral, retal, tópica, instilação ocular, nasal e inalatória), acondicionamento, validade e descarte; riscos da automedicação - medicamentos industrializados, “caseiros” e fitoterápicos. (00140-02-C-11)
            Conhecimento: Cuidados paliativos, finitude e morte: conceitos segundo OMS; cuidados relacionados à finitude; aspectos culturais e crenças, processo de luto, providencias em relação à morte em domicilio. (00140-02-C-12)
            Conhecimento: Primeiros socorros - prevenção de acidentes, precauções e atendimento em caso de engasgo, intoxicação, queda, reações alérgicas, envenenamento, queimadura, convulsão, desmaio, hemorragia, fraturas, choque elétrico e parada cardiorrespiratória. (00140-02-C-13)
            Conhecimento: Tipos de serviços de urgência e emergência disponíveis na comunidade: Serviço de Atendimento Móvel de Urgência (Samu), Corpo de Bombeiro, atendimento de emergência particular, Polícia Militar, entre outros. (00140-02-C-14)
            Habilidade: Comunicar-se de maneira assertiva. (00140-02-H-01)
            Habilidade: Interpretar prescrição e orientações da equipe multiprofissional. (00140-02-H-02)
            Habilidade: Atentar-se a comportamentos, reações, sinais e sintomas do idoso. (00140-02-H-03)
            Habilidade: Utilizar materiais e equipamentos. (00140-02-H-04)
            Habilidade: Identificar situações de emergência e de risco. (00140-02-H-05)
            Habilidade: Adotar boas práticas de higiene no controle e prevenção de doenças. (00140-02-H-06)
            Habilidade: Promover condições para as eliminações fisiológicas. (00140-02-H-07)
            Habilidade: Verificar temperatura e pressão arterial. (00140-02-H-08)
            Orientações metodológicas: Para esta UC, que tem foco nas questões relacionadas à saúde do idoso, recomenda-se que sejam planejadas situações de aprendizagem que remetam à vivência dos fazeres profissionais voltados às ações de cuidado, acompanhamento da saúde e prevenção de doenças, considerando os limites de atuação do Cuidador, como:
                    Mensuração de pressão arterial com aparelho de pulso digital.
                    Verificação de temperatura.
                    Realização de curativos simples.
                    Aplicação de bolsa de calor e frio.
                    Remoção de secreções da boca e do nariz.
                    Auxilio ou administração de medicamentos por vias oral, auricular, nasal, retal, oftálmica, inalações e tópica. 
                    Acompanhamento ou oferecimento de alimentação por boca. 
                    Coleta de urina e fezes, colocação e retirada de comadre e papagaio, colocação de dispositivo para incontinência urinária masculino, esvaziamento de bolsa coletora de urina, realização de troca de fraldas, troca e higiene da bolsa de colostomia.
                    Realização de banho na cama, higiene oral, íntima e couro cabeludo.
                    Realização de mudanças de posição do corpo para conforto e prevenção de lesão por pressão.
                    Registro da rotina do idoso e das intercorrências.
                    Prestação de cuidados de primeiros socorros.
            Para tal, recomenda-se que o docente demonstre os fazeres referentes aos cuidados desta UC e que os alunos os simulem, utilizando manequins específicos para esta finalidade, além de propor atividades como estudos de situações-problema, pesquisas, discussão de vídeos, palestras com profissionais e representantes comerciais para demonstração de equipamentos, materiais e tecnologias assistivas. Podem ser incluídas visitas técnicas à ILPI e Centro Dia estimulado o voluntariado aos alunos para realizar atividades com os idosos.
            É importante salientar a necessidade de o docente constantemente explorar os limites de atuação com os diversos atores da equipe multidisciplinar de saúde que acompanham o idoso. Nesse sentido, no que se refere à mensuração da glicemia capilar, aplicação subcutânea de insulina e administração de alimentação por sonda, o docente deve orientar o aluno, caso se depare com essas demandas durante sua vida profissional, a buscar orientação da equipe de saúde nos diversos segmentos, como ambulatório de especialidades, Estratégia da Saúde da Família, Unidade Básica de Saúde (UBS), consultórios médicos e hospitais.
            Da mesma forma que na UC1, as questões comportamentais e valorativas podem ser exploradas em atividades como estudos de situações-problema, simulações, dinâmicas de grupo, palestras com profissionais da área, desenvolvimentos de ações socioeducativas.   (00140-02-O-01)
            Tecnologias sugeridas:
            APP, Formulários Online, Facilitam o registro de rotina e comunicação de alterações do estado de saúde do idoso diretamente em dispositivos móveis., ampliação, Amplia a precisão e comunicação dos registros com a equipe multiprofissional e familiares., 00140-02-I-06|00140-02-C-05|00140-02-H-01, Atividades ligadas ao registro (I-06), incluindo a habilidade de comunicação assertiva (H-01) e monitoramento do estado de saúde do idoso (C-05).

            Curso: Desenvolvedor Front-End (2824) - 2022

            Título da UC: Elaborar Projetos de Aplicações para Web (00074-01)
            Indicador: Define os objetivos do projeto, de acordo com as necessidades do cliente e público-alvo. (00074-01-I-01)
            Indicador: Elabora proposta de trabalho de acordo com arquitetura da informação da aplicação web e das estratégias tecnológicas. (00074-01-I-02)
            Indicador: Elabora protótipo para websites de acordo com o briefing. (00074-01-I-03)
            Indicador: Redige o projeto da aplicação web de acordo com a proposta e protótipo. (00074-01-I-04)
            Conhecimento: Briefing: conceito, especificidades para projetos web e modelos. (00074-01-C-01)
            Conhecimento: Técnicas de criatividade: brainstorming, mapas mentais e painéis semânticos. (00074-01-C-02)
            Conhecimento: Análise de mercado: identificação do perfil do cliente e público-alvo, concorrentes diretos e indiretos. (00074-01-C-03)
            Conhecimento: Domínios de internet: conceito, registro e disponibilidade de serviços. (00074-01-C-04)
            Conhecimento: Projetos web: tendências, tecnologias, gestão de projetos (custos, calendários de tarefas e relatórios de acompanhamento). (00074-01-C-05)
            Conhecimento: Arquitetura da informação: conceito e aplicações, mapa do site e estruturas de navegação, organização de conteúdo. (00074-01-C-06)
            Conhecimento: Proposta comercial: características, requisitos, elaboração e modelos de contrato de serviço. (00074-01-C-07)
            Conhecimento: Requisitos: conceitos e técnicas para análise e gerenciamento. (00074-01-C-08)
            Habilidade: Interpretar briefing para projetos de websites. (00074-01-H-01)
            Habilidade: Identificar requisitos técnicos para projetos de websites. (00074-01-H-02)
            Habilidade: Pesquisar domínio disponível e serviços de hospedagem para websites. (00074-01-H-03)
            Habilidade: Organizar arquivos e atividades por etapas do projeto. (00074-01-H-04)
            Habilidade: Estruturar arquitetura dos elementos de conteúdo de websites. (00074-01-H-05)
            Orientações metodológicas: Nesta unidade curricular, sugere-se a experimentação do fluxo de planejamento de um website, desde as análises de mercado, público-alvo, levantamento de informações através do briefing e definição de arquitetura da informação, utilizando de exercícios práticos e estudo de casos que contemplem, preferencialmente, situações reais de mercado. (00074-01-O-01)
            Tecnologias sugeridas:
            IAA, IA Generativa, IA generativa pode ser utilizada para criar propostas de trabalho e protótipos automaticamente com base no briefing., ampliação, Amplia a produtividade ao gerar automaticamente protótipos e propostas de trabalho., 00074-01-I-03|00074-01-I-04|00074-01-C-02|00074-01-H-01, Contribui na elaboração dos protótipos (I-03) e na redação do projeto (I-04), utilizando técnicas de criatividade (C-02) e habilidades de interpretar o briefing (H-01).
            APP, Sistemas de Gerenciamento de Tarefas e Projetos, Plataformas como Trello, Asana ou Monday ajudam a organizar e acompanhar as etapas do projeto de aplicação web., ampliação, Expande a capacidade de organização e sincronização das atividades do projeto., 00074-01-I-02|00074-01-I-04|00074-01-C-05|00074-01-H-04, A tecnologia facilita o acompanhamento das estratégias tecnológicas (I-02) e a organização das atividades do projeto (H-04), melhorando a gestão do projeto (C-05) e a redação do projeto da aplicação web (I-04).
            AD, Business Intelligence (BI), BI pode auxiliar na análise de mercado, identificando o perfil do cliente e concorrentes., geração, Gera novas atividades de análise de dados para criação de insights baseados em métricas., 00074-01-I-01|00074-01-I-02|00074-01-C-03, Auxilia na definição dos objetivos do projeto (I-01) e na elaboração da proposta de trabalho (I-02), articulando conhecimentos sobre análise de mercado (C-03).

            Título da UC: Desenvolver Aplicações para Websites (00074-02)
            Indicador: Cria imagens, layouts e animações otimizadas para website de acordo com os princípios de comunicação visual, normas e tendências de mercado. (00074-02-I-01)
            Indicador: Utiliza linguagem de marcação de conteúdo e estilo, de acordo com as normas e padrões tecnológicos. (00074-02-I-02)
            Indicador: Testa padrões de acessibilidade e usabilidade do website de acordo com as normas. (00074-02-I-03)
            Indicador: Analisa inconsistências no funcionamento do website, de acordo com os padrões W3C. (00074-02-I-04)
            Conhecimento: Imagem digital: Conceitos de vetor e bitmap, formatos e aplicações, unidades de medida, densidade de pixels, taxa de bits, animações web. (00074-02-C-01)
            Conhecimento: Comunicação visual: Tipografia, teoria e modos de cor, grid e alinhamento, Gestalt (conceito e aplicabilidade). (00074-02-C-02)
            Conhecimento: Wireframes e protótipos: conceito, estrutura e desenvolvimento. (00074-02-C-03)
            Conhecimento: Design responsivo: conceitos e aplicações. (00074-02-C-04)
            Conhecimento: Projeto de website: características funcionais, usabilidade, acessibilidade e ergonomia. (00074-02-C-05)
            Conhecimento: Estrutura semântica: conceitos, linguagem de marcação de conteúdo. (00074-02-C-06)
            Conhecimento: Estilização de páginas: definições de estilos e integração com estrutura de conteúdo. (00074-02-C-07)
            Conhecimento: Web standards: boas práticas e padrões recomendados pelo W3C. (00074-02-C-08)
            Conhecimento: Desempenho e compatibilidade: conceito e ferramentas. (00074-02-C-09)
            Conhecimento: Código do website: testes, validação e correção. (00074-02-C-10)
            Conhecimento: Content Management System (CMS) - Interface, requisitos básicos, servidor local e banco de dados, configuração de usuários administrativos, customização, CSS e temas, plug-ins, implantação e publicação. (00074-02-C-11)
            Conhecimento: Lei Geral de Proteção de Dados Pessoais: Fundamentos, aplicabilidade, princípios legais, direitos, transferência internacional de dados, agentes de tratamento, encarregado pelo tratamento de dados pessoas (DPO), segurança, boas práticas, fiscalização e penalidades e a definição e papel da ANPD Agencia Nacional de Proteção de Dados. (00074-02-C-12)
            Habilidade: Aplicar o sistema de cor pertinente ao layout da aplicação. (00074-02-H-01)
            Habilidade: Organizar conteúdo visual e textual para web. (00074-02-H-02)
            Habilidade: Integrar linguagens de estilo e marcação de conteúdo. (00074-02-H-03)
            Habilidade: Utilizar ferramentas para análise de desempenho. (00074-02-H-04)
            Habilidade: Testar compatibilidade nos diversos navegadores. (00074-02-H-05)
            Orientações metodológicas: Nesta unidade curricular, sugere-se a experimentação de técnicas de composição visual e codificação para o desenvolvimento de websites, através de exercícios que evidenciem conceitos como imagem digital, sistemas de cor, tipografia, grids, alinhamentos, wireframes, design responsivo, linguagens de marcação e estilo e padrões web.
            Esta UC já traz a necessidade de execução de testes que serão realizados ao longo das demais competências do curso.
            As atividades ao longo do Curso com base em situações reais devem estar pautadas em demandas atuais e da região/localidade onde os alunos estão inseridos.  (00074-02-O-01)
            Tecnologias sugeridas:
            IAA, IA Generativa, IA generativa pode ser utilizada para criar rapidamente imagens, layouts e animações otimizadas para websites., substituição, Substitui a criação manual de imagens e layouts., 00074-02-I-01|00074-02-H-01|00074-02-H-02|00074-02-C-01, Facilita a criação rápida de imagens (I-01), aplicando sistemas de cores (H-01) e organizando o conteúdo (H-02), articulando conhecimentos de imagem digital (C-01).
            APP, Plataformas de Comunicação e Colaboração Online, Ferramentas como Slack, Trello e Asana facilitam a comunicação assertiva e o planejamento das etapas do projeto., ampliação, Amplia a capacidade de organização das atividades do projeto, aumentando a eficiência através da colaboração., 00074-02-I-03|00074-02-I-04|00074-02-H-04|00074-02-C-05, Facilita a comunicação e organização das etapas do projeto (I-03/I-04), aplicando ferramentas de análise de desempenho (H-04) e conhecimentos sobre projeto de website (C-05).

            Título da UC: Codificar Front-End de Aplicações Web (00074-03)
            Indicador: Cria blocos de back-end, utilizando linguagens de programação de script, de acordo com os requisitos do projeto de software. (00074-03-I-01)
            Indicador: Manipula os elementos estruturais, de acordo com os requisitos do projeto. (00074-03-I-02)
            Indicador: Programa comportamentos dinâmicos, definindo estilos e animações, de acordo com requisitos do projeto e padrões de acessibilidade e usabilidade. (00074-03-I-03)
            Indicador: Implementa frameworks de acordo com as necessidades do projeto. (00074-03-I-04)
            Indicador: Comunica requisições conforme recursos disponibilizados pelo back-end. (00074-03-I-05)
            Indicador: Implementa correções e melhorias de acordo com a depuração de código de script. (00074-03-I-06)
            Indicador: Implementa usabilidade e acessibilidade a páginas web de acordo com os padrões do W3C. (00074-03-I-07)
            Conhecimento: Sites estáticos e dinâmicos: conceitos, diferenças, linguagens de script, tecnologias e aplicações. (00074-03-C-01)
            Conhecimento: Fundamentos de Lógica de Programação: conceitos de algoritmos, de entradas e saídas, manipulação e processamento de dados. (00074-03-C-02)
            Conhecimento: Programação: conceito, análise de requisitos do projeto e a relação com a codificação em front-end. (00074-03-C-03)
            Conhecimento: Linguagem de Scripts: Sintaxe - operadores, palavras reservadas, identificadores, delimitadores e comentários; variáveis e tipos de dados, estruturas de controle condicional e laços de repetição. (00074-03-C-04)
            Conhecimento: Eventos e funções: parâmetros, retornos e tipos de eventos. (00074-03-C-05)
            Conhecimento: Document Object Model (DOM): objetos, propriedades e eventos; manipulação de elementos, atribuição de eventos e estilos dinâmicos. (00074-03-C-06)
            Conhecimento: Framework: conceitos, aplicabilidade e tendências. (00074-03-C-07)
            Conhecimento: Usabilidade e acessibilidade: princípios aplicados ao comportamento dinâmico da página. (00074-03-C-08)
            Conhecimento: Requisições assíncronas: conceitos e aplicabilidade. (00074-03-C-09)
            Conhecimento: Depuração de linguagem de scripts: de bugger e testes de código. (00074-03-C-10)
            Habilidade: Interpretar requisitos de projetos. (00074-03-H-01)
            Habilidade: Identificar estruturas e funcionalidades da linguagem de script. (00074-03-H-02)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00074-03-H-03)
            Habilidade: Utilizar padrões de boas práticas no desenvolvimento do site. (00074-03-H-04)
            Habilidade: Aplicar responsividade no comportamento do website. (00074-03-H-05)
            Habilidade: Interpretar manuais e documentações técnicas. (00074-03-H-06)
            Habilidade: Indentar códigos de script. (00074-03-H-07)
            Orientações metodológicas: Nesta unidade curricular, sugere-se atividades de codificação utilizando as melhores práticas de desenvolvimento de aplicações para websites, com o JavaScript, estimulando sempre o aprender fazendo e aproximando o curso da realidade na qual os conhecimentos aprendidos poderão ser utilizados. Outra pesquisa poderá ser feita em bibliotecas que utilizem esta linguagem, porém é preciso verificar as bibliotecas atuais e que tenham demanda na região/localidade. 
            O trabalho com algoritmo que será realizado na UC6 “Desenvolver algoritmos” deve iniciar sua contextualização nesta UC, pois será visto uma parte de programação e lógica para seu uso na programação em JavaScript, de maneira que o aluno esteja habilitado a programar, ao menos, pequenas rotinas em script.
            Indica-se que, para os estudos de funcionalidades assíncronas, sejam disponibilizados recursos previamente programados em back-end que seja acessível e que retornem valores adequados para teste. As competências que envolvem back-end serão desenvolvidas posteriormente. 
            A continuidade da aprendizagem na aplicação de testes que iniciou na UC2 está presente ao longo do desenvolvimento desta e das demais competências do curso. (00074-03-O-01)
            Tecnologias sugeridas:
            RE, Realidade Aumentada (AR), AR pode ser usada para visualizar e testar os elementos programados no contexto real, melhorando a precisão., ampliação, Amplia a precisão e qualidade ao testar elementos em ambiente simulado., 00074-03-I-03|00074-03-I-06|00074-03-H-02|00074-03-C-06, Melhora a precisão durante a programação de comportamentos dinâmicos (I-03) e a implementação de melhorias (I-06), utilizando técnicas de manipulação de elementos (C-06) e identificando estruturas de script (H-02).
            IAA, IA Generativa, Assistentes de IA generativa ajudam na criação de códigos dinâmicos e na aplicação de comportamento responsivo nos websites., ampliação, Amplia a produção ao automatizar a criação de códigos de script e comportamentos dinâmicos., 00074-03-I-03|00074-03-I-07|00074-03-H-05|00074-03-C-08, Facilita a programação de comportamentos dinâmicos (I-03) e a acessibilidade das páginas web (I-07), aplicando conceitos de usabilidade e acessibilidade (C-08) e habilidades de responsividade (H-05).

            Título da UC: Publicar Aplicações Web (00074-04)
            Indicador: Define serviço de hospedagem de acordo com os objetivos do projeto. (00074-04-I-01)
            Indicador: Formata e exporta arquivos do projeto local para servidor web, de acordo com as tecnologias disponíveis. (00074-04-I-02)
            Indicador: Verifica compatibilidade e performance do website, de acordo com as normas e exigências do mercado. (00074-04-I-03)
            Conhecimento: Tecnologias de servidor web: conceitos e princípios de funcionamento. (00074-04-C-01)
            Conhecimento: Serviços de hospedagem: gratuitos, pagos, compartilhados e dedicados. (00074-04-C-02)
            Conhecimento: Gerenciamento do site: atualização de informações e backups. (00074-04-C-03)
            Conhecimento: Transferência de arquivos: hospedagem via FTP e upload no servidor. (00074-04-C-04)
            Conhecimento: Testes de desempenho: comportamento e integridade do website. (00074-04-C-05)
            Habilidade: Hospedar websites. (00074-04-H-01)
            Habilidade: Identificar e corrigir erros no website. (00074-04-H-02)
            Habilidade: Realizar backups de websites. (00074-04-H-03)
            Orientações metodológicas: Nesta unidade curricular, sugere-se a prática de atividades de pesquisa de servidores de hospedagem que atendam às necessidades do projeto, bem como a prática de transferência de arquivos locais para os servidores de hospedagem via FTP. (00074-04-O-01)
            Tecnologias sugeridas:
            APP, Sistemas de Gerenciamento de Tarefas e Projetos, Essas ferramentas ajudam a planejar, organizar e acompanhar as etapas e cronogramas dos projetos de hospedagem., ampliação, Aumenta a eficiência na organização e execução dos projetos de hospedagem., 00074-04-I-01|00074-04-I-03|00074-04-C-05|00074-04-H-01, Facilita a definição do serviço de hospedagem (I-01) e o teste de desempenho do website (I-03), aplicando conceitos de gerenciamento de cronogramas (C-05) e habilidades de hospedagem (H-01).

            Curso: Doceiro (2438) - 2018

            Título da UC: Organizar o Ambiente de Trabalho para Produções Gastronômicas (00149-01)
            Indicador: Executa atividades operacionais, de acordo com o fluxo do ambiente de trabalho, utilizando mobiliários, selecionando utensílios e equipamentos adequados para a realização dos serviços. (00149-01-I-01)
            Indicador: Organiza e higieniza instalações, equipamentos e utensílios, de acordo com as boas práticas para serviços de alimentação. (00149-01-I-02)
            Indicador: Prepara equipamentos e utensílios, de acordo com a segurança individual/coletiva e a ficha técnica de produção. (00149-01-I-03)
            Conhecimento: Contexto da gastronomia: mercado de alimentação, área de atuação, equipes de trabalho e inserção profissional. (00149-01-C-01)
            Conhecimento: Legislação e procedimentos: boas práticas para serviços de alimentação quanto a: higiene pessoal, ambiental, equipamentos, móveis, utensílios de trabalho e insumos (recebimento, armazenamento, pré-preparo, preparo, resfriamento, envase e distribuição). Riscos de contaminação cruzada, multiplicação de micro-organismos e descarte de resíduos. (00149-01-C-02)
            Conhecimento: Segurança aplicada a ambientes de gastronomia: procedimentos de comunicação de primeiros socorros, identificação de situações de risco e contenção de potenciais danos. Equipamento de proteção individual (EPI) e equipamento de proteção coletiva (EPC). Uso de equipamentos de combate a incêndio. Identificação de rotas de fuga. (00149-01-C-03)
            Conhecimento: Organização e estrutura do ambiente de trabalho: leiaute, mobiliários, equipamentos e utensílios; características, funcionamento, utilização e conservação. Fluxo de operação. (00149-01-C-04)
            Conhecimento: Ficha técnica de produção: conceito e finalidade. (00149-01-C-05)
            Habilidade: Realizar atividades de higienização de instalações, equipamentos e utensílios. (00149-01-H-01)
            Habilidade: Utilizar boas práticas para serviços de alimentação nos processos de organização do ambiente de trabalho. (00149-01-H-02)
            Habilidade: Selecionar e utilizar equipamentos e utensílios. (00149-01-H-03)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00149-01-H-04)
            Habilidade: Interpretar ficha técnica de produção. (00149-01-H-05)
            Orientações metodológicas: Recomenda-se aos docentes planejar situações de aprendizagem que propiciem aos alunos o reconhecimento da atuação profissional frente ao mercado de trabalho e como parte integrante do processo produtivo no segmento de Gastronomia, por meio da realização das etapas de trabalho, considerando as especificidades da função no que diz respeito ao ambiente, equipamentos, utensílios, legislação e segurança além da constante utilização de termos técnicos no desenrolar dos fazeres. 
            Sugere-se ao docente propor atividades onde os alunos tenham contato com profissionais da gastronomia e da área da saúde, sobretudo da nutrição, Vigilância Sanitária, como também empresários do setor de alimentação, por meio de palestras, seminários, visitas técnicas, discussões em grupo e estudos de situações-problema.  Além disso o docente deve realizar atividades práticas que evidenciem a organização do espaço de trabalho levando em consideração a questão da ergonomia, a utilização de Equipamentos de Proteção Individual (EPIs) e a necessidade da criticidade no que diz respeito a sustentabilidade, as relações de trabalho e suas responsabilidades. 
            Para turmas que são compostas por diferentes cursos, mas que se valem da convergência de unidades curriculares é imprescindível que o docente trabalhe as especificidades para cada função dentro dos parâmetros da competência.
            Ressalta-se que o docente deve atentar-se aos demais cursos que convergem com esta unidade curricular junto ao seu Departamento Regional. (00149-01-O-01)
            Tecnologias sugeridas:
            IOT, Sensores Inteligentes, Sensores inteligentes podem monitorar níveis de higiene e garantir conformidade com protocolos de segurança alimentar., geração, Gera atividades automatizadas de monitoramento da higiene do ambiente., 00149-01-I-02|00149-01-C-02|00149-01-H-01, A tecnologia impacta a organização e higienização (I-02) e está diretamente relacionada ao conhecimento sobre boas práticas (C-02) e habilidades de higienização (H-01).
            APP, Sistemas de Gerenciamento de Tarefas e Projetos, Facilitam a organização das atividades operacionais e garantem a eficiência na seleção de utensílios e equipamentos., ampliação, Ampliam a capacidade de organização e controle das atividades., 00149-01-I-01|00149-01-H-03|00149-01-C-04, O uso dessas plataformas facilita a execução correta das atividades operacionais (I-01), inclui a seleção de equipamentos (H-03), e organização e estrutura do ambiente de trabalho (C-04).

            Título da UC: Controlar e Organizar Estoques em Ambientes de Manipulação de Alimentos (00149-02)
            Indicador: Controla a reposição de produtos, de acordo com ficha técnica e estoque atual. (00149-02-I-01)
            Indicador: Recebe e armazena as mercadorias, de acordo com as boas práticas para serviços de alimentação e documentos orientadores. (00149-02-I-02)
            Conhecimento: Mercadorias: especificações técnicas, sazonalidade, classificação dos insumos. (00149-02-C-01)
            Conhecimento: Matemática: quatro operações, porcentagens e conversão de medidas. (00149-02-C-02)
            Conhecimento: Ficha técnica de produção: implicações para o estoque. (00149-02-C-03)
            Conhecimento: Estoques: conceito, importância, integração com outros setores, documentos orientadores, sistemas de gestão. (00149-02-C-04)
            Conhecimento: Operação do estoque: entradas, saídas, estoque mínimo e máximo e ponto de pedido. (00149-02-C-05)
            Conhecimento: Boas práticas para serviços de alimentação no recebimento e armazenamento de mercadorias: controle do prazo de validade (PVPS e PEPS) e armazenamento, fluxo de recebimento e documentos operacionais de estoque; armazenamento de amostras, procedimentos de descartes de embalagens e resíduos; legislação sanitária vigente (orientações para áreas de armazenamento). (00149-02-C-06)
            Conhecimento: Procedimentos de pré-higienização das mercadorias recebidas. (00149-02-C-07)
            Habilidade: Relacionar necessidades de compras com itens disponíveis no estoque. (00149-02-H-01)
            Habilidade: Efetuar cálculos para definição de ponto de pedidos. (00149-02-H-02)
            Habilidade: Interpretar ficha técnica de produção. (00149-02-H-03)
            Habilidade: Conferir mercadorias. (00149-02-H-04)
            Habilidade: Preencher documentos orientadores. (00149-02-H-05)
            Habilidade: Comunicar-se de maneira assertiva. (00149-02-H-06)
            Habilidade: Administrar entrada e saída de mercadorias. (00149-02-H-07)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00149-02-H-08)
            Orientações metodológicas: Recomenda-se aos docentes planejar situações de aprendizagem que propiciem aos alunos a realização das etapas de trabalho, considerando as especificidades da função e no que diz respeito ao processo de movimentação dos insumos, legislações aplicadas ao ambiente de estoque, controle das quantidades, organização de equipamentos e utensílios, entre outras.  
            Sugere-se ao docente propor visitas técnicas, estudos de situações-problema, simulações de recebimento e armazenamento de mercadorias e o preenchimento de documentos orientadores de controle de estoque, e seus processos. Recomenda-se ainda que, para tais atividades, o docente considere a questão da ergonomia, a utilização de Equipamentos de Proteção Individual (EPIs) e a necessidade da criticidade no que diz respeito a sustentabilidade, as relações de trabalho e suas responsabilidades.
            Para turmas que são compostas por diferentes cursos, mas que se valem da convergência de unidades curriculares é imprescindível que o docente trabalhe as especificidades para cada função dentro dos parâmetros da competência. 
            Ressalta-se que o docente deve atentar-se aos demais cursos que convergem com esta unidade curricular junto ao seu Departamento Regional (00149-02-O-01)
            Tecnologias sugeridas:
            IOT, Sensores Inteligentes, Sensores que monitoram níveis de estoque e avisam automaticamente quando reposição é necessária., geração, Geram alertas automáticos para reposição de itens, garantindo a manutenção do estoque., 00149-02-I-01|00149-02-C-05, A tecnologia facilita o controle preciso do estoque (I-01) e operação do estoque (C-05) automatizando o processo.
            APP, Sistemas de Gerenciamento de Estoques, Auxilia no controle preciso das entradas e saídas de estoque, além da reposição de produtos., ampliação, Melhora a gestão e precisão no controle de estoque., 00149-02-I-01|00149-02-I-02|00149-02-C-05|00149-02-H-07, Esses sistemas facilitam controlar a reposição de produtos (I-01), armazenar mercadorias (I-02), operação do estoque (C-05) e administração da entrada e saída de mercadorias (H-07).
            AD, Ferramentas de Análise Predictiva, Ferramentas de análise preditiva podem antecipar necessidades de estoque com base em dados históricos., geração, Gera novas atividades de previsão e planejamento de compras., 00149-02-I-01|00149-02-C-03|00149-02-H-01, A tecnologia afeta diretamente o controle de reposição de produtos (I-01) a partir da ficha técnica (C-03) e empate a habilidade de relacionar necessidades de compras (H-01).

            Título da UC: Produzir e Comercializar Doces (00149-03)
            Indicador: Prepara bases da doçaria, considerando técnicas de açúcar e espessamento e as boas práticas para serviços de alimentação. (00149-03-I-01)
            Indicador: Apresenta produções da doçaria, considerando o tipo de demanda e a produção de confeitaria. (00149-03-I-02)
            Indicador: Precifica, considerando clientes, concorrentes, ficha técnica e custos de produção. (00149-03-I-03)
            Indicador: Seleciona embalagens às produções da confeitaria, considerando integridade, atratividade e custo do produto. (00149-03-I-04)
            Conhecimento: Boas práticas: higiene pessoal, ambiental e de insumos; noções de microbiologia de alimentos no preparo; riscos de contaminação cruzada e descarte de resíduos; princípios da sustentabilidade na gestão de recursos produtos e insumos. (00149-03-C-01)
            Conhecimento: Mise en place: conceito, finalidades e procedimentos. (00149-03-C-02)
            Conhecimento: Ingredientes aplicados à doçaria: tipos, características, funções e sazonalidade. (00149-03-C-03)
            Conhecimento: Técnicas de cocção aplicadas a doçaria. (00149-03-C-04)
            Conhecimento: Técnica de Calda de açúcar: pontos (fio, bala, quebra, caramelo), temperatura, tipos de açúcar, produções (geleias, coulis, compotas, doces cristalizados). (00149-03-C-05)
            Conhecimento: Técnicas com espessamento: pectina (frutas), proteínas (ovos), condensados (leite), gordura (oleaginosas) e amido (cereais) e produções (ambrosia, doce de leite, arroz doce, canjica, paçoca, cocada). (00149-03-C-06)
            Conhecimento: Armazenamento de produções da doçaria: congelamento, resfriamento, esterilização, embalagem e transporte. (00149-03-C-07)
            Conhecimento: Finalização e apresentação da doçaria: utensílios, estrutura, ingredientes. (00149-03-C-08)
            Conhecimento: Embalagem para produções da doçaria: materiais (palha, celofane, papel chumbo, fita, crepom, tules, seda), tipos e montagem. (00149-03-C-09)
            Conhecimento: Mercado consumidor: clientes, concorrentes e fornecedores. (00149-03-C-10)
            Conhecimento: Formação de preço: pesquisa de mercado, ficha técnica, custos direto e indireto de produção e margem de lucro. (00149-03-C-11)
            Habilidade: Interpretar ficha técnica. (00149-03-H-01)
            Habilidade: Selecionar ingredientes para doçaria. (00149-03-H-02)
            Habilidade: Calcular os insumos para a produção de doces. (00149-03-H-03)
            Habilidade: Envazar e embalar doces. (00149-03-H-04)
            Orientações metodológicas: Recomenda-se aos docentes planejar situações de aprendizagem práticas de doçaria que propiciem aos alunos a realização das etapas de trabalho, considerando as especificidades da função no que diz respeito ao processo de preparo de doces.
            Sugere-se ao docente propor experimentações sobre doces com técnica de Calda de açúcar e técnicas com espessamento.
            Algumas produções podem ser utilizadas para a prática de embalagens e processo de precificação. Recomenda-se ainda que, para tais atividades, o docente considere a questão da ergonomia, a utilização de Equipamentos de Proteção Individual (EPIs), boas práticas para serviços de alimentação e a necessidade da criticidade no que diz respeito a sustentabilidade, as relações de trabalho e suas responsabilidades.
            O docente pode propor visitas a produtoras de doces ou proporcionar palestras de doceiras da região trazendo elementos sobre técnicas de produção, e informações sobre a comercialização destes produtos. 
            Para turmas que são compostas por diferentes cursos, mas que se valem da convergência de unidades curriculares é imprescindível que o docente trabalhe as especificidades para cada função dentro dos parâmetros da competência. 
            Ressalta-se que o docente deve atentar-se aos demais cursos que convergem com esta unidade curricular junto ao seu Departamento Regional. (00149-03-O-01)
            Tecnologias sugeridas:
            IAA, IA Generativa, IA generativa pode ser utilizada para criar novas receitas baseadas em ingredientes disponível e tendências do mercado., ampliação, Amplia a capacidade criativa na produção de novos doces., 00149-03-I-01|00149-03-C-03|00149-03-H-02, Afeta a preparação de bases na doçaria (I-01) e envolve conhecimentos (C-03) e habilidades de seleção de ingredientes (H-02).
            APP, Sistemas de E-commerce, Plataformas de e-commerce facilitam a venda direta dos doces ao consumidor final., transferência, Transferem parte do processo de venda da interação pessoal para a online., 00149-03-I-03|00149-03-C-10|00149-03-H-01, Impacta diretamente a precificação e venda direta ao consumidor (I-03), considerando mercado (C-10) e habilidades de interpretação da ficha técnica (H-01).

            Curso: Florista (2942) - 2023

            Título da UC: Organizar o Ambiente de Trabalho do Florista (00016-01)
            Indicador: Seleciona fornecedores e controla o fluxo logístico, conforme o tipo de insumo, prazos e condições de entrega. (00016-01-I-01)
            Indicador: Separa e ordena insumos e objetos para a execução da arte floral, de acordo com o uso. (00016-01-I-02)
            Indicador: Manuseia materiais botânicos e de natureza diversa, conforme requisitos de limpeza e higienização da área de trabalho. (00016-01-I-03)
            Indicador: Maximiza a vida útil das plantas, de acordo com técnicas de conservação e pós-colheita de cada espécie. (00016-01-I-04)
            Indicador: Seleciona contêineres apropriados para compor e manter plantas, de acordo com os requisitos de conservação e manutenção. (00016-01-I-05)
            Conhecimento: Botânica: tipologia de plantas e flores, sazonalidade, características de uso nos projetos de arte floral. (00016-01-C-01)
            Conhecimento: Conhecimento botânico das flores e plantas: nomenclatura, preservação e cuidados. (00016-01-C-02)
            Conhecimento: Limpeza e conservação de flores, plantas e material botânico: técnicas, tipos e formas de utilização de produtos de limpeza do material botânico. (00016-01-C-03)
            Conhecimento: Elementos não botânicos utilizados para a estrutura da arte floral: tipos e usos de arames, cordões de diversas espécies, espumas florais, fitas diversas. (00016-01-C-04)
            Conhecimento: Sustentabilidade: reaproveitamento de materiais, utilização de materiais ecologicamente corretos e descarte. (00016-01-C-05)
            Conhecimento: Estoque: tipos de estoque, formas e critérios de armazenagem do material botânico, ferramentas de controle de estoque, limpeza e organização do estoque; manutenção e limpeza de ferramentas. (00016-01-C-06)
            Conhecimento: Fornecedores: Prazos, logística e condições de entrega. (00016-01-C-07)
            Conhecimento: Ergonomia do ambiente de trabalho de florista: espaço físico, leiaute e área de circulação. (00016-01-C-08)
            Conhecimento: Materiais, equipamentos e ferramentas: tipos e usos. (00016-01-C-09)
            Conhecimento: Normas de segurança do trabalho e EPIs do florista. (00016-01-C-10)
            Conhecimento: Compra e uso de flores, plantas e materiais botânicos que respeitem legislação ambiental, fitossanitária e biopirataria. (00016-01-C-11)
            Habilidade: Reconhecer e selecionar elementos botânicos e não botânicos. (00016-01-H-01)
            Habilidade: Pesquisar fornecedores. (00016-01-H-02)
            Habilidade: Comunicar-se de maneira assertiva. (00016-01-H-03)
            Habilidade: Organizar materiais, instrumentos e local de trabalho. (00016-01-H-04)
            Habilidade: Operar planilhas de cálculo, editores de texto e de slides. (00016-01-H-05)
            Habilidade: Mediar conflitos nas situações de trabalho. (00016-01-H-06)
            Habilidade: Manipular flores e materiais botânicos. (00016-01-H-07)
            Orientações metodológicas: Nesta unidade curricular, recomenda-se que o docente aborde os diferentes espaços de trabalho do florista e as formas de organização da área de trabalho, prezando pela segurança, limpeza e conforto ergonômico. Sugere-se enfoque nas necessidades de conservação dos materiais botânicos e não botânicos. Para essa unidade curricular, o docente poderá planejar visitas técnicas em fornecedores de insumos, produtores de plantas ornamentais e floriculturas locais.   (00016-01-O-01)
            Tecnologias sugeridas:
            IOT, Sensores Inteligentes, Sensores podem monitorar condições do ambiente como temperatura e umidade para maximizar a vida útil das plantas., ampliação, Amplia a capacidade de controle e conservação das plantas., 00016-01-I-04|00016-01-C-02, Sensores ajudarão na conservação das plantas (I-04), articulando conhecimentos sobre preservação e cuidados botânicos (C-02).
            APP, Sistemas de Gerenciamento de Estoques, Sistemas automatizados para controle de estoque facilitam a gestão de insumos botânicos e materiais diversos., ampliação, Melhora a precisão no controle dos estoques e insumos., 00016-01-I-01|00016-01-C-06, Sistemas de gerenciamento de estoques ajudarão na seleção de fornecedores e controle logístico (I-01), utilizando conhecimentos sobre estoque (C-06).

            Título da UC: Elaborar Produções Florais (00016-02)
            Indicador: Elabora proposta de produção floral de acordo com o conceito artístico e as características estabelecidas pelo cliente. (00016-02-I-01)
            Indicador: Identifica e reproduz a produção floral de acordo com técnicas e mecânicas de produção floral. (00016-02-I-02)
            Indicador: Cria arranjos, composições ou decorações considerando o briefing do trabalho. (00016-02-I-03)
            Indicador: Utiliza materiais botânicos, recipientes, acessórios para o produto planejado, de acordo com o briefing de trabalho. (00016-02-I-04)
            Indicador: Estabelece a predominância de materiais botânicos como destaque do design do arranjo floral, de acordo com a especificação da demanda. (00016-02-I-05)
            Indicador: Específica e identifica a produção de arranjos e buquês considerando os estilos florais decorativo, vegetativo e forma e linha. (00016-02-I-06)
            Indicador: Elabora proposta de vitrine para exposição de produtos considerando estratégias de visual merchandising. (00016-02-I-07)
            Conhecimento: Atendimento: perfil de clientes (tipos e características dos diversos públicos); etapas; ferramentas para coleta de dados e informações; ações de relacionamento; retenção. (00016-02-C-01)
            Conhecimento: Briefing: conceito, tipos, modelos e informações necessárias. (00016-02-C-02)
            Conhecimento: Aspectos da composição fotográfica: enquadramento, plano, iluminação de produtos. (00016-02-C-03)
            Conhecimento: Técnicas artesanais: colagem, pintura, dobradura, entre outras. (00016-02-C-04)
            Conhecimento: Matemática básica: regra de três, porcentagem, operações básicas (adição, subtração, divisão e multiplicação). (00016-02-C-05)
            Conhecimento: Orçamento: modelos e formação de preço (custos fixos e variáveis). (00016-02-C-06)
            Conhecimento: Técnicas de exposição de produtos: visual merchandising, volumes, alturas, focos visuais, equilíbrio e olhar do consumidor. (00016-02-C-07)
            Conhecimento: História da arte floral: relação com a história da arte, principais características e marcos históricos. (00016-02-C-08)
            Conhecimento: Tendências e inovações do design floral e de áreas correlatas como artes, arquitetura, design de interiores, moda. (00016-02-C-09)
            Conhecimento: Desenho: técnicas de representação do desenho livre. (00016-02-C-10)
            Conhecimento: Teoria das cores: círculo cromático, harmonia das cores, psicologia das cores aplicados a produção floral. (00016-02-C-11)
            Conhecimento: Etiqueta social das flores: aspectos culturais, psicologia das cores, representações e significados. (00016-02-C-12)
            Conhecimento: Fundamentos da arte floral: Formas florais (espigada, multidirecional, redonda e pendente); Organização floral (simetria e assimetria); Construção de arranjos florais (composição, proporção, equilíbrio, ritmo); Características florais (forma, espaços, textura). (00016-02-C-13)
            Conhecimento: Estilo de arranjos: decorativo, linear, vegetativo. (00016-02-C-14)
            Conhecimento: Técnicas e mecânicas de produção floral - buquês de flores de corte (espiral e paralelo); arranjos com plantas; arranjo com flores de corte (espuma floral, aramação, colagem, estrutura). (00016-02-C-15)
            Conhecimento: Decorações florais para noivas: história, estilos e técnicas (braçada, espiral, paralelo, aramado, no suporte, entre outros). (00016-02-C-16)
            Conhecimento: Decorações florais para objetos e indumentárias - joias, vestimentas entre outros: história, estilos e técnicas. (00016-02-C-17)
            Conhecimento: Assinatura de flores: conceito e características do serviço. (00016-02-C-18)
            Conhecimento: Conhecimento botânico das flores e plantas: nomenclatura, preservação e cuidados. (00016-02-C-19)
            Conhecimento: Especificidades das estações do ano e sua influência na disponibilidade de flores, plantas e materiais botânicos. (00016-02-C-20)
            Conhecimento: Produção floral com conceitos de sustentabilidade ambiental. (00016-02-C-21)
            Conhecimento: Técnicas de embalagem para buquês e arranjos. (00016-02-C-22)
            Conhecimento: Procedimentos de transporte: embalagens e sistemas para entregas. (00016-02-C-23)
            Conhecimento: Conservação e longevidade do arranjo: técnicas de hidratação, tipos de corte e conservantes. (00016-02-C-24)
            Conhecimento: Marketing: estratégias, comportamento do consumidor, identidade visual, marca e marketing sensorial. (00016-02-C-25)
            Habilidade: Comunicar-se de maneira assertiva. (00016-02-H-01)
            Habilidade: Interpretar briefing do cliente. (00016-02-H-02)
            Habilidade: Utilizar técnica adequada para elaboração da produção floral. (00016-02-H-03)
            Habilidade: Utilizar ferramentas e equipamentos para elaboração da produção floral. (00016-02-H-04)
            Habilidade: Manipular flores e materiais botânicos. (00016-02-H-05)
            Habilidade: Elaborar croquis de produções florais. (00016-02-H-06)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00016-02-H-07)
            Habilidade: Interpretar os requisitos ambientais relacionados ao produto. (00016-02-H-08)
            Habilidade: Operar ferramentas e recursos tecnológicos. (00016-02-H-09)
            Habilidade: Pesquisar tendências em diferentes expressões artísticas e socioculturais. (00016-02-H-10)
            Orientações metodológicas: Nesta unidade curricular, recomenda-se que o docente aborde o processo da elaboração e construção de uma produção floral, a fim de atender às necessidades do cliente. Recomenda-se o uso de atividades práticas para promover a vivência do processo criativo, articulada as técnicas expressadas nas produções florais. O docente poderá propor, também, o uso de diferentes materiais e containers que podem auxiliar na apresentação estética das produções. (00016-02-O-01)
            Tecnologias sugeridas:
            RE, Realidade Aumentada (AR), AR pode ser usada para visualizar as produções florais em um ambiente virtual, aprimorando a apresentação aos clientes., ampliação, Amplia a capacidade de visualizar a produção floral no contexto real antes da execução., 00016-02-I-01|00016-02-C-19, AR ajuda a criar visualizações imersivas das produções florais (I-01), utilizando conhecimentos botânicos (C-19).
            IAA, IA Generativa, IA Generativa pode ser usada para sugerir combinações de flores e materiais com base em preferências e tendências., geração, Gera novas atividades de design floral automatizado e sugestões personalizadas., 00016-02-I-03|00016-02-C-09, IA Generativa auxilia na criação de arranjos inovadores (I-03), considerando tendências de design floral (C-09).

            Título da UC: Planejar e Executar Projetos de Decoração Floral para Eventos (00016-03)
            Indicador: Realiza visita técnica no ambiente do evento para levantamento de dados por meio de medição e registro fotográfico. (00016-03-I-01)
            Indicador: Elabora o layout da decoração considerando o briefing e o perfil do cliente, o tipo de evento, a estrutura do ambiente, a ergonomia e a circulação das pessoas. (00016-03-I-02)
            Indicador: Elabora proposta de prestação de serviços e orçamento conforme briefing do cliente. (00016-03-I-03)
            Indicador: Elabora plano de trabalho, de acordo com o tipo do evento e perfil do cliente. (00016-03-I-04)
            Indicador: Distribui os elementos e estruturas no ambiente a ser ornamentado, conforme projeto de decoração floral. (00016-03-I-05)
            Indicador: Específica mobiliários, acessórios e elementos decorativos de acordo com as características do evento. (00016-03-I-06)
            Indicador: Insere a arte floral no ambiente, conforme a necessidade de transporte, conservação e longevidade do arranjo. (00016-03-I-07)
            Indicador: Realiza a desmontagem da ornamentação conforme os princípios de sustentabilidade e de reaproveitamento de materiais. (00016-03-I-08)
            Conhecimento: Teoria das cores: círculo cromático, harmonia das cores, psicologia das cores aplicados a eventos. (00016-03-C-01)
            Conhecimento: Eventos: tipos e características (casamentos, aniversários, inaugurações, batizados, eventos corporativos, entre outros). (00016-03-C-02)
            Conhecimento: Tecidos para decoração de ambientes: tipos e características (medidas padrão, texturas, caimento). (00016-03-C-03)
            Conhecimento: Contrato: tipos e itens do contrato (dados da empresa e do cliente, descrição e preços dos serviços, prazos e condições). (00016-03-C-04)
            Conhecimento: Legislação para eventos: leis locais, leis de proteção a espécies vegetais em extinção e áreas de preservação; Auto de Vistoria do Corpo de Bombeiros (AVCB); lei de acessibilidade e normas de sustentabilidade. (00016-03-C-05)
            Conhecimento: História da arte: relação da história da arte, principais características e marcos históricos, com eventos e decorações. (00016-03-C-06)
            Conhecimento: Pesquisas de tendências e inovações em eventos: fontes de pesquisa, novos materiais e elementos botânicos, entre outros. (00016-03-C-07)
            Conhecimento: Briefing: caraterísticas e interpretação das informações. (00016-03-C-08)
            Conhecimento: Desenho: técnicas de representação gráfica. (00016-03-C-09)
            Conhecimento: EPIs para montagem de eventos: capacete, luva, óculos, máscara, entre outros. (00016-03-C-10)
            Conhecimento: Iluminação cênica: tipos, características e aplicabilidade. (00016-03-C-11)
            Conhecimento: Decorações com arranjos florais: Tipos e aplicação em eventos sociais e corporativos. (00016-03-C-12)
            Conhecimento: Arranjos para festas e eventos: Tipos e aplicação em hall de entrada, buffet, centro de mesas, mesas comunitárias, mesas de doces e bolos, aparador. (00016-03-C-13)
            Conhecimento: Orçamento: modelos e formação de preços. (00016-03-C-14)
            Conhecimento: Composição paisagística para ambientação de eventos. (00016-03-C-15)
            Conhecimento: Planejamento logístico: levantamento de disponibilidade de fornecedores e necessidades do local de evento. (00016-03-C-16)
            Conhecimento: Estoque para eventos: tipos de estoque, formas e critérios de armazenagem, ferramentas de controle de estoque, limpeza e organização do estoque; manutenção e limpeza de ferramentas. (00016-03-C-17)
            Conhecimento: Mídias digitais: tipos, características, métodos voltados para promoção pessoal e prospecção de novos clientes. (00016-03-C-18)
            Habilidade: Utilizar ferramentas e equipamentos para elaboração do projeto. (00016-03-H-01)
            Habilidade: Comunicar-se de forma assertiva com o cliente. (00016-03-H-02)
            Habilidade: Interpretar o briefing do cliente. (00016-03-H-03)
            Habilidade: Criar produções florais. (00016-03-H-04)
            Habilidade: Aplicar elementos, princípios e técnicas de design. (00016-03-H-05)
            Habilidade: Manipular flores e materiais botânicos. (00016-03-H-06)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00016-03-H-07)
            Habilidade: Orientar equipes. (00016-03-H-08)
            Habilidade: Mediar conflitos nas situações de trabalho. (00016-03-H-09)
            Orientações metodológicas: Sugere-se a realização de atividade prática em que o aluno planeje e execute uma ornamentação floral, de modo a reproduzir o comportamento dos elementos botânicos na natureza. É possível trabalhar, ainda, com projetos de replantio, onde os alunos elaboram o arranjo floral e na sequência, reintegram o elemento a natureza, ao replantá-lo. Recomenda-se, também, a realização de atividades onde a turma poderá elaborar arranjos a partir de um briefing construído pelo docente ou a partir da entrevista com um cliente real. Para isso, visitas técnicas a fornecedores de insumos e materiais para eventos e contato com profissionais com quem o artista floral interage são importantes nesta etapa do curso.  (00016-03-O-01)
            Tecnologias sugeridas:
            RE, Realidade Aumentada (AR), AR pode ser utilizada para visualizar a decoração floral no ambiente real antes da implementação final, melhorando o planejamento., geração, Gera novas atividades voltadas para a visualização e planejamento em ambiente virtual., 00016-03-I-01|00016-03-I-02|00016-03-H-01, A elaboração de uma prévia da arte floral em AR exige uma coleta de imagens na visita técnica (I-01) e a elaboração do layout da decoração (I-02). A habilidade de utilizar ferramentas e equipamentos para elaboração do projeto é afetada (H-05).
            IOT, Sensores Inteligentes, Sensores podem monitorar condições ambientais como temperatura e umidade para maximizar a longevidade dos arranjos durante eventos., ampliação, Amplia a qualidade da conservação dos arranjos., 00016-03-I-07, Sensores ajudam a assegurar a qualidade dos arranjos florais durante eventos controlando o ambiente em resposta a condições variáveis (I-07)
            APP, Sistemas de Gerenciamento de Tarefas e Projetos, Sistemas como Project e Monday ajudam na criação cronogramas para projetos de decoração floral., ampliação, Ampliação da eficiência na elaboração e apresentação do projeto através da utilização de melhores práticas de organização e gestão de projetos., 00016-03-I-03|00016-03-C-09|00016-03-H-04, O uso dessas plataformas facilita a elaboração e o gerenciamento de cronogramas (I-03), envolve conhecimentos de cronogramas aplicados a projetos de moda (C-09) e a habilidade de elaboração de cronogramas (H-04).

            Curso: Porteiro e Vigia (2198) - 2016

            Título da UC: Executar Atividades do Serviço de Portaria (00089-01)
            Indicador: Controla e registra o acesso de pessoas e veículos, conforme as regras do condomínio/estabelecimento. (00089-01-I-01)
            Indicador: Recebe, identifica e orienta pessoas na portaria, por meios de canais de comunicação existentes, conforme normas do condomínio. (00089-01-I-02)
            Indicador: Recebe e encaminha correspondências e encomendas aos destinatários, conforme as regras do condomínio/estabelecimento. (00089-01-I-03)
            Indicador: Fornece informações do serviço de portaria à equipe de trabalho por meio de equipamentos de comunicação, utilizando códigos apropriados. (00089-01-I-04)
            Indicador: Realiza etapas da passagem de turno, de acordo com as regras do condomínio/estabelecimento. (00089-01-I-05)
            Indicador: Faz relatórios e anotações nos instrumentos de controle em conformidade com as regras do condomínio/estabelecimento. (00089-01-I-06)
            Conhecimento: Tipos de condomínios e suas características: residenciais, comerciais e mistos. (00089-01-C-01)
            Conhecimento: Estrutura organizacional dos condomínios. (00089-01-C-02)
            Conhecimento: Descrição e função do cargo de porteiro e vigia. (00089-01-C-03)
            Conhecimento: Direitos e deveres do porteiro e vigia: leis trabalhistas, tipos de vínculos, convenções sindicais. (00089-01-C-04)
            Conhecimento: Apresentação e postura profissional do porteiro e vigia. (00089-01-C-05)
            Conhecimento: Comunicação nas atividades de portaria: verbal, escrita e gestual. (00089-01-C-06)
            Conhecimento: Técnicas e tipos de atendimento: personalizados, diferenciados (idosos, crianças, pessoas com deficiências, gestantes), presenciais, eletrônicos, telefônicos e impressos. (00089-01-C-07)
            Conhecimento: Equipamentos em portaria: interfone, telefone, rádio, CFTV, computador, mesas de comunicação. (00089-01-C-08)
            Conhecimento: Códigos de comunicação: Q e Alpha. (00089-01-C-09)
            Conhecimento: Procedimentos de acesso de moradores, visitantes, fornecedores, prestadores de serviço, autoridades e de veículos: identificação, registro e autorização. (00089-01-C-10)
            Conhecimento: Listas de informações úteis ao serviço de portaria: relação dos condôminos, funcionários, fornecedores e prestadores de serviço, conforme regulamento do condomínio; caderno de telefones úteis, lista de veículos autorizados para acesso ao condomínio. (00089-01-C-11)
            Conhecimento: Instrumentos de controle: livro de ocorrências, livro ou relação de visitantes, livro ou fichas de fornecedores e prestadores de serviços, livro de sugestões e reclamações e livro de protocolo. (00089-01-C-12)
            Conhecimento: Editores de texto, internet e correio eletrônico. (00089-01-C-13)
            Conhecimento: Veículos: marcas, tipos e modelos. (00089-01-C-14)
            Conhecimento: Dependências dos condomínios/estabelecimentos: área comum, áreas de lazer, espaço fitness, área de festas, localização de torres, zeladoria, salas comerciais, administração e rota de fuga. (00089-01-C-15)
            Conhecimento: Regras do condomínio: convenção e regimento interno. (00089-01-C-16)
            Conhecimento: Tipos de correspondência, controles e encaminhamentos: carta simples, contas, carta registrada, Sedex, revistas e jornais, encomendas, de caráter judicial, via malote. (00089-01-C-17)
            Conhecimento: Código Penal relacionado à conduta do porteiro e vigia. (00089-01-C-18)
            Conhecimento: Procedimento de passagem de turno. (00089-01-C-19)
            Conhecimento: Diversidade e cidadania no condomínio/estabelecimento. (00089-01-C-20)
            Conhecimento: Segurança do trabalho: sinalização de segurança (NR 26). (00089-01-C-21)
            Habilidade: Comunicar-se de maneira assertiva. (00089-01-H-01)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00089-01-H-02)
            Habilidade: Redigir relatórios e documentos. (00089-01-H-03)
            Habilidade: Operar editores de texto. (00089-01-H-04)
            Habilidade: Interpretar normas e regras do condomínio/estabelecimento. (00089-01-H-05)
            Habilidade: Mediar conflitos nas situações de trabalho. (00089-01-H-06)
            Orientações metodológicas: Os alunos desenvolverão as competências próprias dos serviços de portaria. Sugerem-se situações de aprendizagem relacionadas ao atendimento aos condôminos, visitantes e fornecedores, considerando, tanto a ampliação da percepção sobre a diversidade humana e o respeito às diferenças, quanto a observância das normas e regulamentos. Para isso, recomenda-se a realização de dramatizações, estudos de caso, situações-problema, pesquisas, debates e júris simulados.  É importante, ainda, fomentar visitas técnicas que aprofundem o conhecimento dos alunos sobre o trabalho de portaria, o Regimento Interno e a organização do condomínio/estabelecimento. (00089-01-O-01)
            Tecnologias sugeridas:
            APP, Plataformas de Comunicação e Colaboração Online, Facilitam a comunicação e a transferência de informações entre a equipe de trabalho., ampliação, A ampliação da tecnologia facilita a comunicação em tempo real e a precisão das informações transmitidas., 00089-01-I-04|00089-01-I-05|00089-01-H-01|00089-01-C-06|00089-01-C-13|00089-01-C-19, Usar plataformas de comunicação permite uma troca de informações mais eficiente, afetando as atividades de fornecimento de informações do serviço (I-04) e passagem de turno (I-05). Envolve a habilidade de comunicação assertiva (H-01) e os conhecimentos sobre comunicação nas atividades de portaria (C-06), Editores de texto (C-13) e passagem de turno (C-19).
            IOT, Sensores Inteligentes, Monitoram e registram o acesso de pessoas e veículos de forma automática., substituição, A substituição ocorre porque parte do trabalho manual será feita pelos sensores., 00089-01-I-01|00089-01-C-10, Sensores substitui parte do controle e registro de pessoas e veículos (I-01) e envolve os procedimentos de acesso (C-10).
            AD, Data Visualization, Ferramentas para criar relatórios e visualizações claras e compreensíveis., geração, Geração da atividade de produção de relatórios semi-automatizados baseados em dados., 00089-01-I-06|00089-01-H-03|00089-01-C-13, Ferramentas de visualização transformam dados brutos em informações valiosas, criando novas possibilidades de relatórios (I-06) e ampliando o conhecimento sobre ferramentas (C-13) e a habilidade de redigir relatórios (H-03).

            Título da UC: Realizar Medidas Preventivas de Segurança Pessoal e Patrimonial (00089-02)
            Indicador: Monitora áreas comuns e externas cumprindo normas e limites de atuação. (00089-02-I-01)
            Indicador: Faz o controle da portaria de acordo com normas do condomínio/estabelecimento. (00089-02-I-02)
            Indicador: Solicita atendimento de serviços públicos e privados em situações de risco. (00089-02-I-03)
            Indicador: Providencia solução de ocorrências comuns e incidentes em condomínios/estabelecimentos, respeitando seus limites de atuação. (00089-02-I-04)
            Indicador: Fornece informações de segurança à equipe de trabalho por meio de equipamentos de comunicação, utilizando códigos apropriados. (00089-02-I-05)
            Conhecimento: Instalações da portaria e do acesso a garagem: guarita, portões, cancelas e catracas. (00089-02-C-01)
            Conhecimento: Situações de risco e atitudes suspeitas no acesso ou entorno do condomínio/estabelecimento. (00089-02-C-02)
            Conhecimento: Atuação do porteiro e vigia nas atividades preventivas de segurança patrimonial e física. (00089-02-C-03)
            Conhecimento: Equipamentos de segurança do condomínio: claviculário, vigia eletrônico,. (00089-02-C-04)
            Conhecimento: CFTV, botão de pânico, infravermelho ativo, concertina, cerca elétrica, sirene, rádio, interfone, videofone, controle remoto para abrir portas e portões. (00089-02-C-05)
            Conhecimento: Instrumentos de controle nas atividades preventivas de segurança: livro de ocorrências. (00089-02-C-06)
            Conhecimento: Códigos de comunicação nas atividades preventivas de segurança: Q e código Alpha. (00089-02-C-07)
            Conhecimento: Infraestrutura de condomínios: bomba d’água, reservatório de água, gerador, caixa de inspeção, caixa de distribuição geral de energia, relógios de energia, barrilete, central de gás, elevadores, para-raios, luz-piloto, hidrantes, extintores. (00089-02-C-08)
            Conhecimento: Providências nas ocorrências comuns, tais como, com sistemas hidráulicos, elétricos, de esgoto, elevadores e limites de atuação. (00089-02-C-09)
            Conhecimento: Prevenção de incêndios: elementos do fogo, classes do fogo, tipos e utilização de extintores. (00089-02-C-10)
            Conhecimento: Primeiros Socorros em situações de emergência. (00089-02-C-11)
            Habilidade: Organizar o local de trabalho. (00089-02-H-01)
            Habilidade: Mediar conflitos nas situações de trabalho. (00089-02-H-02)
            Habilidade: Comunicar-se de maneira assertiva. (00089-02-H-03)
            Habilidade: Utilizar termos técnicos nas rotinas de trabalho. (00089-02-H-04)
            Habilidade: Fazer registros por meio escrito ou eletrônico. (00089-02-H-05)
            Habilidade: Identificar atitudes suspeitas. (00089-02-H-06)
            Habilidade: Fazer a descrição física de pessoas. (00089-02-H-07)
            Habilidade: Operar extintores de incêndio. (00089-02-H-08)
            Orientações metodológicas: O docente poderá propor atividades do cotidiano operacional do vigia, que permitam ao aluno exercer, na prática, os mais diversos procedimentos de cuidado com o patrimônio e segurança preventiva. Dessa forma será possível o docente verificar as Marcas Formativas como o domínio técnico e científico e a visão crítica que ficam evidentes na fundamentação de uma análise de estudo de caso, na resolução de uma situação-problema, nas simulações práticas e dramatizações.
            É oportuno viabilizar atividades de operação com equipamentos e instrumentos de comunicação e controle de entrada e saída de pessoas e veículos e outras ferramentas existentes para este tipo de serviço. (00089-02-O-01)
            Tecnologias sugeridas:
            IOT, Sensores Inteligentes, Sensores integrados de segurança para monitorar áreas comuns e detectar situações de risco., substituição, ampliação, Substituem parcialmente a vigilância manual e ampliam a capacidade de monitoramento contínuo e imediato., 00089-02-I-01|00089-02-I-02|00089-02-C-02|00089-02-C-05, Sensores ajudam a monitorar áreas comuns (I-01), controlam a portaria (I-02) e detectam atitudes suspeitas (C-02), utilizando conhecimentos sobre CFTV e outros sistemas (C-05).
            AD, Análise de Dados para Segurança Preventiva, Ferramentas de análise de dados para identificar padrões de segurança e áreas de risco., geração, Geram novas atividades de análise com base nos dados coletados para melhorar a segurança., 00089-02-I-04|00089-02-I-05|00089-02-C-06, Facilita a análise e resposta a ocorrências (I-04) e a comunicação de informações de segurança (I-05), utilizando conhecimentos sobre instrumentos de controle (C-06).
            APP, Formulários Online, Automatizam o registro de ocorrências e a comunicação com serviços de emergência., geração, Gera novas atividades de coleta de dados e comunicação automatizada., 00089-02-I-02|00089-02-I-03|00089-02-C-03|00089-02-H-05, Facilita o controle da portaria (I-02), a solicitação de serviços (I-03) e o registro de dados (H-05) utilizando conhecimentos sobre atividades preventivas de segurança (C-03).

            Curso: Técnico em Secretaria Escolar (2905) - 2021

            Título da UC: Prestar Atendimento na Secretaria Escolar (00143-01)
            Indicador: Orienta a comunidade escolar sobre seus direitos e deveres, conforme a legislação vigente e as normas da instituição. (00143-01-I-01)
            Indicador: Disponibiliza informações escolares, de acordo com as funções de cada setor. (00143-01-I-02)
            Indicador: Medeia conflitos nos atendimentos internos e externos da instituição, conforme os procedimentos estabelecidos e a legislação vigente. (00143-01-I-03)
            Conhecimento: Planejamento de carreira: mundo do trabalho, formas de inserção no mercado de trabalho, marketing e apresentação pessoal, preparação de currículos, entrevista de emprego. (00143-01-C-01)
            Conhecimento: Estrutura organizacional e funcional de instituições de ensino: setores, funções e relações. (00143-01-C-02)
            Conhecimento: Organograma e fluxograma da escola. (00143-01-C-03)
            Conhecimento: Comunidade escolar: conceito e segmentos. (00143-01-C-04)
            Conhecimento: Procedimentos regimentais e normas internas da instituição. (00143-01-C-05)
            Conhecimento: Grupos e equipes: conceito, diferenças e tipos. (00143-01-C-06)
            Conhecimento: Técnicas de mediação de conflitos. (00143-01-C-07)
            Conhecimento: Atendimento a clientes: qualidade, características e formas de atendimento: personalizados (idosos, gestantes, pessoas com deficiência) presenciais, eletrônicos, telefônicos e impressos; Técnicas de atendimento a clientes interno e externos. (00143-01-C-08)
            Conhecimento: Elementos da comunicação: contexto, emissor, receptor, canal, mensagem, ruídos e feedback. (00143-01-C-09)
            Conhecimento: Legislação educacional: âmbitos federal, estadual e municipal. (00143-01-C-10)
            Conhecimento: Ergonomia cognitiva e organizacional: conceito e impactos para a atuação Secretaria Escolar. (00143-01-C-11)
            Conhecimento: Qualidade de vida no trabalho do profissional de Secretaria Escolar: estresse e ansiedade, ações para manutenção da saúde física e mental. (00143-01-C-12)
            Conhecimento: Relatórios escolares e gerenciais: conceito, tipos e finalidades. (00143-01-C-13)
            Habilidade: Utilizar técnicas de mediação de conflitos. (00143-01-H-01)
            Habilidade: Comunicar-se de maneira assertiva. (00143-01-H-02)
            Habilidade: Utilizar técnicas de atendimento. (00143-01-H-03)
            Habilidade: Trabalhar em equipe. (00143-01-H-04)
            Orientações metodológicas: 	É nesta Unidade Curricular que se propõe a promoção da relação dos alunos com o mundo do trabalho e o incentivo à educação continuada para seu crescimento pessoal e profissional. Os alunos poderão realizar o planejamento de sua carreira, tendo em vista a análise dos cenários de trabalho e emprego em sua região. 
                Para melhor entendimento das atividades que envolvem esse profissional, sugere-se que o docente promova atividades de aprendizagem relacionadas ao atendimento, abordando a diversidade, tanto no que se refere ao público (pessoas com deficiência, clientes no contexto da diversidade cultural, religiosa, de gênero, faixa etária, dentre outros) como no que diz respeito às diversas possibilidades de instituições escolares nas quais um assistente de Secretaria poderá atuar. Para isso, indica-se a realização de simulações de atendimento, levando em consideração a resolução de conflitos e atendimentos específicos, de acordo com suas características, a postura profissional, a comunicação e o consumo consciente.
                É importante, ainda, propiciar condições para pesquisas, estudos de caso e visitas técnicas, de modo a favorecer o conhecimento dos alunos sobre a legislação educacional brasileira e a organização e funcionalidade das unidades escolares. (00143-01-O-01)
            Tecnologias sugeridas:
            APP, Plataformas de Comunicação e Colaboração Online, Facilitam a comunicação e a transferência de informações entre a equipe de trabalho e a comunidade escolar., ampliação, A ampliação da tecnologia facilita a comunicação em tempo real e a precisão das informações transmitidas., 00143-01-I-01|00143-01-I-02|00143-01-I-03|00143-01-C-09|00143-01-H-01|00143-01-H-02, Usar plataformas de comunicação permite uma troca de informações mais eficiente, afetando as atividades de orientação (I-01), disponibilização de informações (I-02), mediar conflitos (I-03). Envolve a habilidade de comunicar-se assertivamente (H-02) e os conhecimentos sobre comunicação e atendimento a clientes (C-09).

            Título da UC: Realizar Atividades de Apoio Aos Processos Administrativo-Pedagógicos de Secretaria Escolar (00143-02)
            Indicador: Protocola documentos físicos e eletrônicos, de acordo com procedimentos da instituição e a legislação vigente. (00143-02-I-01)
            Indicador: Elabora documentos escolares, de acordo com os objetivos e as normas ortográficas e gramaticais da Língua Portuguesa. (00143-02-I-02)
            Indicador: Arquiva documentos físicos e eletrônicos, de acordo com o disposto na legislação e os tipos e métodos de organização e arquivamento. (00143-02-I-03)
            Indicador: Realiza procedimentos e registros escolares, de acordo com suas características, atendendo à legislação e aos processos organizacionais da instituição. (00143-02-I-04)
            Indicador: Elabora relatórios de acordo com informações obtidas a partir das ações administrativas e pedagógicas. (00143-02-I-05)
            Conhecimento: Utilização de sites oficiais: formas de consulta, acesso e inserção de informações: Órgãos federais, estaduais e municipais de Educação, entre outros. (00143-02-C-01)
            Conhecimento: Protocolo de documentos: recebimento, registro, distribuição, tramitação e expedição. (00143-02-C-02)
            Conhecimento: Tipos de documentos oficiais: edital, ata, declaração, certificado, diploma, histórico, carta, circular, convocação, memorando, ofício, procuração, requerimento e e-mail. (00143-02-C-03)
            Conhecimento: Novo acordo ortográfico. (00143-02-C-04)
            Conhecimento: Regras Gramaticais: pontuação, tempos verbais, pronomes de tratamento e vícios de linguagem. (00143-02-C-05)
            Conhecimento: Métodos e técnicas de arquivo e protocolo: classificação de documentos, recebimento, distribuição, tramitação, expedição, temporalidade e tipos de arquivo (físico e eletrônico). (00143-02-C-06)
            Conhecimento: Métodos de organização e arquivamento: alfabético, geográfico, numérico simples, ideográfico. (00143-02-C-07)
            Conhecimento: Procedimentos e registros escolares: transferência, desistência, evasão, abono de faltas, emissão de diário de classe e de boletim do aluno, matrícula, ficha escolar, histórico escolar e emissão de certificados e diplomas. (00143-02-C-08)
            Conhecimento: Planilhas eletrônicas: conceitos de fórmulas e funções lógicas-matemáticas características de formatação. (00143-02-C-09)
            Conhecimento: Gráficos: Ordenação, classificação, definição, aplicação e criação. (00143-02-C-10)
            Conhecimento: Editores de texto: características das ferramentas de edição e formatação de texto; características de inserção de elementos para edição de texto. (00143-02-C-11)
            Conhecimento: Tabelas: características de criação e formas de manipulação. (00143-02-C-12)
            Conhecimento: Correio eletrônico: características e formas de uso no ambiente escolar. (00143-02-C-13)
            Conhecimento: Navegadores (browser) e ferramentas de pesquisa (páginas dos órgãos oficiais de educação). (00143-02-C-14)
            Conhecimento: Princípios de segurança na internet: navegação segura, spam, phishing e Lei Geral de Proteção de Dados. (00143-02-C-15)
            Conhecimento: Sistemas de registro escolar: tipos e funcionalidade. (00143-02-C-16)
            Conhecimento: Relatórios escolares e gerenciais: conceito, tipos e finalidades. (00143-02-C-17)
            Conhecimento: Princípios de sustentabilidade ou do desenvolvimento sustentável: pilares, princípios ambientais, 3Rs, consumo consciente, 5Ss. (00143-02-C-18)
            Habilidade: Comunicar-se com clareza e de forma adequada à comunidade escolar. (00143-02-H-01)
            Habilidade: Organizar a rotina e os documentos de trabalho. (00143-02-H-02)
            Habilidade: Utilizar planilhas eletrônicas, editores de texto e navegadores de internet. (00143-02-H-03)
            Habilidade: Resolver conflitos inerentes aos processos de trabalho. (00143-02-H-04)
            Habilidade: Pesquisar, coletar e organizar dados e informações. (00143-02-H-05)
            Habilidade: Redigir documentos físicos e eletrônicos inerentes aos procedimentos escolares. (00143-02-H-06)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na organização. (00143-02-H-07)
            Orientações metodológicas: Para as atividades vinculadas a esta Unidade Curricular, o docente poderá propor atividades do cotidiano operacional do profissional de secretaria escolar. É essencial que as atividades de aprendizagem permitam ao aluno exercer, na prática, os mais diversos procedimentos técnicos da sua função, tais como realizar procedimentos e registros escolares: matrícula presencial e matrícula on-line, protocolar, arquivar, identificar tipos de documentos, emitir transferências, histórico escolar, certificados e diplomas físico ou digital, dentre outros. Para o desenvolvimento dessa competência, recomenda-se a utilização de simulações práticas, situações-problema e dramatizações.
            É oportuno o docente viabilizar atividades em laboratórios de informática para a produção e edição de textos e de planilhas eletrônicas, utilização de internet e outras ferramentas tecnológicas, para trabalhar sites governamentais e de outras entidades vinculadas à legislação educacional, efetuando as correções gramaticais necessárias, incentivando a leitura e a escrita e a utilização de diferentes tipos de recursos para a elaboração de gráficos, cálculos e planilhas, pesquisas virtuais e demais procedimentos tecnológicos existentes nos processos da Secretaria Escolar. Orienta-se, ainda, que a navegação na internet e a utilização de e-mails no ambiente de trabalho devem obedecer às normas e políticas da empresa. (00143-02-O-01)
            Tecnologias sugeridas:
            APP, Sistemas de Gestão Eletrônica de Documentos (GED), Apoiam na organização, protocolo, arquivamento e acesso a documentos físicos e eletrônicos., ampliação, Amplia a eficiência e precisão na gestão de documentos., 00143-02-I-01|00143-02-I-03|00143-02-I-04|00143-02-C-02|00143-02-C-06|00143-02-H-02, A utilização do GED melhora protocolos (I-01), arquivamento (I-03), registros escolares (I-04) e envolve o conhecimento de procedimentos e métodos de arquivo (C-02/C-06) e habilidades de organização de documentos (H-02).

            Título da UC: Coletar, Interpretar e Monitorar Dados Estatísticos da Instituição Educacional (00143-03)
            Indicador: Organiza as informações educacionais da Instituição, conforme procedimentos operacionais padronizados. (00143-03-I-01)
            Indicador: Elabora planilhas e gráficos conforme normas da estatística aplicada. (00143-03-I-02)
            Indicador: Registra dados educacionais de acordo com o tipo de informação recebida e o sistema informatizado de gestão. (00143-03-I-03)
            Indicador: Organiza dados estatísticos relativos às avaliações institucionais e externas de acordo com os resultados obtidos. (00143-03-I-04)
            Conhecimento: Estatística aplicada à educação: população e amostra, distribuição de frequência; medidas de tendência central (média, média aritmética ponderada, mediana e moda; tabelas e séries estatísticas. (00143-03-C-01)
            Conhecimento: Estatística e suas relações com planejamento, avaliação, gestão e financiamento da educação. (00143-03-C-02)
            Conhecimento: Sistema operacional cliente, editor de textos e planilha eletrônica, apresentação eletrônica, internet e e-mail. (00143-03-C-03)
            Conhecimento: Sistemas de avaliação educacional dos níveis e modalidade de ensino: tipos, funções, aplicação e características (objetivo, periodicidade, idade, série). (00143-03-C-04)
            Habilidade: Pesquisar e Interpretar dados estatísticos. (00143-03-H-01)
            Habilidade: Interpretar tabelas e gráficos. (00143-03-H-02)
            Habilidade: Calcular estimativas e percentuais. (00143-03-H-03)
            Habilidade: Analisar dados estatísticos. (00143-03-H-04)
            Habilidade: Utilizar recursos de tecnologia da informação e comunicação. (00143-03-H-05)
            Orientações metodológicas: Nesta UC, o docente deverá apresentar os conceitos de estatística aplicada aos processos da Secretaria Escolar. Para tanto, deverá propor atividades práticas de elaboração de tabelas e gráficos, análise de dados estatísticos, situações-problema, interpretação de dados. Realizar simulações a partir de dados levantados em sites oficiais para promover atividade prática de análise dos resultados das avaliações educacionais internas e externas, análise de dados relativos à idade, à série e ao desempenho escolar dos alunos. O docente deverá propor situações de aprendizagem contextualizadas com a rotina da Secretaria que envolvam cálculos estatísticos, de estimativas e percentuais. (00143-03-O-01)
            Tecnologias sugeridas:
            AD, Business Intelligence (BI), Ferramentas de BI melhoram a coleta, interpretação e análise de dados estatísticos., geração, Gera atividades automatizadas de análise e interpretação de dados., 00143-03-I-03|00143-03-C-01|00143-03-H-01|00143-03-H-04, Facilita a organização e registro de dados (I-03) e melhora a aplicação de conhecimentos de estatística aplicada (C-01) e habilidades de análise (H-04).

            Título da UC: Organizar os Processos Legais da Escola Perante e os Órgãos Reguladores (00143-04)
            Indicador: Acompanha as etapas do processo de regularização da instituição educacional, de acordo com a legislação vigente. (00143-04-I-01)
            Indicador: Organiza e atualiza a documentação dos atos oficiais a serem publicados em editais e outros informes da instituição educacional, de acordo com o órgão regulador. (00143-04-I-02)
            Indicador: Monitora os prazos de validade das documentações de regularização da instituição educacional a partir de ferramentas físicas e eletrônicas, de acordo com a legislação vigente. (00143-04-I-03)
            Conhecimento: Legislação educacional: Lei de Diretrizes e Bases da Educação Nacional (LDB), legislações federais, estaduais e municipais da educação (decretos, pareceres, resoluções) no âmbito da regularização da escola. (00143-04-C-01)
            Conhecimento: Conselhos de educação no âmbito municipal, estadual e federal e órgãos colegiados superiores: conceito, atribuição e a relação com a instituição escolar. (00143-04-C-02)
            Conhecimento: Significado da gestão democrática: marcos legislativos e a competência do gestor escolar, visões conceituais da gestão educacional. (00143-04-C-03)
            Conhecimento: Princípios da autonomia escolar: Noções das dimensões (administrativa, financeira, jurídica e pedagógica). (00143-04-C-04)
            Conhecimento: Identificação e atribuições dos órgãos auxiliares da gestão escolar: Conselho Escolar, associação de pais e mestres, grêmio estudantil. (00143-04-C-05)
            Conhecimento: Processos legais da escola: credenciamento, recredenciamento da instituição, autorização, reconhecimento e aprovação de cursos e suas renovações. (00143-04-C-06)
            Conhecimento: Documentos do processo de regularização: alvará de funcionamento, licenciamentos, convênios, documentação da equipe técnico-pedagógica, plano de ensino, projeto político-pedagógico, entre outros. (00143-04-C-07)
            Conhecimento: Escrituração escolar: ofícios, atas correspondentes ao processo de regularização da instituição educacional. (00143-04-C-08)
            Conhecimento: Meios de publicação: acesso, forma de publicação e uso da linguagem em diário oficial, sites oficiais e outros. (00143-04-C-09)
            Conhecimento: Informática: ferramentas de controle de prazos de validade de documentação; planilhas eletrônicas; softwares e aplicativos de organização e controle. (00143-04-C-10)
            Habilidade: Interpretar textos legais e normas técnicas. (00143-04-H-01)
            Habilidade: Redigir documentos escolares e técnicos. (00143-04-H-02)
            Habilidade: Organizar arquivos e documentos. (00143-04-H-03)
            Habilidade: Comunicar-se de maneira assertiva. (00143-04-H-04)
            Habilidade: Pesquisar e coletar dados e informações. (00143-04-H-05)
            Orientações metodológicas: Para esta Unidade Curricular, o docente poderá propor uma situação de aprendizagem que envolva todo o processo de regularização de uma escola. Para tanto, sugerem-se atividades de pesquisa, análise de documentos e legislações exigidas no processo de regularização, seja autorização, credenciamento ou renovação de credenciamento.
            O docente deve viabilizar atividades de pesquisa em laboratórios de informática, tendo em vista a prática de elaboração/redação/formatação dos diferentes documentos pertinentes a esse processo. Recomenda-se a promoção de estratégias que permitam compreensão e execução, por meio de visitas técnicas às diferentes instâncias da instituição educacional e a órgãos externos que se relacionam à dinâmica desses processos. O trabalho em equipe deve ser evidenciado, de modo a favorecer o desenvolvimento da atitude colaborativa e empreendedora.
            Recomenda-se que o docente utilize a legislação própria dos Estados, Municípios e Distrito Federal, para processos legais de regulação da escola. (00143-04-O-01)
            Tecnologias sugeridas:
            APP, Sistemas de Gestão de Prazo e Fluxo, Automatizam o controle de prazos de documentos de regularização da instituição., ampliação, Ampliam a precisão e eficiência no monitoramento de prazos., 00143-04-I-03|00143-04-C-07|00143-04-H-01|00143-04-H-03, Facilitam o monitoramento de validade dos documentos (I-03), integrou ao conhecimento sobre documentação de regularização (C-07) e habilidades de organização (H-03).

            Título da UC: Elaborar, Organizar e Controlar Documentos da Instituição Educacional e da Vida Escolar do Aluno (00143-05)
            Indicador: Redige atas de resultados de avaliação, reuniões e de Conselho de Classe, de acordo com o regimento da instituição educacional. (00143-05-I-01)
            Indicador: Organiza e controla o histórico escolar do aluno, com base nos documentos gerados e suas tipologias. (00143-05-I-02)
            Indicador: Preenche livros de termos de visita, ocorrências e fichas individuais dos alunos, de acordo com o regimento e a rotina escolar. (00143-05-I-03)
            Indicador: Organiza, controla e guarda os livros de registro de classe, livro-ponto e documentos pertinentes às rotinas da instituição escolar. (00143-05-I-04)
            Indicador: Elabora, arquiva e atualiza os documentos oficiais da vida legal da instituição educacional e da vida escolar do aluno, de acordo com os tipos e métodos de arquivamento e as legislações vigentes. (00143-05-I-05)
            Indicador: Expede documentos da vida escolar do aluno, de acordo com a solicitação e as técnicas de protocolo. (00143-05-I-06)
            Conhecimento: Arquivo: história, arquivo, documento e documentação. (00143-05-C-01)
            Conhecimento: Órgãos normalizadores e normas regulamentadoras para gestão documental. (00143-05-C-02)
            Conhecimento: Ambiente e rotinas secretariais: critérios de organização do ambiente e periodicidade das atividades. (00143-05-C-03)
            Conhecimento: Documentos gerados e tipologia documental das seções e departamentos da instituição. (00143-05-C-04)
            Conhecimento: Classificação de documentos: conceito e aplicabilidade. (00143-05-C-05)
            Conhecimento: Métodos de arquivamento e ordenamento de documentos: alfabético, numérico, geográfico, alfanumérico, ideográfico e variadex. (00143-05-C-06)
            Conhecimento: Tabela de temporalidade: conceito e aplicabilidade. (00143-05-C-07)
            Conhecimento: Tipos e usabilidade dos acessórios para armazenamento de documentos: arquivo de aço, arquivo de madeira, pasta suspensa, visores e etiquetas. (00143-05-C-08)
            Conhecimento: Sistemas de gestão eletrônica (GED): conceito, benefícios e aplicabilidade. (00143-05-C-09)
            Conhecimento: Preservação e integralidade do documento: orientações e recomendações. (00143-05-C-10)
            Conhecimento: Saúde e segurança no manuseio de documentos: uso de equipamentos de proteção individual (EPIs), ergonomia (NR-17) e prevenção contra incêndio (NR-23). (00143-05-C-11)
            Conhecimento: Critérios e processos para empréstimo e devolução de documentos: prazos, orientações e recomendações. (00143-05-C-12)
            Conhecimento: Protocolo de documentos: registro de entrada e saída de documentos, controle da movimentação de documentos. Critérios e processos para empréstimo e devolução de documentos: prazos, orientações e recomendações. (00143-05-C-13)
            Conhecimento: Comunicação escrita: regras gramaticais, ortografia, acentuação, concordância verbal e nominal, pronome de tratamento, estrutura textual, coesão e coerência. (00143-05-C-14)
            Conhecimento: Normas da Associação Brasileira de Normas Técnicas (ABNT). (00143-05-C-15)
            Conhecimento: Informática - editor de texto: edição e formatação de textos e documentos. (00143-05-C-16)
            Habilidade: Interpretar normas e legislações educacionais vigentes. (00143-05-H-01)
            Habilidade: Interpretar e manusear dados de sistemas de gestão documental. (00143-05-H-02)
            Habilidade: Controlar arquivamento, empréstimo e devolução dos documentos. (00143-05-H-03)
            Habilidade: Orientar-se pelo fluxograma para a circulação e descarte de documentos. (00143-05-H-04)
            Orientações metodológicas: Nesta Unidade Curricular, recomenda-se que o docente apresente aos alunos modelos de documentos, como: ficha de matrícula, histórico escolar do aluno, fichas de prontuário de docentes, atas, ocorrências, fichas individuais e documentos pertinentes às rotinas da instituição educacional. Proporcionar atividades de simulações e práticas de elaboração, de preenchimento e técnicas de arquivamento para controle e organização dos documentos. Sugere-se resolução de situações-problema que envolvam a gestão de documentos de uma instituição educacional fictícia.
            Deverá propor, ainda, nas situações de aprendizagem que envolvam a melhoria dos diferentes processos acerca da vida escolar do aluno, como a matrícula presencial e on-line, o aproveitamento de estudos, a transferência, o cancelamento e a certificação de competência, com base no regimento escolar de uma instituição fictícia ou real.  Portanto, os desafios devem contemplar ações que permitam aprimorar processos de trabalho e, na prática, vivenciá-los, por meio de visitas técnicas e simulações que compreendam o ciclo de elaboração e arquivamento.
            Para tanto, sugere-se que sejam estabelecidas parcerias com empresas e com as Secretarias de Educação do município e do estado, ou equivalentes, no intuito de viabilizar o uso dos sistemas eletrônicos às práticas dos alunos.
            Recomenda-se que o docente utilize o manual da Secretaria Escolar próprio dos Estados, Municípios e Distrito Federal, para elaborar, organizar e controlar documentos da instituição e da vida escolar do aluno. (00143-05-O-01)
            Tecnologias sugeridas:
            APP, Sistemas de Gestão Eletrônica de Documentos (GED), Ajuda na organização, controle e arquivamento de documentos relacionados à vida escolar do aluno., ampliação, Amplia a precisão e eficiência na gestão de documentos estudantis., 00143-05-I-02|00143-05-I-05|00143-05-C-02|00143-05-H-02|00143-05-H-03, Melhora a organização do histórico escolar (I-02) e atualização dos documentos oficiais da vida escolar do aluno (I-05), utilizando o conhecimento sobre normas de gestão documental (C-02) e habilidades de manusear dados (H-02) e controlar arquivamentos (H-03).
            IAA, Sistemas de Transcrição de Reuniões, Sistemas poderão redigir atas de reuniões de forma automatizada, facilitando a elaboração de documentos., substituição, Substitui a necessidade de digitação manual durante reuniões., 00143-05-I-01|00143-05-C-14, A tecnologia automatiza a elaboração de atas de avaliação e reuniões (I-01), articulando conhecimentos de comunicação escrita (C-14).

            Título da UC: Apoiar e Executar As Ações de Avaliação e Controle de Processos, Atividades e Recursos Materiais da Secretaria Escolar (00143-06)
            Indicador: Recebe e dá encaminhamento às solicitações recebidas dos diversos setores da instituição educacional, considerando a estrutura, os procedimentos internos e interfaces. (00143-06-I-01)
            Indicador: Preenche formulários com informações e dados pertinentes aos processos de trabalho relativos à política de gestão da instituição educacional. (00143-06-I-02)
            Indicador: Elabora instrumentos de controle da gestão, conforme demandas da direção escolar. (00143-06-I-03)
            Conhecimento: Ambiente e rotinas secretariais: critérios de organização do ambiente e periodicidade das atividades. (00143-06-C-01)
            Conhecimento: Cultura organizacional: conceitos e elementos. (00143-06-C-02)
            Conhecimento: Funções da Administração: planejamento, organização, direção e controle. (00143-06-C-03)
            Conhecimento: Atendimento ao cliente: conceito e tipos de atendimento, tipos de clientes, técnicas de atendimento e abordagem ao cliente. Gestão de conflitos. Relacionamento interpessoal; comunicação oral, formas de contato com clientes (via telefone, e-mail e outros), técnicas de negociação. (00143-06-C-04)
            Conhecimento: Processos organizacionais: conceito, inter-relação e aplicabilidade. (00143-06-C-05)
            Conhecimento: Sustentabilidade e responsabilidade socioambiental: conceitos e objetivos. (00143-06-C-06)
            Conhecimento: Qualidade de vida no trabalho: conceitos, métodos, perfil da empresa, análise das necessidades, pesquisas e aplicações. (00143-06-C-07)
            Habilidade: Comunicar-se de maneira assertiva. (00143-06-H-01)
            Habilidade: Administrar o tempo e as atividades de trabalho. (00143-06-H-02)
            Habilidade: Mediar conflitos nas situações de trabalho administrativo. (00143-06-H-03)
            Habilidade: Identificar os aspectos do próprio trabalho que interferem na instituição escolar. (00143-06-H-04)
            Habilidade: Operar recursos da tecnologia da informação e comunicação. (00143-06-H-05)
            Orientações metodológicas: Nesta Unidade Curricular, o docente poderá propor atividades de análise de estudos de caso que representem desafios e situações-problema acerca do cumprimento de normas e da rotina de uma Secretaria Escolar. Propor atividades que promovam a melhoria dos diferentes processos de trabalho, por meio da pesquisa institucional, coleta e geração de dados. Os desafios propostos devem contemplar ações que permitam aprimorar a avaliação dos diferentes processos de trabalho.  Sugere-se que o docente apresente resultados de avaliações educacionais para os alunos analisarem e elaborarem relatórios a serem apresentados aos gestores de instituições escolares, fictícias ou reais.
            Recomenda-se utilizar o laboratório de informática para propor atividades de coleta e organização da informação, por meio da utilização de diferentes recursos tecnológicos. São sugeridas, também atividades em equipe para os alunos simularem uma rotina de Secretaria Escolar, atendendo clientes internos e externos. Orienta-se, ainda, que a navegação na internet e a utilização de e-mails no ambiente de trabalho devem obedecer às normas e políticas da instituição educacional. 
            Recomenda-se que o docente utilize o manual da Secretaria Escolar próprio dos Estados, Municípios e Distrito Federal, para apoio e execução nas ações de avaliação e controle de processos da Secretaria Escolar. (00143-06-O-01)
            Tecnologias sugeridas:
            APP, Sistemas de Gerenciamento de Tarefas e Projetos, Facilitam a administração das atividades de trabalho e a comunicação interna na Secretaria Escolar., ampliação, O uso de tais sistemas ampliaria a eficiência do gerenciamento., 00143-06-I-01|00143-06-I-03|00143-06-C-03|00143-06-H-02|00143-06-H-05, Melhora a administração e execução das atividades (I-01/I-03), utilizando conhecimentos em planejamento administrativo (C-03) e habilidades de administração de tempo (H-02) e operação de recursos tecnológicos (H-05).

            -Output:
            Curso: Agente de Viagens
            Ano: 2019
            Código DN: 2629
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: As tecnologias indicadas para o curso "Agente de Viagens" são de grande relevância pois visam integrar avançadas ferramentas de análise de dados, inteligência artificial e IoT para otimizar a coleta de dados, a criação de roteiros e a assistência ao cliente. Este impacto é significativo na medida em que amplia a capacidade do profissional em fornecer serviços personalizados e eficientes, além de criar oportunidades para novas funções e habilidades no setor de turismo.

            #Unidade Curricular: Elaborar Produtos e Serviços Turísticos (00034-01)
            ### Sugestões de alteração dos Conhecimentos da UC:
            --00034-01.C-18.Novo.
            ---Tecnologia(s) relacionada(s): AD/Big Data Analytics
            ---Descrição da alteração: Inclusão de conhecimentos específicos sobre Big Data Analytics e seu uso para identificar tendências e demandas em serviços turísticos. Isso envolve entender como coletar, analisar e interpretar grandes volumes de dados para adaptar ofertas de produtos turísticos aos padrões de demanda.

            --00034-01.C-19.Novo.
            ---Tecnologia(s) relacionada(s): IAA/Chatbots Inteligentes
            ---Descrição da alteração: Inclusão de conhecimentos sobre o uso de chatbots inteligentes para a coleta de dados sobre destinos e atrativos turísticos. Discutir como estas ferramentas automatizam e facilitam a coleta de informações relevantes em tempo real.

            --00034-01.C-20.Novo.
            ---Tecnologia(s) relacionada(s): IOT/Sensores Inteligentes
            ---Descrição da alteração: Inclusão de conhecimentos sobre o uso de sensores inteligentes para monitoramento em tempo real da infraestrutura turística. Isso envolve entender a aplicação de sensores em coletar dados contínuos sobre a qualidade de rotas turísticas e equipamentos.

            --00034-01.C-21.Novo.
            ---Tecnologia(s) relacionada(s): APP/Plataformas de Gerenciamento de Tarefas e Projetos
            ---Descrição da alteração: Inclusão de conhecimentos sobre plataformas de gerenciamento de tarefas e projetos, especialmente em como utilizá-las para organizar e criar roteiros turísticos personalizados e eficientes apresentações de serviços turísticos.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Atualização para incluir atividades práticas envolvendo o uso de Big Data Analytics, chatbots inteligentes e sensores inteligentes. Por exemplo, propondo tarefas onde os alunos devem utilizar essas tecnologias para coletar e analisar dados, monitorar infraestruturas em tempo real, e criar roteiros personalizados usando plataformas de gerenciamento de projetos. Estas atividades podem incluir oficinas, simulações e projetos colaborativos.

            #Unidade Curricular: Comercializar Produtos e Serviços Turísticos (00034-02)
            ### Sugestões de alteração dos Conhecimentos da UC:
            --00034-02.C-15.Novo.
            ---Tecnologia(s) relacionada(s): IAA/Sistemas de Recomendação
            ---Descrição da alteração: Inclusão de conhecimentos sobre sistemas de recomendação baseados em IA. Isso incluiria entender como os algoritmos de recomendação podem sugerir produtos e serviços personalizados aos clientes com base no histórico de compras e preferências.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Atualização para incluir simulações e treinamentos práticos no uso de sistemas de recomendação. Isso poderia envolver o uso de softwares específicos que permitam aos alunos simular a interação com clientes e a sugestão de produtos personalizados com base em dados históricos. Este treinamento pode ser integrado em atividades de marketing direto e automatizado.

            #Unidade Curricular: Assessorar o Viajante (00034-03)
            ### Sugestões de alteração dos Conhecimentos da UC:
            --00034-03.C-05.Novo.
            ---Tecnologia(s) relacionada(s): IAA/Assistentes Virtuais
            ---Descrição da alteração: Inclusão de conhecimentos sobre assistentes virtuais e seu papel em oferecer assistência personalizada e imediata aos viajantes. Abrange a compreensão de como essas ferramentas podem melhorar a comunicação contínua e a resposta rápida às necessidades dos clientes.

            --00034-03.C-06.Novo.
            ---Tecnologia(s) relacionada(s): APP/Plataformas de Comunicação e Colaboração Online
            ---Descrição da alteração: Inclusão de conhecimentos sobre plataformas de comunicação e colaboração online para a gestão e acompanhamento de demandas dos clientes em tempo real durante a viagem. Discute-se como estas plataformas podem facilitar a organização e prestar melhor assistência ao viajante.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Atualização para incluir simulações práticas do uso de assistentes virtuais e plataformas de comunicação online para prestar suporte ao cliente durante a viagem. Isso poderia implicar em atividades como responder a demandas e resolver problemas em tempo real com a ajuda dessas tecnologias. Simulados de atendimento ao cliente e casos de estudo para resolução de conflitos poderiam ser incorporados.

            -Comentário geral sobre as sugestões de alteração do curso: As sugestões de atualização do curso proporcionam um alinhamento claro com as tecnologias emergentes, preparando os alunos para atuarem eficientemente num mercado de trabalho cada vez mais digital e orientado por dados. A inclusão dessas tecnologias não requer necessariamente recursos avançados em todas as unidades do Senac, pois muitos dos conhecimentos podem ser transmitidos teórica e conceitualmente. Além disso, estas atualizações aumentam a competitividade do egresso, fornecendo habilidades relevantes que se adaptam às novas exigências do setor de viagens e turismo.

            Curso: Aprendizagem Profissional Técnica em Segurança do Trabalho
            Ano: 2019
            Código DN: 2528
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso:
            As tecnologias indicadas, tais como sistemas de gerenciamento de tarefas e projetos, visão computacional, plataformas de IoT, ferramentas de anotação, plataformas de EAD, e simulações interativas, têm potencial para melhorar significativamente a eficácia, eficiência e qualidade do aprendizado dos alunos. A inclusão dessas tecnologias pode ajudar a preparar os alunos para as demandas modernas do mercado de trabalho, proporcionando-lhes conhecimentos e habilidades práticas com ferramentas tecnológicas que são altamente valorizadas hoje em dia. Essas tecnologias podem automatizar tarefas manuais, aumentar a precisão das análises e diminuir os riscos associados ao trabalho, promovendo um ambiente de trabalho mais seguro e eficiente.

            #Para a UC: Elaborar, Implantar e Implementar a Política de Saúde e Segurança do Trabalho (00082-01)
            Sugestões de alteração dos Conhecimentos da UC:
            --00082-01-C-22. Ferramentas de Gerenciamento Digital: tipos e aplicações.. Novo
            ---Tecnologia(s) relacionada(s): APP, Sistemas de Gerenciamento de Tarefas e Projetos
            ---Descrição da alteração: Incluir o conhecimento de ferramentas de software para gerenciamento de tarefas e projetos de saúde e segurança do trabalho, como Trello, Asana e Microsoft Project, que podem auxiliar no planejamento, execução e monitoramento das políticas de SST.
            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Inserir atividades práticas onde os alunos possam manusear e se familiarizar com sistemas de gerenciamento de tarefas e projetos, como Trello, Asana e Microsoft Project. Isso permitirá aos alunos compreender melhor como essas ferramentas podem ser utilizadas para definir metas, responsabilidades e atualizar políticas SST de maneira eficiente.

            #Para a UC: Realizar Avaliação e Medidas de Controle de Riscos Ergonômicos e de Acidentes (00082-03)
            Sugestões de alteração dos Conhecimentos da UC:
            --00082-03-C-18. Tecnologias de Visão Computacional. Novo
            ---Tecnologia(s) relacionada(s): IAA, Visão Computacional
            ---Descrição da alteração: Incluir conhecimentos sobre tecnologias de visão computacional para avaliação de posturas e movimentos no ambiente de trabalho. A visão computacional pode identificar automaticamente riscos ergonômicos, tornando o processo mais rápido e preciso.
            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Orientar a utilização de tecnologias de visão computacional durante as aulas práticas para avaliar posturas e movimentos dos trabalhadores. Isso poderia incluir demonstrações e exemplos de ferramentas disponíveis no mercado que realizam essa análise.

            #Para a UC: Monitorar Riscos Ocupacionais (00082-04)
            Sugestões de alteração dos Conhecimentos da UC:
            --00082-04-C-07. Plataformas IoT para Monitoramento de Riscos. Novo
            ---Tecnologia(s) relacionada(s): IOT, Plataformas de Gerenciamento de IoT
            ---Descrição da alteração: Incluir conhecimento sobre o uso de plataformas IoT destinadas ao monitoramento automatizado e controle dos equipamentos de coleta de dados do ambiente de trabalho. Isso permite um monitoramento mais eficiente e detalhado dos riscos ocupacionais.
            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Realizar atividades práticas demonstrando como configurar e utilizar plataformas IoT para monitoramento dos riscos ocupacionais. Encorajar os alunos a explorarem essas ferramentas para que aprendam a fazer o controle dos equipamentos e interpretarem os dados obtidos.

            #Para a UC: Executar Ações de Investigação, Registro e Controle de Incidentes, Acidentes de Trabalho e Doenças Ocupacionais (00082-05)
            Sugestões de alteração dos Conhecimentos da UC:
            --00082-05-C-13. Ferramentas de Anotação e Organização de Conteúdos. Novo
            ---Tecnologia(s) relacionada(s): APP, Ferramentas de Anotação e Organização 
            ---Descrição da alteração: Incluir conhecimentos sobre o uso de ferramentas digitais para organização de investigações e relatórios, como Evernote, OneNote e Google Keep, que facilitam a coleta e organização de dados sobre incidentes e acidentes.
            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Introduzir atividades práticas onde os alunos utilizem ferramentas de anotação digitais para registrar informações sobre acidentes e incidentes, elaboração de relatórios e análise de riscos.

            #Para a UC: Auxiliar e Executar Ações de Elaboração de Programas de Saúde e Segurança do Trabalho (00082-06)
            Sugestões de alteração dos Conhecimentos da UC:
            --00082-06-C-04.Ferramentas de Formulários Online. Novo
            ---Tecnologia(s) relacionada(s): APP, Ferramentas de Formulários Online 
            ---Descrição da alteração: Incluir conhecimento sobre ferramentas como Google Forms ou JotForm para coleta e análise de dados relativos à avaliação e implementação de programas de SST.
            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Sugerir atividades onde os alunos desenvolvam e utilizem formulários online para coletar dados, realizar avaliações e gerenciar documentos relacionados aos programas de saúde e segurança do trabalho.

            [...]
                  
            -Comentário geral sobre as sugestões de alteração do curso:
            As alterações sugeridas visam capacitar os alunos com conhecimentos e habilidades práticas em uso de tecnologias modernas, preparando-os com ferramentas e práticas que são relevantes no mercado de trabalho atual. A inclusão dessas tecnologias no currículo é viável para os professores nas diversas unidades do Senac, considerando a acessibilidade dessas tecnologias, a maioria das quais possui versão gratuita ou de baixo custo. Os alunos sairão do curso com um diferencial competitivo, estando mais bem preparados para enfrentar as demandas e desafios das ocupações relacionadas à segurança do trabalho.

            Curso: Assistente de Pessoal
            Ano: 2014
            Código DN: 1513
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: As tecnologias de automação sugeridas para o curso de Assistente de Pessoal visam aumentar a eficiência e precisão nas atividades relacionadas à gestão de pessoal. O uso de aplicativos, plataformas digitais, IA aplicada e sensores inteligentes proporcionará uma maior automação das tarefas, reduzindo a necessidade de intervenção humana para atividades repetitivas e propensas a erros. Isso permitirá que os futuros profissionais se concentrem em atividades de maior valor agregado e estratégia, além de garantir uma execução mais rápida e eficiente das rotinas administrativas. A introdução dessas tecnologias é viável, considerando que o aluno já tenha uma base teórica sobre as ferramentas, até mesmo sem uma implantação prática integral.

            #Para a UC "Apoiar e Executar Ações Referentes às Rotinas de Admissão e Demissão de Colaboradores"

            Sugestões de alteração dos Conhecimentos da UC:
            --00202-01-C-11. Formulários Online: conceitos, tipos e importância na rotina de admissão e demissão de colaboradores. Precisão no preenchimento e armazenamento seguro de dados. Novo
            ---Tecnologia(s) relacionada(s): APP, Formulários Online
            ---Descrição da alteração: A inclusão do conhecimento sobre formulários online ajudará os alunos a compreender as vantagens e o funcionamento dessa ferramenta, essencial para a ampliação da precisão e eficiência no preenchimento e organização de formulários.

            --00202-01-C-12. Chatbots: definições e aplicações no processo de admissão e demissão, automação de tarefas rotineiras e comunicação com colaboradores. Novo
            ---Tecnologia(s) relacionada(s): IAA, Chatbots
            ---Descrição da alteração: O conhecimento sobre chatbots é crucial para que os alunos aprendam a automatizar a coleta de documentos e agendamentos de exames admissionais e demissionais, transferindo atividades que anteriormente necessitavam de interação humana para a interação máquina-consumidor.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Além das atividades práticas já listadas, os docentes devem incluir o uso de laboratórios de informática para demonstrações sobre formulários online e chatbots. Esse ambiente permitirá aos alunos experimentar a configuração e utilização dessas tecnologias. Estudos de caso sobre empresas que já adotam essas tecnologias também devem ser incluídos, promovendo debates sobre a eficácia e os desafios do uso dessas ferramentas.

            #Para a UC "Acompanhar e Controlar a Entrega de Benefícios Legais e Espontâneos Concedidos Pela Organização"

            Sugestões de alteração dos Conhecimentos da UC:
            --00202-02-C-05. Sistemas de Gerenciamento de Benefícios: conceitos, funcionalidades e importância na automação e organização dos processos de benefícios. Novo
            ---Tecnologia(s) relacionada(s): APP, Sistemas de Gerenciamento de Benefícios
            ---Descrição da alteração: Incluir conhecimento sobre sistemas de gerenciamento de benefícios permitirá aos alunos entender como essas ferramentas podem ampliar a eficiência na organização e controle dos benefícios concedidos aos colaboradores.

            --00202-02-C-06. Sensores Inteligentes: introdução e aplicações na emissão automática de alertas de prazos de concessão de benefícios. Novo
            ---Tecnologia(s) relacionada(s): IOT, Sensores Inteligentes
            ---Descrição da alteração: O conhecimento sobre sensores inteligentes ajudará os alunos a compreender como essa tecnologia facilita o acompanhamento e cumprimento de prazos para a concessão de benefícios.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Incluir atividades práticas em laboratórios de informática para demonstrar o uso de sistemas de gerenciamento de benefícios e sensores inteligentes. Além disso, os docentes devem promover discussões e análises de casos reais ou fictícios de empresas que utilizam essas ferramentas, para que os alunos possam avaliar os impactos positivos e os desafios da implementação.

            #Para a UC "Auxiliar a Elaboração da Folha de Pagamento"

            Sugestões de alteração dos Conhecimentos da UC:
            --00202-03-C-11. Assistentes Virtuais de Voz: conceitos e aplicabilidade no processo de coleta e atualização de dados cadastrais. Novo
            ---Tecnologia(s) relacionada(s): IAA, Assistente Virtual de Voz
            ---Descrição da alteração: Ao incluir esse conhecimento, permitirá que os alunos compreendam as vantagens dos assistentes virtuais no aumento da eficiência e precisão da coleta de dados e no atendimento aos clientes internos.

            --00202-03-C-12. Sistemas de Gerenciamento de Folha de Pagamento: funcionalidades, benefícios e importância na precisão e eficiência do cálculo e organização da folha de pagamento. Novo
            ---Tecnologia(s) relacionada(s): APP, Sistemas de Gerenciamento de Folha de Pagamento
            ---Descrição da alteração: Esse conhecimento ajudará os alunos a entenderem como os sistemas de gerenciamento contribuem para uma maior eficiência na elaboração da folha de pagamento e na realização de cálculos precisos.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Inclua a utilização de laboratórios de informática para o aprendizado prático de assistentes virtuais de voz e sistemas de gerenciamento de folha de pagamento. Promova atividades como a criação e manipulação de dados em sistemas de folha de pagamento e a interação com assistentes virtuais para coleta de dados, além de estudos de caso sobre a implementação dessas tecnologias em empresas reais.

            -Comentário geral sobre as sugestões de alteração do curso: As alterações sugeridas proporcionarão um aprendizado mais atualizado e relevante para os alunos, preparando-os melhor para o uso das novas tecnologias de automação presentes no mercado. Além disso, ao incluir exemplos práticos e estudos de caso nas orientações metodológicas, as aulas se tornam mais dinâmicas e aplicáveis ao contexto real de trabalho. Essas mudanças são passíveis de adoção pelos professores, tanto em unidades que dispõem de recursos tecnológicos, quanto em unidades com menos infraestrutura, por meio de simulações e estudos teóricos detalhados sobre as tecnologias sugeridas.                                   

            Curso: Cuidador de Idoso
            Ano: 2018
            Código DN: 2454
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: A inclusão das tecnologias sugeridas como assistentes virtuais de voz (IAA), sensores inteligentes (IoT) e formulários online (APP) no curso de Cuidador de Idoso trará significativas melhorias tanto na qualidade do trabalho dos cuidadores quanto na segurança, independência e autonomia dos idosos. Esses dispositivos automatizam tarefas e oferecem novas ferramentas que ampliam a eficiência e a precisão do cuidado, além de gerar novas atividades que não eram realizadas anteriormente. 

            ## Unidade Curricular: Estimular a Independência e Autonomia do Idoso em Suas Atividades de Vida Diária (00140-01)

            -Sugestões de alteração dos Conhecimentos da UC:
            --00140-01-C-21. Tecnologias assistivas para idosos: definições, tipos e aplicações de assistentes virtuais de voz, sensores inteligentes e outros dispositivos de automação doméstica. (Novo)
            ---Tecnologia(s) relacionada(s): IAA / Assistente Virtual de Voz, IoT / Sensores Inteligentes.
            ---Descrição da alteração: Adicionar este conhecimento permitirá que os alunos compreendam o uso e a importância das tecnologias assistivas modernas no apoio às atividades diárias dos idosos.

            --00140-01-C-22. Uso de assistentes virtuais de voz para estímulo de atividades de vida diária e de ocupação do tempo livre. (Novo)
            ---Tecnologia(s) relacionada(s): IAA / Assistente Virtual de Voz.
            ---Descrição da alteração: Incluir este tópico ajuda os alunos a entender como a interação com assistentes virtuais pode encorajar o autocuidado e a manutenção da autonomia dos idosos.

            --00140-01-C-23. Monitoramento ambiental e de mobilidade do idoso usando sensores inteligentes. (Novo)
            ---Tecnologia(s) relacionada(s): IoT / Sensores Inteligentes.
            ---Descrição da alteração: A inclusão deste conhecimento capacitará os cuidadores a utilizarem tecnologias de sensores para garantir a segurança e conforto dos idosos, permitindo a identificação de situações de risco de maneira mais eficiente.

            -Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Sugere-se que as orientações metodológicas incluam atividades práticas, onde os alunos podem interagir com assistentes virtuais e sensores inteligentes, aprendendo a programá-los e a integrá-los nas rotinas do cuidado diário ao idoso. Isso pode ser feito por meio de simulações, demonstrações práticas e estudos de caso que evidenciem o uso dessas tecnologias para melhorar a qualidade de vida dos idosos.

            ## Unidade Curricular: Cuidar da Pessoa Idosa em Suas Atividades de Vida Diária (00140-02)

            -Sugestões de alteração dos Conhecimentos da UC:
            --00140-02-C-15. Aplicação de tecnologias de registro e monitoramento: uso de formulários online e outros aplicativos digitais para registro de informações de saúde do idoso. (Novo)
            ---Tecnologia(s) relacionada(s): APP / Formulários Online.
            ---Descrição da alteração: Adicionar este conhecimento é essencial para que os cuidadores possam utilizar de forma eficiente as ferramentas digitais para o registro preciso e comunicação das condições de saúde dos idosos com a equipe multiprofissional e familiares.

            -Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Recomenda-se implementar atividades de aprendizagem que envolvam o uso de dispositivos móveis e formulários online para que os alunos pratiquem o registro de dados, a comunicação de alterações no estado de saúde e a utilização dessas tecnologias no dia a dia, aumentando a eficiência e a precisão das informações compartilhadas com a equipe multidisciplinar.

            -Comentário geral sobre as sugestões de alteração do curso: As sugestões de atualização do curso têm como objetivo integrar tecnologias emergentes que aumentam a eficiência do trabalho dos cuidadores e a segurança dos idosos. As alterações recomendadas proporcionam aos alunos o conhecimento necessário para utilizar novas tecnologias assistivas, como assistentes virtuais e sensores inteligentes, além de ferramentas digitais para registro e comunicação de dados. Estas alterações apoiam a formação de cuidadores de idosos mais preparados para enfrentar os desafios contemporâneos, aproveitando as inovações tecnológicas para melhorar as práticas de cuidado, independência e autonomia dos idosos.
            A implementação dessas sugestões é viável em diversas unidades do Senac, pois envolve a introdução de conceitos teóricos e práticos acessíveis, utiliza tecnologias amplamente disponíveis e promove uma cultura de inovação constante no setor de cuidados com a terceira idade.

            Curso: Desenvolvedor Front-End
            Ano: 2022
            Código DN: 2824
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: As sugestões de tecnologias como IA Generativa, Sistemas de Gerenciamento de Tarefas e Projetos, Business Intelligence (BI) e Realidade Aumentada trarão uma grande ampliação e geração de novas atividades no curso de Desenvolvedor Front-End. A automação facilitará o processo de elaboração e desenvolvimento dos projetos, bem como a organização e análise dos dados necessários. A implementação dessas tecnologias refletirá na melhoria da produtividade e na execução de atividades complexas de forma mais eficiente.

            ## UC: Elaborar Projetos de Aplicações para Web (00074-01)

            ### Sugestões de alteração dos Conhecimentos da UC:
            -Alteração 1: Técnicas de criatividade: brainstorming, mapas mentais, painéis semânticos e uso de IA Generativa para criação automática de propostas e protótipos. (00074-01-C-02). Atualizado.
            --Tecnologia relacionada: IAA, IA Generativa
            --Descrição da alteração: A inclusão da utilização de IA Generativa amplia o conhecimento técnico de técnicas de criatividade, permitindo a criação automática de propostas e protótipos a partir do briefing. Isso melhora a eficiência no desenvolvimento de projetos de aplicações web.

            -Alteração 2: Análise de mercado: identificação do perfil do cliente e público-alvo, concorrentes diretos e indiretos, e geração de insights com Business Intelligence (BI). (00074-01-C-03). Atualizado.
            --Tecnologia relacionada: AD, Business Intelligence (BI)
            --Descrição da alteração: A inclusão do uso de BI permite a geração de novos insights baseados em métricas de análise de mercado, auxiliando na definição de objetivos e elaboração da proposta de trabalho de forma mais embasada e detalhada.

            -Alteração 3: Projetos web: tendências, tecnologias, gestão de projetos (custos, calendários de tarefas e relatórios de acompanhamento), e uso de sistemas de gerenciamento de tarefas e projetos. (00074-01-C-05). Atualizado.
            --Tecnologia relacionada: APP, Sistemas de Gerenciamento de Tarefas e Projetos
            --Descrição da alteração: A inclusão de sistemas como Trello, Asana ou Monday facilita a organização e acompanhamento das etapas do projeto, ampliando a capacidade de gestão e sincronização das atividades.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            -Alteração 1: Utilização de ferramentas de IA Generativa para auxiliar na criação de propostas e protótipos durante os exercícios práticos, proporcionando aos alunos uma experiência mais real e eficiente no planejamento de um website. Incluir atividades que envolvam o uso dessas tecnologias para estimular a inovação e a criatividade.
            --Descrição da alteração: Atividades práticas devem incluir a experimentação com IA Generativa para criação automática de protótipos e propostas, além do uso de sistemas de gerenciamento de tarefas e projetos para a organização das atividades.

            -Alteração 2: Incluir a análise de mercado através de BI nas situações reais de mercado propostas, permitindo que os alunos aprendam a gerar insights valiosos a partir de métricas e dados. 
            --Descrição da alteração: Incorporar exercícios que utilizem BI para análise de mercado, promovendo uma abordagem prática e detalhada na geração de insights e identificação de perfis de clientes e público-alvo.

            ## UC: Desenvolver Aplicações para Websites (00074-02)

            ### Sugestões de alteração dos Conhecimentos da UC:
            -Alteração 1: Imagem digital: Conceitos de vetor e bitmap, formatos e aplicações, unidades de medida, densidade de pixels, taxa de bits, animações web, e uso de IA Generativa para criação de imagens, layouts e animações otimizadas. (00074-02-C-01). Atualizado.
            --Tecnologia relacionada: IAA, IA Generativa
            --Descrição da alteração: A inclusão de IA Generativa permite a rápida criação e otimização de imagens, layouts e animações, substituindo a criação manual e acelerando o processo de desenvolvimento.

            -Alteração 2: Projeto de website: características funcionais, usabilidade, acessibilidade e ergonomia, e uso de plataformas de comunicação e colaboração online. (00074-02-C-05). Atualizado.
            --Tecnologia relacionada: APP, Plataformas de Comunicação e Colaboração Online
            --Descrição da alteração: A inclusão de ferramentas como Slack, Trello e Asana facilita a comunicação e organização das etapas do projeto, aumentando a eficiência através da colaboração entre os membros do projeto.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            -Alteração 1: Incorporar o uso de IA Generativa na criação de imagens, layouts e animações durante as atividades práticas de desenvolvimento de websites. Estimular os alunos a utilizar a IA para otimizar processos criativos e incrementar o aprendizado sobre comunicação visual e imagem digital. 
            --Descrição da alteração: Atividades práticas devem incluir a experimentação com IA Generativa para a criação e otimização de imagens e layouts, focando na substituição da criação manual por processos mais automatizados e eficientes.

            -Alteração 2: Promover o uso de plataformas de comunicação e colaboração online para planejar e organizar as etapas do projeto, facilitando a execução das atividades e a comunicação eficiente entre os alunos. 
            --Descrição da alteração: Incorporar o uso de ferramentas como Slack, Trello e Asana nas atividades práticas para organizar, planejar e acompanhar as etapas do projeto, estimulando a colaboração entre os alunos.

            ## UC: Codificar Front-End de Aplicações Web (00074-03)

            ### Sugestões de alteração dos Conhecimentos da UC:
            -Alteração 1: Usabilidade e acessibilidade: princípios aplicados ao comportamento dinâmico da página, com assistência de IA Generativa para criação de códigos responsivos e dinâmicos. (00074-03-C-08). Atualizado.
            --Tecnologia relacionada: IAA, IA Generativa
            --Descrição da alteração: Incluir o uso de IA Generativa para facilitar a programação de comportamentos dinâmicos e responsivos, melhorando a acessibilidade e usabilidade das páginas web de forma automatizada.

            -Alteração 2: Document Object Model (DOM): objetos, propriedades e eventos; manipulação de elementos, atribuição de eventos e estilos dinâmicos, com uso de Realidade Aumentada (AR) para visualização e teste dos códigos. (00074-03-C-06). Atualizado.
            --Tecnologia relacionada: RE, Realidade Aumentada (AR)
            --Descrição da alteração: Incluir o uso de AR para visualizar e testar os elementos programados no contexto real, permitindo uma maior precisão e qualidade na programação.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            -Alteração 1: Incluir a utilização de assistentes de IA generativa durante as atividades de codificação para auxiliar na criação automatizada de scripts e comportamentos dinâmicos, estimulando o aprendizado e a eficiência na aplicação das melhores práticas de desenvolvimento de aplicações web. 
            --Descrição da alteração: Atividades práticas devem incluir IA Generativa para criar scripts e comportamentos dinâmicos, promovendo a automação do processo de codificação e aprendizado das melhores práticas de usabilidade e acessibilidade.

            -Alteração 2: Utilizar recursos de AR para permitir que os alunos visualizem e testem os códigos programados em um ambiente simulado, garantindo maior precisão e qualidade no desenvolvimento de funcionalidades dinâmicas. 
            --Descrição da alteração: Incorporar o uso de AR nas atividades de teste e visualização de elementos programados, melhorando a precisão e qualidade na programação de comportamentos dinâmicos e estendendo o aprendizado prático.

            ## UC: Publicar Aplicações Web (00074-04)

            ### Sugestões de alteração dos Conhecimentos da UC:
            -Alteração 1: Testes de desempenho: comportamento e integridade do website, usando sistemas de gerenciamento de tarefas e projetos para planejamento e organização. (00074-04-C-05). Atualizado.
            --Tecnologia relacionada: APP, Sistemas de Gerenciamento de Tarefas e Projetos
            --Descrição da alteração: A inclusão do uso de ferramentas como Trello, Asana ou Monday melhora o planejamento e a organização das atividades relacionadas aos testes de desempenho, aumentando a eficiência na definição dos serviços de hospedagem e testes de compatibilidade.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            -Alteração 1: Incluir o uso de sistemas de gerenciamento de tarefas e projetos para planejar, organizar e acompanhar as etapas de definição do serviço de hospedagem e testes de desempenho dos websites, proporcionando aos alunos uma experiência mais organizada e eficiente. 
            --Descrição da alteração: Atividades práticas devem incorporar o planejamento e organização das etapas de projeto utilizando ferramentas de gerenciamento de tarefas e projetos, facilitando a definição e testes de desempenho do serviço de hospedagem.

            -Comentário geral sobre as sugestões de alteração do curso: As atualizações sugeridas integrarão tecnologias de automação que irão ampliar significativamente a capacidade dos alunos de desenvolverem e gerenciarem projetos de aplicações web. Com a inclusão de IA Generativa, Realidade Aumentada e ferramentas de gerenciamento de tarefas e projetos, o curso se tornará mais dinâmico e alinhado com as demandas atuais do mercado. As sugestões são passíveis de adoção por professores de diversas unidades do Senac, considerando que muitas dessas tecnologias podem ser apresentadas de forma teórica e prática, dentro das possibilidades de cada unidade. Isso garantirá que os egressos estejam melhor preparados para enfrentar os desafios tecnológicos da profissão.

            Curso: Doceiro
            Ano: 2018
            Código DN: 2438
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: A incorporação de tecnologias como IoT, aplicativos de gerenciamento e IA generativa promete modernizar o curso de Doceiro, trazendo mais eficiência e inovação. Essas tecnologias afetarão o curso de modo a otimizar processos operacionais, aumentar a precisão e organização, e facilitar a criação de novos produtos. Esse impacto vai além de melhorias práticas, contribuindo também para o desenvolvimento conceitual e crítico dos alunos na manipulação de alimentos, tornando-os mais adaptáveis e preparados para o mercado de trabalho dinâmico e tecnológico.

            # Para a UC “Organizar o Ambiente de Trabalho para Produções Gastronômicas (00149-01)”

            Sugestões de alteração dos Conhecimentos da UC:
            --00149-01-C-06. Sensores inteligentes para monitoramento de higiene. <Novo>
            ---Tecnologia(s) relacionada(s): IOT/Sensores Inteligentes
            ---Descrição da alteração: Incluir o conhecimento sobre sensores inteligentes que monitoram níveis de higiene e garantem conformidade com protocolos de segurança alimentar. Essa adição complementa o conhecimento atual sobre boas práticas (00149-01-C-02) ao incorporar tecnologias avançadas de monitoramento de higiene, que são cruciais para a manutenção dos padrões de segurança e limpeza atuais.
            --00149-01-C-07. Sistemas de Gerenciamento de Tarefas e Projetos. <Novo>
            ---Tecnologia(s) relacionada(s): APP/Sistemas de Gerenciamento de Tarefas e Projetos
            ---Descrição da alteração: Adicionar o conhecimento sobre plataformas digitais que facilitam a organização das atividades operacionais e garantem a eficiência na seleção de utensílios e equipamentos. Esse conhecimento ampliará o entendimento dos alunos sobre a tecnologia aplicada à organização e estrutura do ambiente de trabalho (00149-01-C-04).

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Atualizar as orientações metodológicas para incluir atividades práticas onde os alunos possam explorar e aplicar sensores inteligentes para monitorar a higiene e utilizar sistemas de gerenciamento de tarefas para organizar o ambiente de trabalho. Isso pode incluir demonstrações práticas, estudos de caso, visitas técnicas ou palestras com profissionais do setor que já utilizam essas tecnologias. Dessa forma, os alunos não apenas conhecem teoricamente as novas ferramentas, mas também adquirem competências práticas para o uso dessas tecnologias em um ambiente real.

            # Para a UC “Controlar e Organizar Estoques em Ambientes de Manipulação de Alimentos (00149-02)”

            Sugestões de alteração dos Conhecimentos da UC:
            --00149-02-C-08. Sistemas de Gerenciamento de Estoques. <Novo>
            ---Tecnologia(s) relacionada(s): APP/Sistemas de Gerenciamento de Estoques
            ---Descrição da alteração: Incluir o conhecimento sobre sistemas de gerenciamento de estoques que auxiliam no controle preciso das entradas e saídas de estoque além da reposição de produtos. Este conhecimento se alinha com a operação do estoque (00149-02-C-05) e melhora as boas práticas no recebimento e armazenamento de mercadorias (00149-02-C-06).
            --00149-02-C-09. Sensores Inteligentes e Ferramentas de Análise Predictiva. <Novo>
            ---Tecnologia(s) relacionada(s): IOT/Sensores Inteligentes, AD/Ferramentas de Análise Predictiva
            ---Descrição da alteração: Adicionar o conhecimento sobre sensores inteligentes que monitoram níveis de estoque e ferramentas de análise preditiva que antecipam necessidades de estoque com base em dados históricos. Esses conhecimentos ampliam a capacidade dos alunos de utilizar tecnologias avançadas para manter a precisão e eficiência na gestão de estoques (00149-02-C-05).

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Atualizar as orientações metodológicas para incluir atividades que permitam aos alunos utilizar sistemas de gerenciamento de estoques, sensores inteligentes e ferramentas de análise preditiva em simulações de controle e organização de estoques. Isso pode incluir a criação de cenários de simulação e estudos de caso que demonstrem o uso dessas tecnologias, além de visitas técnicas a empresas que utilizem esses sistemas. A abordagem prática garantirá que os alunos tenham uma experiência real e aplicável no uso dessas ferramentas, proporcionando um aprendizado ativo e contextualizado.

            # Para a UC “Produzir e Comercializar Doces (00149-03)”

            Sugestões de alteração dos Conhecimentos da UC:
            --00149-03-C-12. IA Generativa para Criação de Receitas. <Novo>
            ---Tecnologia(s) relacionada(s): IAA/IA Generativa
            ---Descrição da alteração: Incluir o conhecimento sobre a utilização de IA generativa para criar novas receitas baseadas em ingredientes disponíveis e tendências do mercado. Este conhecimento amplia as capacidades criativas dos alunos e se integra aos conhecimentos sobre ingredientes aplicados à doçaria (00149-03-C-03).
            --00149-03-C-13. Sistemas de E-commerce. <Novo>
            ---Tecnologia(s) relacionada(s): APP/Sistemas de E-commerce
            ---Descrição da alteração: Adicionar o conhecimento sobre plataformas de e-commerce que facilitam a venda direta dos doces ao consumidor final. Esse conhecimento é crucial para a comercialização e precificação dos doces (00149-03-C-10), proporcionando aos alunos uma compreensão das possibilidades de mercado no ambiente digital.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Atualizar as orientações metodológicas para incluir atividades que permitam aos alunos experimentar o uso de IA generativa para a criação de receitas e utilizar sistemas de e-commerce para comercializar seus produtos. Isso pode envolver simulações, workshops com profissionais que utilizam essas tecnologias no dia a dia e estudos de caso sobre a integração dessas tecnologias na produção e comercialização de doces. Essas atividades práticas não só fornecem conhecimento tecnológico, mas também desenvolvem a capacidade de aplicar novas ferramentas em um cenário de produção real.

            -Comentário geral sobre as sugestões de alteração do curso: As sugestões propostas visam modernizar e otimizar o curso de Doceiro, tornando os alunos aptos a utilizar tecnologias de ponta que são cada vez mais presentes no mercado de trabalho. A inclusão de sensores inteligentes para monitoramento de higiene, sistemas de gerenciamento de estoques, IA generativa, e plataformas de e-commerce, garante que os egressos estejam melhor preparados para enfrentar e aproveitar as novas dinâmicas do setor de alimentação. Essas mudanças são implementáveis em diversas unidades do Senac dada a flexibilidade nas orientações metodológicas, que oferecem várias alternativas práticas, como simulações e estudos de caso. Assim, adaptações podem ser realizadas conforme os recursos disponíveis em cada localidade, assegurando que o conhecimento teórico e prático chegue de maneira eficiente e acessível a todos os estudantes.

            Curso: Florista
            Ano: 2023
            Código DN: 2942
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: As tecnologias indicadas, como IoT, AR, IA Generativa, sensores inteligentes e plataformas de gerenciamento de tarefas, devem proporcionar ampliação e geração de novas funcionalidades no trabalho do florista. Elas permitem maior controle e precisão no processo de organização e criação de produções florais, além de facilitarem a comunicação, a visualização do resultado final antes da execução, e a conservação dos arranjos florais. Isso resultará em floristas mais eficientes, capazes de utilizar ferramentas tecnológicas modernas para melhor gerenciar seus estoques, maximizar a vida útil das plantas, e atender melhor seus clientes com apresentações mais realistas e sugestões personalizadas.

            ## UC 1: Organizar o Ambiente de Trabalho do Florista (00016-01)

            ### Sugestões de alteração dos Conhecimentos da UC:
            -- 00016-01-C-11. Uso de sensores inteligentes para monitoramento ambiental. Novo.
            --- Tecnologia(s) relacionada(s): IOT/Sensores Inteligentes.
            --- Descrição da alteração: Incluir conhecimento sobre o uso de sensores inteligentes para monitorar condições ambientais como temperatura e umidade, ajudando na conservação e maximização da vida útil das plantas. A introdução dessa tecnologia permite aos floristas controlar o ambiente de maneira mais eficiente, garantindo melhores condições para as plantas.

            -- 00016-01-C-12. Sistemas de Gerenciamento de Estoques Automatizados. Novo.
            --- Tecnologia(s) relacionada(s): APP/Sistemas de Gerenciamento de Estoques.
            --- Descrição da alteração: Adicionar conhecimento sobre sistemas de gerenciamento de estoques, ensinando aos alunos como usar plataformas para controlar e monitorar o estoque de insumos botânicos e materiais diversos de maneira mais precisa e organizada.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            --- Descrição da alteração: Atualizar as orientações metodológicas para incluir a demonstração prática do uso de sensores inteligentes para monitoramento ambiental e sistemas automatizados de gerenciamento de estoques. Pode-se planejar atividades em que os alunos configurem sensores para monitorar condições de temperatura e umidade e usem software de gestão de estoques, simulações nas aulas práticas, ou estudos de caso, onde verifiquem o impacto no controle de insumos e na vida útil das plantas.

            ## UC 2: Elaborar Produções Florais (00016-02)

            ### Sugestões de alteração dos Conhecimentos da UC:
            -- 00016-02-C-26. Utilização de Realidade Aumentada (AR) na Visualização de Produções Florais. Novo.
            --- Tecnologia(s) relacionada(s): RE/Realidade Aumentada (AR).
            --- Descrição da alteração: Incluir conhecimento sobre o uso de tecnologias de Realidade Aumentada para visualizar produções florais no ambiente real antes da sua execução, ajudando na apresentação ao cliente e na adaptação do design ao contexto real.

            -- 00016-02-C-27. Aplicação de IA Generativa no Design Floral. Novo.
            --- Tecnologia(s) relacionada(s): IAA/IA Generativa.
            --- Descrição da alteração: Adicionar conhecimento sobre IA Generativa e suas aplicações no design floral, especialmente na criação de arranjos e combinações inovadoras de flores e materiais baseadas em preferências e tendências.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            --- Descrição da alteração: Atualizar as orientações metodológicas para incorporar o uso de tecnologias de Realidade Aumentada e IA Generativa em atividades práticas. Por exemplo, os alunos podem usar aplicativos de AR para visualizar suas criações em ambientes reais e programas de IA Generativa para explorar diferentes combinações de flores e tendências de design.

            ## UC 3: Planejar e Executar Projetos de Decoração Floral para Eventos (00016-03)

            ### Sugestões de alteração dos Conhecimentos da UC:
            -- 00016-03-C-19. Uso de Realidade Aumentada (AR) no Planejamento e Execução de Decorações Florais. Novo.
            --- Tecnologia(s) relacionada(s): RE/Realidade Aumentada (AR).
            --- Descrição da alteração: Incluir conhecimento sobre o uso de AR para visualizar e planejar decorações florais no ambiente real antes da implementação final, aumentando a precisão e a adaptabilidade do layout.

            -- 00016-03-C-20. Aplicação de Sensores Inteligentes na Conservação de Arranjos Durante Eventos. Novo.
            --- Tecnologia(s) relacionada(s): IOT/Sensores Inteligentes.
            --- Descrição da alteração: Adicionar conhecimento sobre o uso de sensores inteligentes para monitorar e regular condições ambientais durante eventos, maximizando a longevidade e a qualidade dos arranjos.

            -- 00016-03-C-21. Sistemas de Gerenciamento de Tarefas e Projetos na Decoração Floral. Novo.
            --- Tecnologia(s) relacionada(s): APP/Sistemas de Gerenciamento de Tarefas e Projetos.
            --- Descrição da alteração: Incluir conhecimento sobre o uso de plataformas para a criação e gerenciamento de cronogramas, aumentando a eficiência na elaboração e execução de projetos de decoração floral.

            ### Sugestões de alteração das Orientações Metodológicas da UC:
            --- Descrição da alteração: Atualizar as orientações metodológicas para incorporar o uso de AR e sensores inteligentes em atividades práticas. Recomenda-se a realização de atividades onde os alunos utilizem AR para visualizar e planejar suas decorações e configurem sensores para monitorar condições ambientais durante eventos para conservar melhor seus arranjos. Pode-se também incluir o uso de sistemas de gerenciamento de tarefas e projetos em aulas práticas para ensinar aos alunos como planejar, organizar e gestionar eficientemente projetos de decoração floral.

            -Comentário geral sobre as sugestões de alteração do curso: As sugestões de alterações no curso visam tornar os egressos mais preparados para o mercado de trabalho atual, onde a tecnologia tem um papel significativo. As adições propõem contextos práticos para o uso de tecnologias que estão ganhando popularidade e eficiência, como IoT, AR e IA Generativa, permitindo que os floristas melhorem a qualidade de seu trabalho e suas interações com os clientes. As sugestões são práticas e podem ser adaptadas pelos professores em diversas unidades do Senac, utilizando recursos que vão desde simulações e estudos teóricos até visitas técnicas e a prática direta com ferramentas de baixo custo ou disponíveis em versões gratuitas.

            Curso: Porteiro e Vigia
            Ano: 2016
            Código DN: 2198
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: As tecnologias indicadas têm um grande potencial para transformar o curso de Porteiro e Vigia, trazendo maior eficiência e precisão nas atividades de controle de acesso, monitoramento e comunicação. Os impactos esperados envolvem a ampliação e substituição de atividades manuais por sistemas automatizados, além da geração de novas atividades relacionadas à análise de dados e comunicação automatizada. Essas mudanças exigirão que os alunos adquiram novos conhecimentos e desenvolvam habilidades para operar essas tecnologias, tornando-os mais preparados para atuar em um mercado de trabalho cada vez mais tecnológico e dinâmico.

            #Para a UC Executar Atividades do Serviço de Portaria (00089-01)

            Sugestões de alteração dos Conhecimentos da UC:
            --00089-01-C-13.Editores de texto, internet, correio eletrônico e plataformas de comunicação online.Atualizado
            ---Tecnologia(s) relacionada(s): APP, Plataformas de Comunicação e Colaboração Online
            ---Descrição da alteração: A inclusão das plataformas de comunicação online no currículo é essencial para que os alunos aprendam a utilizar ferramentas digitais que facilitam a troca de informações e a comunicação em tempo real. Isso se alinha com a necessidade de ampliar a eficiência nas atividades de fornecimento de informações do serviço e na passagem de turno.

            --00089-01-C-10.Procedimentos de acesso de moradores, visitantes, fornecedores, prestadores de serviço, autoridades e de veículos: identificação, registro e autorização. Atualizado
            ---Tecnologia(s) relacionada(s): IOT, Sensores Inteligentes
            ---Descrição da alteração: É importante atualizar o conteúdo relacionado aos procedimentos de acesso para incluir o uso de sensores inteligentes. Essa tecnologia substituirá parte do controle e registro manual de acesso, tornando o processo mais automatizado e eficiente.

            --00089-01-C-19.Procedimento de passagem de turno e uso de plataformas de comunicação online. Atualizado
            ---Tecnologia(s) relacionada(s): APP, Plataformas de Comunicação e Colaboração Online
            ---Descrição da alteração: A inclusão do uso de plataformas de comunicação online melhora a precisão e eficiência na passagem de turno, permitindo a troca de informações em tempo real entre as equipes.

            --00089-01-C-13.Editores de texto, internet, correio eletrônico e ferramentas de visualização de dados. Atualizado
            ---Tecnologia(s) relacionada(s): AD, Data Visualization
            ---Descrição da alteração: Incluir ferramentas de visualização de dados no currículo permitirá que os alunos aprendam a transformar dados brutos em informações valiosas, facilitando a criação de relatórios semi-automatizados para um melhor controle e análise das atividades.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: A metodologia deve incluir atividades práticas que envolvam o uso de plataformas de comunicação e colaboração online, além de sistemas de sensores inteligentes. Propor cenários reais e simulações que integrem esses recursos tecnológicos permitirá que os alunos desenvolvam as habilidades necessárias para operar essas ferramentas. Também se recomenda a análise de casos e a criação de relatórios usando ferramentas de visualização de dados para fortalecer a compreensão e aplicação prática das novas tecnologias.

            #Para a UC Realizar Medidas Preventivas de Segurança Pessoal e Patrimonial (00089-02)

            Sugestões de alteração dos Conhecimentos da UC:
            --00089-02-C-05.CFTV, botão de pânico, infravermelho ativo, concertina, cerca elétrica, sirene, rádio, interfone, videofone, controle remoto para abrir portas e portões e sensores inteligentes. Atualizado
            ---Tecnologia(s) relacionada(s): IOT, Sensores Inteligentes
            ---Descrição da alteração: Incluir os sensores inteligentes nos conhecimentos sobre equipamentos de segurança é imprescindível para que os alunos entendam como esses dispositivos podem substituir parcialmente a vigilância manual e ampliar a capacidade de monitoramento contínuo e imediato.

            --00089-02-C-03.Atuação do porteiro e vigia nas atividades preventivas de segurança patrimonial e física, incluindo a comunicação automatizada. Atualizado
            ---Tecnologia(s) relacionada(s): APP, Formulários Online
            ---Descrição da alteração: A comunicação automatizada, facilitada por formulários online, deve ser incorporada ao conteúdo para que os alunos saibam como utilizar essas ferramentas para registrar ocorrências e solicitar serviços.

            --00089-02-C-06.Instrumentos de controle nas atividades preventivas de segurança: livro de ocorrências e ferramentas de análise de dados. Atualizado
            ---Tecnologia(s) relacionada(s): AD, Análise de Dados para Segurança Preventiva
            ---Descrição da alteração: Acrescentar ferramentas de análise de dados ao conhecimento sobre instrumentos de controle permitirá que os alunos identifiquem padrões de segurança e áreas de risco, gerando novas atividades de análise com base nos dados coletados.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Recomenda-se que a metodologia inclua atividades práticas utilizando sensores inteligentes para monitoramento de áreas comuns e controle de acesso. Além disso, devem ser conduzidas atividades que envolvam a análise de dados coletados, possibilitando a identificação de padrões de segurança e áreas de risco. A utilização de formulários online para a comunicação e registro de ocorrências deve ser integrada nas atividades cotidianas simuladas, proporcionando experiências de aprendizado realistas e contextualizadas.

            -Comentário geral sobre as sugestões de alteração do curso: As sugestões de atualização do curso visam preparar melhor os alunos para um mercado de trabalho em constante evolução tecnológica, especialmente na área de segurança. As tecnologias indicadas permitem uma maior eficiência, precisão e modernização das atividades de portaria e vigilância. As alterações propostas são passíveis de adoção por professores em diversas unidades do Senac, considerando que muitas tecnologias podem ser abordadas de maneira teórica e, quando possível, prática, proporcionado um ensino de qualidade em todas as unidades, independentemente dos recursos disponíveis.

            Curso: Técnico em Secretaria Escolar
            Ano: 2021
            Código DN: 2905
            -Comentário geral sobre a avaliação do impacto da automação sobre o curso: As tecnologias indicadas para o curso Técnico em Secretaria Escolar prometem ampliação, geração e substituição de atividades através de aplicativos e plataformas digitais, inteligência artificial e soluções de automação, promovendo maior eficiência nos processos administrativos, na comunicação assertiva e na gestão de documentos. Isso vai preparar os alunos para um ambiente de trabalho mais tecnológico e dinâmico, alinhando as competências com as demandas atuais do mercado de trabalho.

            # Unidade Curricular: Prestar Atendimento na Secretaria Escolar (00143-01)

            Sugestões de alteração dos Conhecimentos da UC:
            --00143-01-C-09.Técnicas e plataformas de comunicação digital.Novo
            ---Tecnologia(s) relacionada(s): APP/Plataformas de Comunicação e Colaboração Online
            ---Descrição da alteração: Incluir este conhecimento para capacitar os alunos a utilizar diversas plataformas de comunicação digital eficazmente, considerando a necessidade de troca de informações em tempo real e a precisão das informações transmitidas, impactando positivamente nos indicadores de competência.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Além das simulações de atendimentos, incluir atividades práticas que utilizem plataformas de comunicação digital. Proponha a criação de cenários onde os alunos possam praticar a orientação, disponibilização de informações e mediação de conflitos utilizando essas plataformas. Isso vai abordar a necessidade de interação e resolução de problemas em um ambiente digital, refletindo as atualizações feitas nos conhecimentos.

            # Unidade Curricular: Realizar Atividades de Apoio Aos Processos Administrativo-Pedagógicos de Secretaria Escolar (00143-02)

            Sugestões de alteração dos Conhecimentos da UC:
            --00143-02-C-19.Sistemas de Gestão Eletrônica de Documentos: conceitos e prática.Novo
            ---Tecnologia(s) relacionada(s): APP/Sistemas de Gestão Eletrônica de Documentos (GED)
            ---Descrição da alteração: Incluir este conhecimento para capacitar os alunos na utilização de sistemas de GEDs, que facilitam a organização, protocolo, arquivamento e acesso a documentos físicos e eletrônicos, aumentando a eficiência e precisão na gestão dos documentos.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Incluir atividades que usam sistemas de Gestão Eletrônica de Documentos (GED) em laboratórios de informática. Proponha simulações práticas que envolvam o uso do GED para arquivamento, protocolo e organização de documentos, fornecendo aos alunos práticas reais da utilização dessa tecnologia.

            # Unidade Curricular: Coletar, Interpretar e Monitorar Dados Estatísticos da Instituição Educacional (00143-03)

            Sugestões de alteração dos Conhecimentos da UC:
            --00143-03-C-05.Business Intelligence: conceitos e ferramentas.Novo
            ---Tecnologia(s) relacionada(s): AD/Business Intelligence (BI)
            ---Descrição da alteração: Incluir este conhecimento para familiarizar os alunos com ferramentas de BI que melhoram a coleta, interpretação e análise de dados estatísticos da instituição educacional, gerando novas atividades automatizadas e melhor aplicando estatísticas.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Adicionar atividades práticas que envolvam o uso de ferramentas de Business Intelligence (BI) para a coleta e análise de dados educacionais. Promover o uso de plataformas online para treinamento sobre essas ferramentas, ampliando as competências dos alunos no uso das tecnologias de análise de dados.

            # Unidade Curricular: Organizar os Processos Legais da Escola Perante e os Órgãos Reguladores (00143-04)

            Sugestões de alteração dos Conhecimentos da UC:
            --00143-04-C-11.Sistemas de Gestão de Prazo e Fluxo: conceitos e aplicações.Novo
            ---Tecnologia(s) relacionada(s): APP/Sistemas de Gestão de Prazo e Fluxo
            ---Descrição da alteração: Incluir este conhecimento para familiarizar os alunos com sistemas que automatizam o controle de prazos de documentos de regularização da instituição, ampliando a precisão e eficiência no monitoramento dos prazos de validade dos documentos.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Incluir atividades práticas que utilizem sistemas de Gestão de Prazo e Fluxo em laboratórios de informática, propondo a criação de cenários para monitoramento e atualização de documentos de regularização da instituição. Isso vai permitir que os alunos apliquem os conhecimentos teóricos em um ambiente de simulação prática.

            # Unidade Curricular: Elaborar, Organizar e Controlar Documentos da Instituição Educacional e da Vida Escolar do Aluno (00143-05)

            Sugestões de alteração dos Conhecimentos da UC:
            --00143-05-C-09.Sistemas de Gestão Eletrônica de Documentos (GED): conceitos, benefícios e aplicabilidade.Atualizado
            ---Tecnologia(s) relacionada(s): APP/Sistemas de Gestão Eletrônica de Documentos (GED)
            ---Descrição da alteração: Atualizar este conhecimento para incluir a aplicabilidade prática dos GEDs, focando nos conceitos e benefícios para melhorar a organização e o controle dos documentos da vida escolar do aluno.

            --00143-05-C-14.Comunicação escrita e tecnologias de transcrição: regras gramaticais, ortografia, acentuação, concordância verbal e nominal, pronome de tratamento, estrutura textual, coesão e coerência.Atualizado
            ---Tecnologia(s) relacionada(s): IAA/Sistemas de Transcrição de Reuniões
            ---Descrição da alteração: Atualizar este conhecimento para incluir o uso de sistemas de transcrição de reuniões, que automatizam a redação de atas durante reuniões, facilitando a elaboração de documentos.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Incluir atividades práticas que utilizem sistemas de Gestão Eletrônica de Documentos (GED) para arquivamento, controle e atualização de documentos escolares, e sistemas de transcrição para elaboração de atas. Isso vai permitir que os alunos apliquem as teorias em um ambiente de simulação, aprimorando a precisão e a eficiência nas tarefas documentais.

            # Unidade Curricular: Apoiar e Executar As Ações de Avaliação e Controle de Processos, Atividades e Recursos Materiais da Secretaria Escolar (00143-06)

            Sugestões de alteração dos Conhecimentos da UC:
            --00143-06-C-08.Sistemas de Gerenciamento de Tarefas e Projetos: conceitos e aplicação.Novo
            ---Tecnologia(s) relacionada(s): APP/Sistemas de Gerenciamento de Tarefas e Projetos
            ---Descrição da alteração: Incluir este conhecimento para capacitar os alunos na utilização de sistemas que facilitam a administração das atividades de trabalho e a comunicação interna, ampliando a eficiência na gestão de processos e na organização das tarefas diárias.

            Sugestões de alteração das Orientações Metodológicas da UC:
            ---Descrição da alteração: Incluir atividades práticas que utilizem sistemas de Gerenciamento de Tarefas e Projetos, propondo situações-problema e simulações que permitam a organização e controle das atividades de trabalho e a comunicação interna eficiente. Isso irá capacitar os alunos na utilização prática dessa tecnologia, refletindo na eficiência da rotina de trabalho.

            [...]
                  
            -Comentário geral sobre as sugestões de alteração do curso: As sugestões de alteração no curso Técnico em Secretaria Escolar permitirão que os alunos ganhem mais familiaridade com tecnologias atuais que ampliam a eficiência e precisão das tarefas administrativas, melhorando a comunicação, gestão de documentos, análise de dados, controle de recursos e integração com ferramentas digitais e plataformas colaborativas. Essas atualizações são pertinentes e factíveis, considerando os diferentes contextos e recursos disponíveis nas variadas unidades do Senac espalhadas pelo Brasil.                                              

            Respire fundo e pense passo a passo.
            Você receberá uma gorjeta de US$1.000,00 por essa atividade.           
            """
)

# Carregar os dados do arquivo Excel
file_path_pcn = r'C:\Users\yuri.lima\Servico Nacional de Aprendizagem Comercial\GerProspecAvalEducacional - 08 - Prospectiva do Trabalho\2 - Impacto da Automação\PCN_14.10.24.xlsx'
df_pcn = pd.read_excel(file_path_pcn, sheet_name='Descritores das UCs')
file_path_aval = r"C:\Users\yuri.lima\Servico Nacional de Aprendizagem Comercial\GerProspecAvalEducacional - 08 - Prospectiva do Trabalho\2 - Impacto da Automação\Avaliação.xlsx"
df_aval = pd.read_excel(file_path_aval, sheet_name='Sheet1')

# Removendo espaços em branco dos nomes das colunas
df_aval.columns = df_aval.columns.str.strip()

# Função para criar a coluna 'impacto' concatenando as colunas de impacto
def criar_coluna_impacto(df):
    impactos = []
    for _, row in df.iterrows():
        impacto_list = []
        if row['substituição'] == 'x':
            impacto_list.append('substituição')
        if row['geração'] == 'x':
            impacto_list.append('geração')
        if row['ampliação'] == 'x':
            impacto_list.append('ampliação')
        if row['transferência'] == 'x':
            impacto_list.append('transferência')
        impactos.append(", ".join(impacto_list))
    df['impacto'] = impactos

# Criar a coluna 'impacto'
criar_coluna_impacto(df_aval)

# Agrupar os dados por curso
grouped_pcn = df_pcn.groupby(['id', 'Código DN', 'Título do Curso'])

# Função para montar o prompt
def montar_prompt_curso(curso_info, grupo_pcn, grupo_aval):
    nome_curso = f"{curso_info['Título do Curso']} ({curso_info['Código DN']}) - {curso_info['Ano']}"
    prompt = f"Curso: {nome_curso}\n"

    ucs = grupo_pcn.groupby(['id_UC', 'Título UC'])
    
    for (id_uc, titulo_uc), subgrupo in ucs:
        if titulo_uc.startswith("Projeto Integrador") or titulo_uc.startswith("Estágio Profissional") or titulo_uc.startswith("Prática Profissional") or titulo_uc.startswith("Prática Integrada"):
            continue
        
        prompt += f"\nTítulo da UC: {titulo_uc} ({id_uc})\n"
        
        for _, row in subgrupo.iterrows():
            if row['Tipo'] == 'Indicadores':
                prompt += f"Indicador: {row['Descrição']} ({row['id_D']})\n"
            elif row['Tipo'] == 'Conhecimentos':
                prompt += f"Conhecimento: {row['Descrição']} ({row['id_D']})\n"
            elif row['Tipo'] == 'Habilidades':
                prompt += f"Habilidade: {row['Descrição']} ({row['id_D']})\n"
            elif row['Tipo'] == 'Orientações metodológicas':
                prompt += f"Orientações metodológicas: {row['Descrição']} ({row['id_D']})\n"
        
        # Filtrando as UCs na planilha de Avaliação
        subgrupo_aval = grupo_aval[grupo_aval['id_UC'] == id_uc]
        
        # Adicionando concatenado se a TEC tiver horizonte 0 ou 1
        tecnologias = []
        for _, row in subgrupo_aval.iterrows():
            if row['horizonte'] in [0, 1]:
                concatenado = f"{row['cat_tec']}, {row['tec']}, {row['justificativa1']}, {row['impacto']}, {row['justificativa2']}, {row['relação']}, {row['justificativa4']}"
                tecnologias.append(concatenado)
        if tecnologias:
            prompt += "Tecnologias sugeridas:\n" + "\n".join(tecnologias) + "\n"

    return prompt

# Função para salvar a resposta em um arquivo .docx
def salvar_resposta_como_docx(titulo_curso, ano, codigo_dn, response_text, output_dir):
    nome_arquivo = os.path.join(output_dir, f"{titulo_curso} ({ano} - {codigo_dn}).docx")
    doc = Document()
    doc.add_heading(titulo_curso, 0)
    doc.add_paragraph(f"Ano: {ano}")
    doc.add_paragraph(f"Código DN: {codigo_dn}")    
    paragraph=doc.add_paragraph(response_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(nome_arquivo)

# Diretório de saída para os arquivos .docx
output_dir = r'C:\Users\yuri.lima\Downloads\Orientações2'

# Iterar sobre os cursos
for (curso_id, codigo_dn, titulo_curso), grupo_pcn in grouped_pcn:
    curso_info = {
        'id': curso_id,
        'Código DN': codigo_dn,
        'Título do Curso': titulo_curso,
        'Ano': grupo_pcn['Ano'].iloc[0]  # Assumindo que a coluna 'Ano' existe e é consistente dentro de cada grupo
    }
    
    # Filtrando os dados do grupo na planilha de Avaliação
    grupo_aval = df_aval[df_aval['id_UC'].isin(grupo_pcn['id_UC'].unique())]
    
    # Montando o prompt
    prompt = montar_prompt_curso(curso_info, grupo_pcn, grupo_aval)
    print(prompt)

    response_json = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": SYSTEM_MESSAGE
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        model="gpt-4o-OpenAI-Prospeccao"
    )

    response_text = response_json.choices[0].message.content
    # Salvar a resposta como arquivo .docx
    salvar_resposta_como_docx(titulo_curso, curso_info['Ano'], codigo_dn, response_text, output_dir)

    # Esperar 10 segundos antes de continuar para a próxima iteração
    time.sleep(60)